# -*- coding: utf-8 -*-
"""Embed figures into the final report - fixed version"""
import sys, os
sys.path.insert(0, r'E:\xicha gis 智能定位')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

BASE = r'E:\xicha gis 智能定位'

# Map: caption_text -> image_path
FIGURES = [
    ('（图1：综合可达性分布图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'v2_real_data', 'p8_fig3_spatial.png')),
    ('（图2：街景障碍物分布图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'v2_real_data', 'p8_fig7_saii_analysis.png')),
    ('（图3：三类幻觉维度示意图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig1_framework.png')),
    ('（图5：研究框架技术路线图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig1_framework.png')),
    ('（图6：时间贫困指数空间聚类图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig6_time_poverty.png')),
    ('（图7：综合可达性与幻觉指数双变量地图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'v2_real_data', 'p8_fig3_spatial.png')),
    ('（图8：四象限散点图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'conference_paper', 'figures', 'fig4_illusion_scatter.png')),
    ('（图9：弱势群体剥夺热点图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'conference_paper', 'figures', 'fig6_deprived_communities.png')),
    ('（图10：步行环境四维评分雷达图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig5_radar.png')),
    ('（图11：供需匹配三维框架图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'conference_paper', 'figures', 'fig9_supply_demand.png')),
    ('（图12：时间贫困与幻觉指数相关性图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig6_time_poverty.png')),
    ('（图13：昼夜达标率对比图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'conference_paper', 'figures', 'fig8_day_night.png')),
    ('（图14：四区障碍评分对比箱线图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'v2_real_data', 'p8_fig7_saii_analysis.png')),
    ('（图15：四区对比机制路径图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'conference_paper', 'figures', 'fig5_type_analysis.png')),
    ('（图16：治理建议实施路径图）', os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig1_framework.png')),
]

# Namespaces
NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing/inline/1'
PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingDrawing/2006/main'

def make_inline_pic(image_path, rId, width_inches=5.5):
    ext_cx = int(width_inches * 914400)
    ext_cy = int(width_inches * 914400 * 0.6)

    drawing = etree.Element(f'{{{W_NS}}}drawing')
    inline = etree.SubElement(drawing, f'{{{WP_NS}}}inline', attrib={
        f'{{{WP_NS}}}distT': '0',
        f'{{{WP_NS}}}distB': '0',
        f'{{{WP_NS}}}distL': '0',
        f'{{{WP_NS}}}distR': '0',
    })
    etree.SubElement(inline, f'{{{WP_NS}}}extent', attrib={'cx': str(ext_cx), 'cy': str(ext_cy)})
    etree.SubElement(inline, f'{{{WP_NS}}}docPr', attrib={'id': '1', 'name': os.path.basename(image_path)})
    etree.SubElement(inline, f'{{{WP_NS}}}cNvGraphicFramePr')

    graphic = etree.SubElement(inline, f'{{{NS}}}graphic')
    etree.SubElement(graphic, f'{{{NS}}}graphicData', attrib={'uri': f'{{{PIC_NS}}}picture'})

    pic = etree.SubElement(graphic, f'{{{PIC_NS}}}pic')
    nvPicPr = etree.SubElement(pic, f'{{{PIC_NS}}}nvPicPr')
    etree.SubElement(nvPicPr, f'{{{PIC_NS}}}cNvPr', attrib={'id': '1', 'name': os.path.basename(image_path)})
    cNvPicPr = etree.SubElement(nvPicPr, f'{{{PIC_NS}}}cNvPicPr')
    etree.SubElement(cNvPicPr, f'{{{NS}}}picLocks', attrib={'noChangeAspect': '1'})

    blipFill = etree.SubElement(pic, f'{{{PIC_NS}}}blipFill')
    etree.SubElement(blipFill, f'{{{NS}}}blip', attrib={f'{{{R_NS}}}embed': rId})
    stretch = etree.SubElement(blipFill, f'{{{NS}}}stretch')
    etree.SubElement(stretch, f'{{{NS}}}fillRect', attrib={'l': '0', 't': '0', 'r': '0', 'b': '0'})

    spPr = etree.SubElement(pic, f'{{{PIC_NS}}}spPr', attrib={'bwMode': 'auto'})
    xfrm = etree.SubElement(spPr, f'{{{NS}}}xfrm')
    etree.SubElement(xfrm, f'{{{NS}}}off', attrib={'x': '0', 'y': '0'})
    etree.SubElement(xfrm, f'{{{NS}}}ext', attrib={'cx': str(ext_cx), 'cy': str(ext_cy)})
    prstGeom = etree.SubElement(spPr, f'{{{NS}}}prstGeom', attrib={'prst': 'rect'})
    etree.SubElement(prstGeom, f'{{{NS}}}avLst')

    return drawing

def embed_figures_in_doc(input_path, output_path):
    doc = Document(input_path)
    img_counter = [0]

    for i, para in enumerate(doc.paragraphs):
        text = para.text
        for caption, img_path in FIGURES:
            if caption in text:
                if not os.path.exists(img_path):
                    print(f"  [SKIP] Not found: {img_path}")
                    continue

                img_counter[0] += 1
                rId = doc.part.relate_to(
                    open(img_path, 'rb').read(),
                    'image/png',
                    part_name=f'/word/media/fig_{img_counter[0]}.png'
                )

                drawing = make_inline_pic(img_path, rId, 5.5)

                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run._r.append(drawing)

                cap_num = caption.replace('（', '').replace('）', '')
                cap_para = doc.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_para.add_run(f'图 {cap_num}')
                cap_run.font.name = '宋体'
                cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                cap_run.font.size = Pt(10)
                cap_run.font.italic = True

                print(f"  [OK] [{img_counter[0]}] {caption} <- {os.path.basename(img_path)}")

    print(f"\nTotal figures embedded: {img_counter[0]}")
    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == '__main__':
    input_doc = os.path.join(BASE, '报告_final.docx')
    output_doc = os.path.join(BASE, '报告_final_含图.docx')
    embed_figures_in_doc(input_doc, output_doc)
