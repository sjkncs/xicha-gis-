# -*- coding: utf-8 -*-
"""Embed figures into the final report - simplified reliable approach"""
import sys, os
sys.path.insert(0, r'E:\xicha gis 智能定位')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

BASE = r'E:\xicha gis 智能定位'

# Map figure captions to actual file paths
FIGURE_MAP = {
    '（图1：综合可达性分布图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'v2_real_data', 'p8_fig3_spatial.png'),
    '（图2：街景障碍物分布图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'v2_real_data', 'p8_fig7_saii_analysis.png'),
    '（图3：三类幻觉维度示意图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig1_framework.png'),
    '（图5：研究框架技术路线图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig1_framework.png'),
    '（图6：时间贫困指数空间聚类图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig6_time_poverty.png'),
    '（图7：综合可达性与幻觉指数双变量地图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'v2_real_data', 'p8_fig3_spatial.png'),
    '（图8：四象限散点图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'conference_paper', 'figures', 'fig4_illusion_scatter.png'),
    '（图9：街道障碍物空间分布图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'v2_real_data', 'p8_fig7_saii_analysis.png'),
    '（图10：弱势群体剥夺热点图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'conference_paper', 'figures', 'fig6_deprived_communities.png'),
    '（图11：步行环境四维评分雷达图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig5_radar.png'),
    '（图12：供需匹配三维框架图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'conference_paper', 'figures', 'fig9_supply_demand.png'),
    '（图13：时间贫困与幻觉指数相关性图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig6_time_poverty.png'),
    '（图14：昼夜达标率对比图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'conference_paper', 'figures', 'fig8_day_night.png'),
    '（图15：四区障碍评分对比箱线图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'v2_real_data', 'p8_fig7_saii_analysis.png'),
    '（图16：四区障碍评分对比箱线图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'v2_real_data', 'p8_fig7_saii_analysis.png'),
    '（图17：四区对比机制路径图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'conference_paper', 'figures', 'fig5_type_analysis.png'),
    '（图18：四区对比机制路径图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'conference_paper', 'figures', 'fig5_type_analysis.png'),
    '（图19：治理建议实施路径图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig1_framework.png'),
    '（图20：治理建议框架图）': os.path.join(BASE, 'projects', '15min-urban-accessibility', 'paper', 'figures', 'fig1_framework.png'),
}

def make_inline_pic(image_path, width_inches=5.5):
    """Create an inline picture XML element using lxml with proper namespaces"""
    ext_cx = int(width_inches * 914400)
    ext_cy = int(width_inches * 914400 * 0.6)

    nsmap = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing/inline/1',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    }

    # Build from scratch using lxml with namespace
    NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing/inline/1'
    PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingDrawing/2006/main'

    drawing = etree.Element(f'{{{W_NS}}}drawing')

    inline = etree.SubElement(drawing, f'{{{WP_NS}}}inline', attrib={
        f'{{{WP_NS}}}distT': '0',
        f'{{{WP_NS}}}distB': '0',
        f'{{{WP_NS}}}distL': '0',
        f'{{{WP_NS}}}distR': '0',
    })

    extent = etree.SubElement(inline, f'{{{WP_NS}}}extent', attrib={
        'cx': str(ext_cx),
        'cy': str(ext_cy),
    })

    docPr = etree.SubElement(inline, f'{{{WP_NS}}}docPr', attrib={
        'id': '1',
        'name': os.path.basename(image_path),
    })

    cNvGraphicFramePr = etree.SubElement(inline, f'{{{WP_NS}}}cNvGraphicFramePr')

    graphic = etree.SubElement(inline, f'{{{NS}}}graphic')
    graphicData = etree.SubElement(graphic, f'{{{NS}}}graphicData', attrib={
        'uri': f'{{{PIC_NS}}}picture'
    })

    pic = etree.SubElement(graphicData, f'{{{PIC_NS}}}pic')

    nvPicPr = etree.SubElement(pic, f'{{{PIC_NS}}}nvPicPr')
    cNvPr = etree.SubElement(nvPicPr, f'{{{PIC_NS}}}cNvPr', attrib={
        'id': '1',
        'name': os.path.basename(image_path),
    })
    cNvPicPr = etree.SubElement(nvPicPr, f'{{{PIC_NS}}}cNvPicPr')
    picLocks = etree.SubElement(cNvPicPr, f'{{{NS}}}picLocks', attrib={
        'noChangeAspect': '1',
    })

    blipFill = etree.SubElement(pic, f'{{{PIC_NS}}}blipFill')
    blip = etree.SubElement(blipFill, f'{{{NS}}}blip', attrib={
        f'{{{R_NS}}}embed': 'rId999',
    })
    stretch = etree.SubElement(blipFill, f'{{{NS}}}stretch')
    fillRect = etree.SubElement(stretch, f'{{{NS}}}fillRect', attrib={
        'l': '0', 't': '0', 'r': '0', 'b': '0',
    })

    spPr = etree.SubElement(pic, f'{{{PIC_NS}}}spPr', attrib={
        'bwMode': 'auto',
    })
    xfrm = etree.SubElement(spPr, f'{{{NS}}}xfrm')
    off = etree.SubElement(xfrm, f'{{{NS}}}off', attrib={'x': '0', 'y': '0'})
    ext = etree.SubElement(xfrm, f'{{{NS}}}ext', attrib={
        'cx': str(ext_cx), 'cy': str(ext_cy),
    })
    prstGeom = etree.SubElement(spPr, f'{{{NS}}}prstGeom', attrib={
        'prst': 'rect',
    })
    avLst = etree.SubElement(prstGeom, f'{{{NS}}}avLst')

    return drawing

def add_image_after_paragraph(doc, paragraph, image_path, caption, width_inches=5.5):
    """Add image + caption paragraph after a given paragraph"""
    if not os.path.exists(image_path):
        print(f"  [SKIP] Not found: {image_path}")
        return False

    # Add image to document part first
    with open(image_path, 'rb') as f:
        img_bytes = f.read()

    img_part_name = f'/word/media/fig_{len(doc.part.rels)}.png'
    rId = doc.part.relate_to(
        img_bytes,
        'image/png' if image_path.lower().endswith('.png') else 'image/jpeg',
        part_name=img_part_name
    )

    # Create paragraph with the image
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = img_para.add_run()
    drawing = make_inline_pic(image_path, width_inches)
    # Fix the rId
    for blip_el in drawing.iter(f'{{{R_NS}}}blip'):
        blip_el.set(f'{{{R_NS}}}embed', rId)
        break

    run._r.append(drawing)

    # Add caption paragraph
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.font.name = '宋体'
    cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    cap_run.font.size = Pt(10.5)
    cap_run.font.italic = True

    print(f"  [OK] {caption} <- {os.path.basename(image_path)}")
    return True

# Namespaces
NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

def embed_figures_in_doc(input_path, output_path):
    doc = Document(input_path)

    matched = 0
    skipped = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text
        for caption, img_path in FIGURE_MAP.items():
            if caption in text:
                print(f"Found [{i}]: {caption}")
                if os.path.exists(img_path):
                    # Add image + caption as new paragraphs at the end
                    # We track where to insert them
                    para_idx = i

                    # Read image and add to doc
                    with open(img_path, 'rb') as f:
                        img_bytes = f.read()

                    rId = doc.part.relate_to(
                        img_bytes,
                        'image/png',
                        part_name=f'/word/media/fig_{rId}_{matched}.png'
                    )

                    # Create image paragraph
                    drawing = make_inline_pic(img_path, 5.5)
                    # Fix rId
                    for blip_el in drawing.iter(f'{{{R_NS}}}blip'):
                        blip_el.set(f'{{{R_NS}}}embed', rId)
                        break

                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_para.add_run()
                    run._r.append(drawing)

                    # Caption
                    cap_num = caption.replace('（', '').replace('）', '')
                    cap_para = doc.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap_para.add_run(f'图 {cap_num}')
                    cap_run.font.name = '宋体'
                    cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    cap_run.font.size = Pt(10)
                    cap_run.font.italic = True

                    matched += 1
                    print(f"  [OK] {caption} <- {os.path.basename(img_path)}")
                else:
                    print(f"  [SKIP] Not found: {img_path}")
                    skipped += 1

    print(f"\nSummary: {matched} embedded, {skipped} skipped")
    doc.save(output_path)
    print(f"Saved to: {output_path}")

if __name__ == '__main__':
    input_doc = os.path.join(BASE, '报告_final.docx')
    output_doc = os.path.join(BASE, '报告_final_含图.docx')
    embed_figures_in_doc(input_doc, output_doc)
