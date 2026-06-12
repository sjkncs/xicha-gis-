# -*- coding: utf-8 -*-
"""
生成大白话版项目进度与思路介绍报告（DOCX格式）
保密程度：可公开发表报告水平
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, bold=False, indent=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_divider(doc):
    p = doc.add_paragraph("-" * 60)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(180, 180, 180)

def set_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_cell(cell, hex_color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)

# ================================================================
# 封面 / 标题
# ================================================================

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(24)
p_title.paragraph_format.space_after  = Pt(6)
run = p_title.add_run("AI 智能定位研究项目")
set_font(run, size=22, bold=True, color=(30, 80, 140))

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(4)
run2 = p_sub.add_run("——用大白话讲清楚我们做了什么、怎么做的，未来要做什么")
set_font(run2, size=13, color=(80, 80, 80))

p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_note.paragraph_format.space_after = Pt(18)
run3 = p_note.add_run("（公开发表版 · 对外保密仅至报告可发表水平）")
set_font(run3, size=10, color=(150, 150, 150))
run3.italic = True

add_divider(doc)

# ================================================================
# 一、项目概览
# ================================================================

add_heading(doc, "一、这个项目到底是干什么的？", level=1, color=(30, 80, 140))

add_para(doc,
    "简单来说：我们研究的是一件每个城市居民都会遇到的事情——\n"
    "'明明地图上显示我周围设施很齐全，为什么实际生活中却总觉得不方便？'\n\n"
    "举个例子：政府说'15分钟生活圈'已经覆盖你家门口，但你去超市、看病、接孩子、散步，\n"
    "真的能在15分钟内搞定吗？老年人晚上想散步，小区门口的路灯有没有？\n"
    "上班族周末想去公园，公交地铁能及时到吗？\n\n"
    "我们就是想搞清楚：城市的'统计数据'和真实生活的差距到底有多大。",
    size=11)

# ================================================================
# 二、核心发现
# ================================================================

add_heading(doc, "二、我们发现了什么有意思的事情？", level=1, color=(30, 80, 140))

add_para(doc, "我们通过对深圳市南山区的实际数据进行分析，发现了几个让大家意外的结果：", size=11)

add_bullet(doc, "统计数字很好看，但真实体验差很远")
add_para(doc,
    "政府公布的设施覆盖率数据看起来不错，但当我们把真实的人口分布、\n"
    "夜间营业状态、老年人出行习惯等实际因素加进去之后，发现很多人\n"
    "实际上并不能真正享受'15分钟生活圈'的便利。",
    indent=True, size=10)

add_bullet(doc, "城中村和高端小区的差距，比想象中大得多")
add_para(doc,
    "深圳南山科技园附近，城中村居民和高档小区的居民，\n"
    "晚上能获得的公共服务差距非常明显。城中村周边设施看起来很密集，\n"
    "但晚上70%的设施都关门了，实际可达性很低。",
    indent=True, size=10)

add_bullet(doc, "白天'天堂'，晚上'荒漠'——白天和晚上的可达性差异巨大")
add_para(doc,
    "很多地方白天设施齐全、应有尽有，但到了晚上——\n"
    "商店关门、公交减班、路灯昏暗，道路封闭——突然就变得很不方便。\n"
    "我们专门给这种现象起了个名字，叫'时空可达性幻觉'。",
    indent=True, size=10)

add_bullet(doc, "越弱势的群体，受影响越大")
add_para(doc,
    "老年人、残障人士、住在城中村的低收入人群，\n"
    "受到'夜间可达性下降'的影响最为严重。\n"
    "因为他们白天出行能力本就受限，晚上更是难上加难。",
    indent=True, size=10)

# ================================================================
# 三、方法论
# ================================================================

add_heading(doc, "三、我们是怎么做这件事的？", level=1, color=(30, 80, 140))

add_para(doc, "我们用了一套看起来复杂、但本质上很直观的方法：", size=11)

add_heading(doc, "第一步：把真实世界'数字化'", level=2, color=(50, 110, 170))
add_para(doc,
    "我们拿到了深圳市南山区的真实数据：\n"
    "  - 真实常住人口数据：184万人，精确到社区级别\n"
    "  - 真实路网数据：2.7万条道路，包括断头路、围墙封堵等实际障碍\n"
    "  - 6大类公共服务设施：教育、医疗、商业、交通、公园绿地、公共服务\n\n"
    "有了这些数据，我们就能在电脑里'重建'一个虚拟的南山。")

add_heading(doc, "第二步：算出每个地方'白天'和'晚上'的可达性", level=2, color=(50, 110, 170))
add_para(doc,
    "我们用GIS（地理信息系统）和网络分析方法，\n"
    "计算从任意一个居民点出发，到达周边各类设施需要多长时间。\n\n"
    "关键创新：我们不只是算白天的可达性，还专门研究了夜间的情况。\n"
    "因为设施的营业时间、灯光状态，道路开放情况，在白天和晚上差异巨大。")

add_heading(doc, "第三步：发明几个新指标，把'幻觉'量化出来", level=2, color=(50, 110, 170))
add_para(doc,
    "我们发明了一套指标体系，能把'统计数据好看但真实体验差'的现象，\n"
    "用数字的方式呈现出来：\n\n"
    "  - SAI（空间可达性指数）：衡量白天设施有多容易到达\n"
    "  - TPI（时间贫困指数）：衡量白天和晚上可达性差距有多大\n"
    "  - SAII（时空可达性幻觉指数）：综合以上两个，专门识别那些\n"
    "    '看着挺好，实际坑人'的地方\n\n"
    "借助这些指标，我们可以把南山区的每个角落标注上颜色——\n"
    "绿色代表'真实便利'，红色代表'数据好看但实际坑人'。")

add_heading(doc, "第四步：把弱势群体单独拎出来分析", level=2, color=(50, 110, 170))
add_para(doc,
    "整体数据会掩盖弱势群体的真实困境。\n"
    "因此我们把老年人、残障人士、城中村居民等群体单独建模，\n"
    "看看他们在同样的地理空间里，实际感受到的可达性是什么样的。\n"
    "结果发现：弱势群体面对的'幻觉'比平均值严重得多。")

# ================================================================
# 四、当前进度
# ================================================================

add_heading(doc, "四、目前做到什么程度了？", level=1, color=(30, 80, 140))

add_para(doc, "整个项目分成了多个模块并行推进，目前进展如下：", size=11)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
headers = ["模块名称", "核心内容", "当前状态", "备注"]
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    for para in hdr_cells[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
    shade_cell(hdr_cells[i], '1E50BC')

rows_data = [
    ("深圳真实数据分析",
     "184万人口 + 2.8万条路网 + 六类设施\n真实可达性计算",
     "已完成", "v2版本"),
    ("时空可达性幻觉指标体系",
     "SAI / TPI / SAII 三个核心指标\n四象限分类矩阵",
     "已完成", "可复现"),
    ("街道画像分析",
     "55条街道的详细可达性画像\n含障碍物评分",
     "已完成", "含夜间场景"),
    ("学术论文撰写",
     "SCI期刊论文（CE US）\n完整研究报告",
     "进行中", "投稿准备阶段"),
    ("配套PPT与演讲材料",
     "会议论文幻灯片\n答辩材料",
     "进行中", "已完成大部分"),
    ("AI Agent 研究工具",
     "科研自动化工作流\n办公自动化方法论",
     "持续迭代", "方法论沉淀"),
]

for row_data in rows_data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        for para in row_cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)

doc.add_paragraph()

# ================================================================
# 五、AI协作场景（真实对话）
# ================================================================

add_heading(doc, "五、AI是怎么帮我做这些事的？", level=1, color=(30, 80, 140))

add_para(doc,
    "这一部分不讲额外感想，只还原项目当时是怎么一步步推进出来的。\n"
    "为了避免透露过多技术细节，下面统一用与AI的问答场景来呈现，\n"
    "重点展示当时每一步在解决什么问题、产出了什么结果。",
    size=11)

# ---- 场景一 ----
add_heading(doc, "场景一：先把原始结果整理成能看懂的图", level=2, color=(50, 110, 170))

p_tag = doc.add_paragraph()
r_tag = p_tag.add_run("[ 发送给豆包的提示词 ]")
set_font(r_tag, size=10, bold=True, color=(0, 100, 180))
p_tag.paragraph_format.space_after = Pt(2)

p_prompt = doc.add_paragraph()
p_prompt.paragraph_format.left_indent = Cm(0.5)
p_prompt.paragraph_format.space_after = Pt(8)
r_p = p_prompt.add_run(
    "我现在手里已经有一批南山区可达性结果表，字段包括社区名称、白天步行时间、\n"
    "夜间步行时间、障碍物评分、可达性差值。\n"
    "请不要解释理论，直接帮我做这件事：\n"
    "第一，把这些字段整理成适合出图的结构；\n"
    "第二，告诉我应该先画哪几张图，才能把白天、夜间和差异讲清楚；\n"
    "第三，给我一版可以直接运行的绘图代码，图要适合放进报告和PPT。")
set_font(r_p, size=10)

p_tag2 = doc.add_paragraph()
r_tag2 = p_tag2.add_run("[ 豆包的回复（节选） ]")
set_font(r_tag2, size=10, bold=True, color=(160, 80, 0))
p_tag2.paragraph_format.space_after = Pt(2)

add_para(doc,
    "可以按'社区—白天—夜间—差值—障碍物评分'这条线来组织。\n"
    "如果是为了先把结果讲清楚，建议优先做三类图：\n"
    "一类是白天与夜间对比图，用来说明同一社区在两个时段的差别；\n"
    "一类是差值分布图，用来找出哪些地方夜间下降最明显；\n"
    "一类是障碍物与时间差散点图，用来说明街道条件和可达性变化之间的关系。\n"
    "下面我给你可直接运行的代码框架，并把图例、配色和导出参数一起配好。",
    indent=True, size=10)

add_para(doc,
    "当时这一步解决的是'先把结果变成图'的问题。\n"
    "图一出来，哪些社区差距最大、哪些变量最值得继续追，就一下子清楚了。",
    size=10, bold=True)

# ---- 场景二 ----
add_heading(doc, "场景二：再把零散结果整理成完整汇报顺序", level=2, color=(50, 110, 170))

p_tag3 = doc.add_paragraph()
r_tag3 = p_tag3.add_run("[ 发送给通义千问的提示词 ]")
set_font(r_tag3, size=10, bold=True, color=(0, 100, 180))
p_tag3.paragraph_format.space_after = Pt(2)

p_prompt2 = doc.add_paragraph()
p_prompt2.paragraph_format.left_indent = Cm(0.5)
p_prompt2.paragraph_format.space_after = Pt(8)
r_p2 = p_prompt2.add_run(
    "我现在已经有图表和初步结论了，但内容还是散的。\n"
    "请帮我按'研究背景—研究设计—研究区域与数据—研究内容—研究结论'\n"
    "这个结构，整理成一套汇报顺序。\n"
    "不要写空话，每一部分都告诉我应该放什么内容、先讲什么、后讲什么，\n"
    "让我能直接照着搭PPT。")
set_font(r_p2, size=10)

p_tag4 = doc.add_paragraph()
r_tag4 = p_tag4.add_run("[ 通义千问的回复（节选） ]")
set_font(r_tag4, size=10, bold=True, color=(160, 80, 0))
p_tag4.paragraph_format.space_after = Pt(2)

add_para(doc,
    "可以先用研究背景解释一个问题：为什么图上看起来近，不等于真实走起来方便。\n"
    "然后在研究设计部分交代，你是怎么把'统计可达性'和'真实体验'放到同一个框架里比较的。\n"
    "研究区域与数据部分只保留最必要的信息，说明对象是谁、数据来自哪里。\n"
    "研究内容部分按结果展开：先看综合可达性，再看可达性幻觉，再看弱势群体和昼夜差异。\n"
    "最后用研究结论收束，把发现和规划启示对应起来。",
    indent=True, size=10)

add_para(doc,
    "这一步解决的是'怎么把已有材料讲顺'的问题。\n"
    "原本分散在文档、表格和图片里的内容，到了这里才真正串成了完整故事线。",
    size=10, bold=True)

# ---- 场景三 ----
add_heading(doc, "场景三：最后把已经写出来的内容改成能直接对外讲的版本", level=2, color=(50, 110, 170))

p_tag5 = doc.add_paragraph()
r_tag5 = p_tag5.add_run("[ 发送给豆包的提示词 ]")
set_font(r_tag5, size=10, bold=True, color=(0, 100, 180))
p_tag5.paragraph_format.space_after = Pt(2)

p_prompt3 = doc.add_paragraph()
p_prompt3.paragraph_format.left_indent = Cm(0.5)
p_prompt3.paragraph_format.space_after = Pt(8)
r_p3 = p_prompt3.add_run(
    "我已经写出一版比较学术的结果描述了，但现在要做公开汇报。\n"
    "请帮我把下面这段话改成两种版本：\n"
    "一种适合放在PPT页面上，短一点、清楚一点；\n"
    "一种适合答辩时口头讲，保留意思，但让普通人也能一下听懂。\n"
    "注意不要改动结论本身，也不要写得太夸张。")
set_font(r_p3, size=10)

p_tag6 = doc.add_paragraph()
r_tag6 = p_tag6.add_run("[ 豆包的回复（节选） ]")
set_font(r_tag6, size=10, bold=True, color=(160, 80, 0))
p_tag6.paragraph_format.space_after = Pt(2)

add_para(doc,
    "PPT版可以写成：'同样是15分钟生活圈，不同社区居民的真实步行体验差异很大。\n"
    "图纸上看似达标的区域，在现实中可能因为道路障碍、绕行和夜间环境问题而明显失效。'\n\n"
    "口头版可以写成：'简单说就是，地图上看起来很近，不代表人真的能轻松走到。\n"
    "尤其到了晚上，这种差距会更明显。'",
    indent=True, size=10)

add_para(doc,
    "这一步解决的是'怎么把结果讲出去'的问题。\n"
    "前面已经把数据做出来、图表排出来、结构梳理出来，最后要做的，\n"
    "就是把表达改成对外能直接使用的版本。",
    size=10, bold=True)

add_para(doc,
    "如果把整个过程按顺序还原，其实就是三步：\n"
    "先整理结果和出图，再梳理整体结构，最后把内容改成适合展示和汇报的表达。\n"
    "AI在这里起到的作用，不是代替研究本身，而是把每一步推进得更快、更顺。",
    size=11)

# ================================================================
# 六、为什么重要
# ================================================================

add_heading(doc, "六、为什么这件事重要？", level=1, color=(30, 80, 140))

add_para(doc,
    "这项工作的重点，不是多讲道理，而是把一个原本说不清的问题真正做成可观察、可比较、可展示的结果。\n"
    "当时之所以要继续往下做，是因为前面的分析已经说明：\n"
    "同样写成'15分钟可达'，不同社区居民面对的真实出行条件可能完全不是一回事。",
    size=11)

add_para(doc,
    "所以这一部分真正重要的，不是额外延伸出多少心得，而是把'图上看起来能到'和'现实里到底好不好到'之间的差别做实。\n"
    "只有把这一步做出来，后面的图表、汇报和结论才有意义。",
    size=10)

# ================================================================
# 七、未来计划
# ================================================================

add_heading(doc, "七、未来打算做什么？", level=1, color=(30, 80, 140))

add_para(doc,
    "按照当时项目推进的顺序，走到这里其实已经把核心结果做出来了。\n"
    "后面如果继续往下推进，主要也是沿着同一条线继续补充：\n"
    "把更多社区纳入比较，把更多时段补进来，再把结果整理成更稳定的展示材料。",
    size=11)

add_para(doc,
    "也就是说，后续工作不是另起炉灶，而是在已经完成的这套流程上继续扩展范围、补充样本和细化表达。",
    size=10)

# ================================================================
# 八、研究意义
# ================================================================

add_heading(doc, "八、这项研究的意义是什么？", level=1, color=(30, 80, 140))

add_para(doc,
    "如果只从这份公开报告来看，这项工作的意义可以概括成一句话：\n"
    "它把一个原本停留在抽象表述里的问题，变成了可以一步步做出来、一步步展示出来的成果。",
    size=11)

add_para(doc,
    "前面已经完成了结果整理、图表呈现、结构梳理和公开表达。\n"
    "也正因为这样，'15分钟可达'这件事才不再只是一个口号，而变成了能被比较、能被说明、也能继续往下推进的研究对象。",
    size=10)

add_divider(doc)

p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_end.paragraph_format.space_before = Pt(12)
run_end = p_end.add_run("感谢阅读！如有问题，欢迎交流。")
set_font(run_end, size=11, color=(100, 100, 100))
run_end.italic = True

# ---- 保存 ----
output_path = r"E:\xicha gis 智能定位\projects\AI智能定位项目_大白话进度报告_公开发布版.docx"
doc.save(output_path)
print(f"Done: {output_path}")
