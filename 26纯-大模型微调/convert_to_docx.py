# -*- coding: utf-8 -*-
"""Convert Markdown report to formatted DOCX preserving structure and styling."""

import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── colour palette ──────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)   # deep-blue for headings
MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)   # mid-blue for sub-headings
ACCENT     = RGBColor(0xCC, 0x33, 0x33)   # red accent for titles
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)  # table-header background
BLACK      = RGBColor(0x00, 0x00, 0x00)

# ── helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, fill_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

# ── helper: set cell text colour ───────────────────────────────────────────────
def set_cell_color(cell, rgb: RGBColor):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = rgb

# ── helper: set cell bold ───────────────────────────────────────────────────────
def set_cell_bold(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True

# ── helper: add a horizontal rule ─────────────────────────────────────────────
def add_hr(doc: Document):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'),  '1')
    bot.set(qn('w:color'), '2E74B5')
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

# ── helper: apply heading style with colour ───────────────────────────────────
def add_heading(doc: Document, text: str, level: int = 1):
    """level 1 = title (centre, large, red), 2 = section (left, mid-blue)"""
    p = doc.add_heading('', level=level)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    if level == 0:                      # document title
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = ACCENT
    elif level == 1:                    # section headings
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = DARK_BLUE
    elif level == 2:                    # sub-section headings
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = MID_BLUE
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = BLACK
    return p

# ── helper: styled paragraph ──────────────────────────────────────────────────
def add_body(doc: Document, text: str, bold=False, italic=False,
             indent=False, colour: RGBColor = BLACK, size: int = 11,
             space_before=0, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size       = Pt(size)
    run.font.bold       = bold
    run.font.italic     = italic
    run.font.color.rgb  = colour
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

# ── helper: bullet list item ────────────────────────────────────────────────────
def add_bullet(doc: Document, text: str, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name  = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size      = Pt(11)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(3)
    return p

# ── helper: code block ─────────────────────────────────────────────────────────
def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(4)
    # light-grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F0F0F0')
    pPr.append(shd)
    run = p.add_run(code)
    run.font.name  = 'Consolas'
    run.font.size   = Pt(9)
    run.font.color.rgb = RGBColor(0x31, 0x31, 0x31)
    return p

# ── helper: inline code run (bold marker) ───────────────────────────────────────
def add_inline_code(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = 'Consolas'
    run.font.size   = Pt(10)
    run.font.color.rgb = RGBColor(0xC7, 0x15, 0x28)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

# ── helper: parse & add table ─────────────────────────────────────────────────
def add_table(doc: Document, md_table: str):
    rows_text = [ln.strip() for ln in md_table.strip().split('\n') if ln.strip()]
    if len(rows_text) < 2:
        return

    def split_row(line: str):
        line = line.strip().strip('|')
        cells = [c.strip() for c in line.split('|')]
        return [c for c in cells
                if c and c not in ('---', '--', ':--', '--:', ':---', '---:')
                and not re.match(r'^[-: ]+$', c)]

    header_cells = split_row(rows_text[0])
    n_cols = len(header_cells)
    if n_cols == 0:
        return

    # count actual data rows (skip separator line at index 1)
    data_rows = [rows_text[i] for i in range(1, len(rows_text))]

    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.style = 'Table Grid'

    # header
    for c_idx, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[c_idx]
        cell.text = cell_text
        shade_cell(cell, 'D6E4F0')
        set_cell_bold(cell)
        set_cell_color(cell, DARK_BLUE)
        for para in cell.paragraphs:
            if para.runs:
                para.runs[0].font.name = '微软雅黑'
                para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                para.runs[0].font.size = Pt(10.5)

    # data rows
    for r_idx, row_text in enumerate(data_rows):
        cells = split_row(row_text)
        for c_idx in range(n_cols):
            cell_text = cells[c_idx] if c_idx < len(cells) else ''
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = cell_text
            if (r_idx + 1) % 2 == 0:
                shade_cell(cell, 'F7FAFD')
            for para in cell.paragraphs:
                if para.runs:
                    para.runs[0].font.name = '宋体'
                    para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    para.runs[0].font.size = Pt(10)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(4.2) if i == 0 else Cm(5.8)

    doc.add_paragraph()

# ── stateful markdown parser ───────────────────────────────────────────────────
def build_docx(md_path: str, out_path: str):
    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    # margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── paragraph spacing defaults ────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name  = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size  = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    i = 0
    while i < len(lines):
        ln = lines[i].rstrip()

        # ── heading levels ────────────────────────────────────────────────────
        if ln.startswith('## '):
            add_hr(doc)
            add_heading(doc, ln[3:].strip(), level=1)
            doc.add_paragraph()
        elif ln.startswith('### '):
            add_heading(doc, ln[4:].strip(), level=2)
        elif ln.startswith('#### '):
            add_heading(doc, ln[5:].strip(), level=3)
        elif ln.startswith('# '):
            # document title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(ln[2:].strip())
            run.font.name  = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size  = Pt(20)
            run.font.bold  = True
            run.font.color.rgb = ACCENT
            doc.add_paragraph()

        # ── table block ──────────────────────────────────────────────────────
        elif ln.startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                tbl_lines.append(lines[i].rstrip())
                i += 1
            add_table(doc, '\n'.join(tbl_lines))
            continue

        # ── code block (```) ──────────────────────────────────────────────────
        elif ln.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            add_code_block(doc, '\n'.join(code_lines))
            i += 1
            continue

        # ── blockquote / note ─────────────────────────────────────────────────
        elif ln.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  'EEF4FB')
            pPr.append(shd)
            run = p.add_run(ln[2:])
            run.font.name  = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size  = Pt(10.5)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            p.paragraph_format.space_after = Pt(6)

        # ── bullet list ──────────────────────────────────────────────────────
        elif ln.startswith('- '):
            add_bullet(doc, ln[2:])

        # ── separator ───────────────────────────────────────────────────────
        elif ln.startswith('---'):
            add_hr(doc)

        # ── empty line ───────────────────────────────────────────────────────
        elif ln == '':
            pass   # swallow blanks

        # ── bold italic inline ───────────────────────────────────────────────
        elif ln.startswith('**思考题'):
            # render as bold heading
            p = doc.add_paragraph()
            run = p.add_run(ln)
            run.font.name  = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size  = Pt(12)
            run.font.bold  = True
            run.font.color.rgb = DARK_BLUE
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)

        # ── inline formatting within body text ────────────────────────────────
        else:
            # replace **bold**, *italic*, `code` inline
            def inline(text):
                parts = []
                # split by `code`
                segments = text.split('`')
                for idx, seg in enumerate(segments):
                    if idx % 2 == 1:          # code segment
                        r = doc.add_run(seg)
                        r.font.name  = 'Consolas'
                        r.font.size   = Pt(9.5)
                        r.font.color.rgb = RGBColor(0xC7, 0x15, 0x28)
                    else:
                        # split by **bold**
                        sub_parts = seg.split('**')
                        for s_idx, s_seg in enumerate(sub_parts):
                            if s_idx % 2 == 1:  # bold
                                r = doc.add_run(s_seg)
                                r.font.name  = '宋体'
                                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                r.font.size  = Pt(11)
                                r.font.bold  = True
                            elif s_seg:
                                r = doc.add_run(s_seg)
                                r.font.name  = '宋体'
                                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                r.font.size  = Pt(11)
                return parts

            # render as new paragraph with inline styles
            p = doc.add_paragraph()
            # parse whole line into segments
            rest = ln
            while rest:
                # find next bold marker
                bold_m = re.search(r'\*\*(.+?)\*\*', rest)
                code_m = re.search(r'`([^`]+)`', rest)
                if not bold_m and not code_m:
                    r = p.add_run(rest)
                    r.font.name  = '宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size  = Pt(11)
                    break
                # closest marker
                first = bold_m
                if code_m and (not bold_m or code_m.start() < bold_m.start()):
                    first = code_m
                # text before marker
                if first.start() > 0:
                    r = p.add_run(rest[:first.start()])
                    r.font.name  = '宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size  = Pt(11)
                # the marker content
                content = first.group(1)
                r = p.add_run(content)
                r.font.name  = 'Consolas' if first == code_m else '宋体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size  = Pt(9.5) if first == code_m else Pt(11)
                r.font.bold  = True  if first == bold_m else False
                r.font.color.rgb = RGBColor(0xC7, 0x15, 0x28) if first == code_m else BLACK
                rest = rest[first.end():]

            p.paragraph_format.space_after = Pt(4)

        i += 1

    # ── page numbers in footer ─────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        para   = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run    = para.add_run('第 ')
        run.font.size  = Pt(9)
        run.font.name  = '宋体'
        run    = para.add_run('&P')
        run.font.size  = Pt(9)
        run    = para.add_run(' 页')
        run.font.size  = Pt(9)
        run.font.name  = '宋体'

    doc.save(out_path)
    print(f'Saved: {out_path}')

if __name__ == '__main__':
    md_path = r'e:\xicha gis 智能定位\26纯-大模型微调\实验项目1_大模型搭建与使用_报告内容.md'
    out_path = r'e:\xicha gis 智能定位\26纯-大模型微调\实验项目1_大模型搭建与使用_报告内容.docx'
    build_docx(md_path, out_path)
