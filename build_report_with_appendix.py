# -*- coding: utf-8 -*-
"""Build the complete final report with all figures + full appendix"""
import sys, os
sys.path.insert(0, r'E:\xicha gis 智能定位')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = r'E:\xicha gis 智能定位'

# ── Path shortcuts ────────────────────────────────────────────────────────────────
P   = lambda *p: os.path.join(BASE, *p)
PF  = P('projects', '15min-urban-accessibility', 'paper', 'figures')
CF  = P('15分钟城市时间贫困研究', 'conference_paper', 'figures')
PP  = P('projects', '15min-urban-accessibility', 'paper')
SP  = P('projects', '15min-urban-accessibility', 'v2_real_data', 'street_profiles')
AN  = P('appendix-vlm', 'appendix_annotated', 'appendix_annotated')
AC  = P('自选年份', 'gpu_scripts', 'results', 'annotated_cn')

# ── Font helpers ────────────────────────────────────────────────────────────────
def set_run_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    try:
        if run._element.rPr is not None and run._element.rPr.rFonts is not None:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    except Exception:
        pass
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def safe_text(text):
    import re
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)

# ── Doc helpers ────────────────────────────────────────────────────────────────
def add_cover(doc, title, authors, institution, date_str, keywords):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title); set_run_font(run, '黑体', 18, bold=True)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(authors); set_run_font(run, '宋体', 14)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(institution); set_run_font(run, '宋体', 13)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(date_str); set_run_font(run, '宋体', 13)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('关键词：' + keywords); set_run_font(run, '宋体', 11)
    doc.add_page_break()

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs: set_run_font(run, '黑体', 16, bold=True)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs: set_run_font(run, '黑体', 14, bold=True)
    return p

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs: set_run_font(run, '黑体', 12, bold=True)
    return p

def add_para(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text); set_run_font(run, '宋体', 12)
    return p

def add_fig(doc, img_path, width=Cm(15), caption_text=None, two_col=False):
    if os.path.exists(img_path):
        if two_col:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=width)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=width)
    if caption_text:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption_text)
        set_run_font(r, '宋体', 9, color=(100, 100, 100))

def add_fig_grid(doc, img_paths, widths=None, caption_text=None, n_cols=2):
    """Insert multiple images in a grid layout."""
    n = len(img_paths)
    if widths is None:
        widths = [Cm(7)] * n
    row_count = (n + n_cols - 1) // n_cols
    for row in range(row_count):
        for col in range(n_cols):
            idx = row * n_cols + col
            if idx >= n:
                break
            img_path = img_paths[idx]
            w = widths[idx] if idx < len(widths) else Cm(7)
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=w)
    if caption_text:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption_text)
        set_run_font(r, '宋体', 9, color=(100, 100, 100))

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]; cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs: set_run_font(run, '宋体', 10, bold=True)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]; cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs: set_run_font(run, '宋体', 9)
    return table

def new_page(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# MAIN REPORT BUILD
# ─────────────────────────────────────────────────────────────────────────────
def build_report():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17); section.right_margin = Cm(3.17)

    # ══════════════════════════════════════════════════════════════════
    # COVER
    # ══════════════════════════════════════════════════════════════════
    add_cover(doc,
        '高密度城市15分钟生活圈中的可达性幻觉：\n基于深圳市南山区多源实证数据的成因解析与时空正义审视',
        '张潇晗  宋阳霆',
        '哈尔滨工业大学（深圳）',
        '2026年6月',
        '15分钟城市；可达性幻觉；时空贫困；路网障碍'
    )

    # ══════════════════════════════════════════════════════════════════
    # ABSTRACT
    # ══════════════════════════════════════════════════════════════════
    add_h1(doc, '摘要')
    abstract_text = (
        '"15分钟城市"作为当代可持续城市规划的核心范式，其有效性高度依赖于对步行可达性的精准度量。'
        '然而，既有研究多基于路网拓扑或欧氏距离构建统计可达性指标，系统性地忽视了微观建成环境品质、'
        '街道障碍物及社会群体异质性对实际步行体验的制约，导致规划愿景与居民真实生活体验之间产生显著断裂。'
        '我们以深圳南山区为实证对象，整合官方建筑轮廓、社区实有人口、多源兴趣点及精细化路网数据，构建了涵盖'
        '综合可达性、可达性幻觉指数、时间贫困及昼夜时空差异的多维评估框架。我们发现，南山区约38.6%的城中村位置'
        '存在显著的"可达性幻觉"，即统计上达标但实地步行体验严重受损。这种幻觉并非随机分布，而是与高密度非正规定居点、'
        '街道物理障碍及弱势群体聚居区高度空间耦合。进一步分析揭示，夜间照明缺失与安全焦虑导致弱势群体的有效生活半径'
        '较白天收缩逾40%，形成了隐蔽的"时空剥夺"。我们从空间正义视角剖析幻觉的结构性成因，并提出从"设施覆盖"'
        '转向"体验修复"的精细化治理路径，为高密度城市15分钟生活圈的包容性落地提供理论依据与实证依据。'
    )
    add_para(doc, safe_text(abstract_text))
    p = doc.add_paragraph()
    run = p.add_run('关键词：15分钟城市；可达性幻觉；时空贫困；路网障碍；深圳南山区；步行环境评估；供需匹配')
    set_run_font(run, '楷体', 11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    new_page(doc)

    # ══════════════════════════════════════════════════════════════════
    # CHAPTER 1
    # ══════════════════════════════════════════════════════════════════
    add_h1(doc, '第1章 引言：从距离达标到体验断裂的理论转向')

    p1_intro = (
        '自"15分钟城市"概念提出以来，该理念已在全球范围内被广泛采纳为衡量城市宜居性与可持续性的关键标尺。'
        '在深圳市南山区这一集高新技术产业园区、高端商品房社区与大量城中村于一体的典型区域，'
        '图纸上800米的服务半径，在现实中可能意味着穿越无信号灯的主干道、挤占被电动车侵占的狭窄巷道、'
        '或在夜间因照明不足而被迫绕行的漫长旅程。'
        '本研究旨在超越传统可达性评估的技术窠臼，将分析单元从抽象的路网节点下沉至具体的街道断面与人群体验，'
        '通过解构可达性幻觉的要素构成与生成机制，重新锚定15分钟城市在高密度语境下的真实内涵。'
        '这不仅是一个方法论问题，更是一个关乎谁的城市、为谁服务的空间正义命题。'
    )
    add_para(doc, safe_text(p1_intro))

    add_h2(doc, '1.1  从距离达标到体验断裂的理论转向')
    p11 = (
        '主流规划实践与学术研究普遍将"可达性"简化为几何距离或网络阻抗的函数，'
        '预设了所有路径具有均等的通行效率、安全品质与心理舒适度。'
        '这种"距离中心主义"的度量方式，在深圳这样的高密度、高异质性中国超大城市中，'
        '极易制造出一种危险的"可达性幻觉"。'
    )
    add_para(doc, safe_text(p11), indent=True)

    add_h2(doc, '1.2  可达性幻觉的概念界定')
    p12 = (
        '传统15分钟城市可达性评估依赖两种方法：基于欧氏（直线）距离的缓冲区分析，'
        '以及基于名义路网的服务覆盖率计算。这两种方法均存在系统性偏差。'
        '欧氏距离忽略了河流、铁路、高速公路、封闭社区围墙等物理障碍对步行路径的阻断效应；'
        '名义路网分析则将所有道路视为等权通行，忽视了人行道缺失、步行设施破损、街道照明不足等步行环境质量差异。'
        '我们把名义承诺可达性（Promised Accessibility）与实际步行可达性（Actual Walkability）'
        '之间的系统性偏差称为"可达性幻觉"。可达性幻觉包含三个互补维度：'
        '空间维度（Spatial Accessibility Illusion）——路网障碍幻觉，指物理障碍和路网绕行使实际步行距离远超直线距离；'
        '时间维度（Temporal Accessibility Illusion）——服务可用性幻觉，指白天可用的服务在夜间可能不可用；'
        '质量维度（Walkability Environment Illusion）——步行环境幻觉，指统计模型基于POI密度计算的可达性'
        '未能反映实际步行路径的质量。'
    )
    add_para(doc, safe_text(p12), indent=True)

    add_fig(doc, P('15分钟城市时间贫困研究','conference_paper','figures','fig2_euclidean_vs_network.png'),
            width=Cm(14), caption_text='图1  欧氏距离可达性 vs 路网可达性对比')

    new_page(doc)

    # ══════════════════════════════════════════════════════════════════
    # CHAPTER 2
    # ══════════════════════════════════════════════════════════════════
    add_h1(doc, '第2章 数据基础与研究框架：真实世界复杂性的量化表征')

    p2_intro = (
        '为捕捉可达性幻觉的多维特征，本研究摒弃了单一数据源的局限性，构建了融合官方统计、开源地理信息与实地验证的'
        '多源异构数据集。研究基底为南山区全域，核心数据包括：官方建筑轮廓数据用于提取建筑密度、楼间距及形态类型；'
        '社区实有人口数据提供了精确到居住小区尺度的人口结构信息；整合后的多源POI数据涵盖医疗、教育、商业、公园、交通及金融服务六大类；'
        '精细化路网数据嵌入人行道有无、坡度、过街设施及夜间照明等属性字段。'
    )
    add_para(doc, safe_text(p2_intro))

    add_h2(doc, '2.1  研究区域')
    add_fig(doc, P('15分钟城市时间贫困研究','conference_paper','figures','fig3_study_area.png'),
            width=Cm(14), caption_text='图2  研究区域：深圳市南山区全域')
    add_para(doc, safe_text(
        '南山区位于深圳市西南部，是中国高新技术企业密度最高、建成环境异质性最强的城区之一。'
        '区域内部既有深圳湾超级总部基地、后海总部基地等高端商务区，'
        '也有大冲村、白石洲等大型城中村，以及大量老旧商品房社区，'
        '构成研究高密度城市可达性幻觉的理想"天然实验室"。'
    ))

    add_h2(doc, '2.2  多源数据体系')
    p21 = (
        '本研究构建的数据体系整合了四大类数据源。建筑轮廓数据来自自然资源部门官方数据库；'
        '社区实有人口数据精确到居住小区尺度；多源POI数据整合了高德、腾讯及百度地图开放接口；'
        '精细化路网数据在OpenStreetMap和商业地图数据的基础上，由研究团队进行了系统性步行属性补充调查。'
    )
    add_para(doc, safe_text(p21), indent=True)

    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig7_building_data.png'),
            width=Cm(14), caption_text='图3  建筑轮廓与形态类型数据')
    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig3_morphology.png'),
            width=Cm(14), caption_text='图4  建筑形态类型分布')

    add_h2(doc, '2.3  三层递进分析框架')
    add_fig(doc, P('15分钟城市时间贫困研究','conference_paper','figures','fig1_framework.png'),
            width=Cm(14), caption_text='图5  研究框架与技术路线图')
    p22 = (
        '研究构建了三层递进的分析框架。第一层为"名义可达性"，采用改进的两步移动搜索法（2SFCA）'
        '计算各社区在理想条件下的服务获取能力。第二层为"实地可达性"，通过叠加街道障碍物惩罚系数、'
        '建筑密度衰减函数及弱势群体出行速度修正，还原真实步行成本。'
        '第三层为"幻觉诊断与时空分异"，通过计算可达性幻觉指数与时间贫困指数，'
        '识别规划失效区域，并引入昼夜阻抗差异模型，揭示时间维度上的可达性动态变化。'
    )
    add_para(doc, safe_text(p22), indent=True)

    add_h2(doc, '2.4  路网构建与网络分析方法')
    p24 = (
        '网络数据集是可达性分析的空间载体，其构建质量直接决定了路径计算结果的可靠性。'
        '本研究采用Network Analyst构建研究区域的多模式步行网络数据集。'
        '网络构建以OpenStreetMap街道数据为基础底图，叠加商业地图数据补充支路与巷道信息，'
        '并由研究团队对城中村密集步行通道进行了专项矢量化，最终形成包含61402条道路边的网络拓扑。'
        '每条网络边被赋予步行阻抗权重，由基础步行速度、障碍物惩罚系数、建筑密度衰减函数'
        '与昼夜照明差异系数共同调制。在最短路径计算方面，Dijkstra算法作为精确算法被选定为基准算法，'
        '同时采用蚁群优化（ACO）、遗传算法（GA）、粒子群优化（PSO）、模拟退火（SA）'
        '及禁忌搜索（TS）五类元启发式算法在100次独立运行中验证结果稳健性。'
    )
    add_para(doc, safe_text(p24), indent=True)

    add_table(doc,
        ['指标', 'Dijkstra', 'ACO', 'GA', 'PSO', 'SA', 'TS'],
        [
            ['平均误差（%）', '0.00', '1.42', '1.78', '1.87', '2.15', '1.87'],
            ['最大误差（%）', '0.00', '3.24', '4.12', '3.98', '4.67', '4.01'],
            ['计算时间（s）', '0.34', '1.23', '2.45', '0.89', '1.56', '1.98'],
            ['收敛率（%）', '100', '98.2', '97.5', '98.7', '96.8', '98.1'],
        ]
    )
    doc.add_paragraph()

    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig2_convergence.png'),
            width=Cm(14), caption_text='图6  元启发式算法收敛性验证')
    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig8_network_analysis_sci.png'),
            width=Cm(15), caption_text='图7  南山区步行网络综合可达性幻觉分析')
    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig_network_analysis_nanshan.png'),
            width=Cm(15), caption_text='图8  南山区步行网络分析详图')

    new_page(doc)

    # ══════════════════════════════════════════════════════════════════
    # CHAPTER 3
    # ══════════════════════════════════════════════════════════════════
    add_h1(doc, '第3章 可达性幻觉的空间图谱与结构性成因')

    p3_intro = (
        '通过对南山区全域402个社区的逐一测算，研究绘制了综合可达性分布图与可达性幻觉指数空间分布图，'
        '二者叠加呈现出令人警醒的二元图景。名义上的高可达性区域主要集中在南山中心区、科技园核心区及蛇口片区，'
        '然而当视线转向可达性幻觉指数图层时，高值区并未均匀分布于低可达性边缘地带，'
        '反而高度聚集于上述高可达性区域内部的城中村及老旧住宅组团。'
        '四象限散点图清晰地揭示了这一悖论：大量样本落入第四象限（高统计可达性、低实地可达性），'
        '其中城中村样本占比超过七成。'
    )
    add_para(doc, safe_text(p3_intro))

    add_fig(doc, P('15分钟城市时间贫困研究','conference_paper','figures','fig4_illusion_scatter.png'),
            width=Cm(14), caption_text='图9-a  可达性幻觉四象限散点分布图')
    add_fig(doc, P('15分钟城市时间贫困研究','conference_paper','figures','fig7_ai_distribution.png'),
            width=Cm(14), caption_text='图9-b  综合可达性指数（AI*）空间分布图')
    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig4_aii_quadrant.png'),
            width=Cm(14), caption_text='图9-c  可达性幻觉指数四象限分析')

    add_h2(doc, '3.1  综合可达性与幻觉指数的二元图景')
    p31 = (
        '研究绘制的综合可达性分布图揭示了南山区设施供给的空间非均衡格局：'
        '高品质设施高度集聚于科技园南区、海岸城商圈及蛇口海上世界周边。'
        '然而当将同一社区的可达性幻觉指数（AII）叠加于名义可达性地图之上时，'
        '一个反直觉的图景浮现：高幻觉区域并非如预期般均匀分布于设施匮乏的边缘地带，'
        '而是高度富集于设施密集的建成区核心——特别是南山中心区、科技园南区内部的城中村及老旧住宅组团。'
        '这一发现颠覆了"设施匮乏导致不可达"的直觉假设，'
        '揭示出"设施近在咫尺但无法抵达"这一更为隐蔽的规划失效形态。'
    )
    add_para(doc, safe_text(p31), indent=True)

    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig_saii_walkability_relationship.png'),
            width=Cm(14), caption_text='图10  SAII与步行可达性关联分析')
    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig9_saii_walkability_sci.png'),
            width=Cm(15), caption_text='图11  综合幻觉指数（SAII）与步行可达性多维关联')

    add_h2(doc, '3.2  街道障碍物与步行环境的微观阻断')
    p32 = (
        '深入剖析其成因，街道障碍物分布与影响图提供了关键线索。'
        '在高幻觉社区，阻碍步行的并非遥远的距离，而是近在咫尺的物理阻隔：'
        '被占道经营压缩至不足一米的人行道、缺乏无障碍坡道的过街天桥、'
        '机动车违停形成的断头路、以及夜间完全无照明的背街小巷。'
        '这些微观障碍在宏观路网模型中被彻底抹平，却在居民的每日通勤中累积成巨大的时间与心理成本。'
        'Getis-Ord Gi*分析识别出的高-高聚类区与城中村边界高度吻合，'
        '证实了可达性幻觉本质上是一种环境不公。'
    )
    add_para(doc, safe_text(p32), indent=True)

    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig_sv_obstacle_mosaic.png'),
            width=Cm(15), caption_text='图12-a  南山区典型建成环境街景障碍物标注样本（12个代表性场景）')
    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig_sv_before_after_mosaic.png'),
            width=Cm(15), caption_text='图12-b  街景标注原始图与障碍物语义分割标注对照示例')
    add_fig(doc, P('15分钟城市时间贫困研究','conference_paper','figures','fig10_streetview_methodology.png'),
            width=Cm(14), caption_text='图12-c  街景障碍物识别与分类方法')

    add_h2(doc, '3.3  步行环境四维评分')
    p33 = (
        '研究基于Qwen3-VL-8B-Instruct对1166个建筑点位的四向街景图像进行深度学习评估，'
        '提取步行环境四维度评分：WS（人行道评分）、SI（安全指数）、AI（便利性指数）、NVS（夜间可见性评分）。'
        '综合评分公式为GTA = 0.40×WS + 0.35×SI + 0.25×NVS。'
        '四类社区的步行环境雷达图对比显示，城中村在NVS方面表现优于高端社区，'
        '验证了非正式夜间经济对生活圈韧性的独特贡献。'
    )
    add_para(doc, safe_text(p33), indent=True)

    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig5_radar.png'),
            width=Cm(14), caption_text='图13  步行环境四维评分雷达图')
    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','dl_integration_results.png'),
            width=Cm(15), caption_text='图14  深度学习步行环境评估集成结果')

    add_h2(doc, '3.4  供需匹配与多维指标关联')
    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig10_poi_distribution_sci.png'),
            width=Cm(15), caption_text='图15  南山区10类生活服务设施空间分布（背景为SAII综合幻觉指数热力图）')
    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig_poi_distribution_detail.png'),
            width=Cm(15), caption_text='图16  POI分类详细分布图')
    add_fig(doc, P('15分钟城市时间贫困研究','conference_paper','figures','fig9_supply_demand.png'),
            width=Cm(14), caption_text='图17  供需匹配三维框架分析')

    add_h2(doc, '3.5  时间贫困与昼夜时空剥夺')
    p34 = (
        '可达性幻觉的后果远不止于出行不便，它直接转化为对居民生活质量的实质性侵蚀。'
        '时间贫困指数（TPI）空间聚类图显示，高时间贫困热点与高可达性幻觉热点几乎完全重叠。'
        '对于商品房社区居民而言，夜间可达性仅比白天下降约10%至15%；'
        '但对于城中村及老旧小区居民，夜间可达性平均收缩幅度高达40%至60%。'
        '这一巨大落差源于多重因素的叠加：夜间照明覆盖率低导致安全感骤降；'
        '非正规商业摊贩收摊后街道界面变得封闭且缺乏自然监视；'
        '公共交通班次减少使得接驳段步行距离被动延长。'
    )
    add_para(doc, safe_text(p34), indent=True)

    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig6_time_poverty.png'),
            width=Cm(14), caption_text='图18  时间贫困指数空间聚类图')
    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig_four_index_spatial_comparison.png'),
            width=Cm(15), caption_text='图19  多维可达性指标空间分布对比（SAII · TPI · AI* · SCR）')
    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig11_four_index_comparison_sci.png'),
            width=Cm(15), caption_text='图20  四指标并列空间分布详图')
    add_fig(doc, P('15分钟城市时间贫困研究','conference_paper','figures','fig8_day_night.png'),
            width=Cm(14), caption_text='图21-a  昼夜时空可达性对比分析')
    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','fig_day_night_comparison_sci.png'),
            width=Cm(15), caption_text='图21  南山区昼夜时空可达性对比分析（6面板大图）')

    new_page(doc)

    # ══════════════════════════════════════════════════════════════════
    # CHAPTER 4
    # ══════════════════════════════════════════════════════════════════
    add_h1(doc, '第4章 跨区对比分析：南山、宝安、福田与龙华')

    p4_intro = (
        '为检验"高密度中心区是唯一可达性幻觉风险源"这一假设，'
        '本研究将街景数据采集范围从南山区扩展至宝安区（西乡、航城、新安，58个样本）、'
        '福田区（香蜜湖、莲花、沙头，48个样本）和龙华区（民治、大浪，48个样本），'
        '形成对照实验设计。采用DeepLabV3+语义分割（294个样本）'
        '评估建成环境语义构成，为跨区比较提供一致的建成环境量化标准。'
    )
    add_para(doc, safe_text(p4_intro))

    add_fig(doc, P('15分钟城市时间贫困研究','conference_paper','figures','fig5_type_analysis.png'),
            width=Cm(14), caption_text='图22-a  四区建成环境类型障碍评分分析')
    add_fig(doc, P('15分钟城市时间贫困研究','conference_paper','figures','fig6_deprived_communities.png'),
            width=Cm(14), caption_text='图22-b  弱势群体剥夺热点与高幻觉区域空间耦合分析')

    add_h2(doc, '4.1  四区障碍物与建成环境的量化对比')
    p41 = (
        'YOLO障碍检测与DeepLabV3语义分割揭示了各区在视觉障碍与建成环境结构上的显著差异。'
        '南山区障碍评分高达8.45分，绿视率仅9.8%，障碍来源多元化；'
        '宝安区障碍评分为8.26分、天空开敞率却高达40.2%，障碍来源以机场高速和产业大街区为主；'
        '福田区障碍评分最低，仅3.62分，人行空间识别占比最高，步行设施配建率最高；'
        '龙华区障碍评分最低，仅2.86分，新城区的视觉阻隔最低，但服务成熟度与夜间可用性仍待持续检验。'
    )
    add_para(doc, safe_text(p41), indent=True)

    add_fig(doc, P('projects','15min-urban-accessibility','paper','figures','section13_results.png'),
            width=Cm(15), caption_text='图23  跨区障碍物检测与分类统计结果')

    add_h2(doc, '4.2  反事实分析与机制解释')
    p42 = (
        '采用反事实假设法可进一步解释跨区幻觉差异。'
        '假设一（南山vs宝安）：控制服务密度与街景可达性指数相同后，宝安的可达性幻觉指数可能更负——'
        '其机制在于宝安机场快速路和产业大街区的绕行成本，高于南山城中村密集巷道提供的通透效益。'
        '假设二（南山vs福田）：控制POI密度与人口密度相同后，福田的可达性幻觉指数预期更接近零——'
        '其机制在于福田的步行空间连续性和人行设施密度系统性优于南山。'
        '假设三（南山vs龙华）：龙华的障碍评分最低，但低成熟度意味着POI密度本身不足，'
        '"幻觉"的主要来源已不是"设施不可达"而是"设施不存在"。'
    )
    add_para(doc, safe_text(p42), indent=True)

    new_page(doc)

    # ══════════════════════════════════════════════════════════════════
    # CHAPTER 5
    # ══════════════════════════════════════════════════════════════════
    add_h1(doc, '第5章 研究局限与治理建议')

    add_h2(doc, '5.1  研究局限')
    p51 = (
        '本研究在数据覆盖、方法假设与推断范围上存在以下局限。'
        '在街景样本方面，1166个建筑点位覆盖约70%的社区质心，其余社区的步行环境评分由空间插值估算。'
        '在步行速度假设方面，研究将所有社区的步行速度统一设定为1.2米/秒，'
        '尚未分年龄、分性别设置差异化步行速度参数。'
        '在路网完整性方面，OSM与高德路网数据中可能缺失部分非正式步行连接。'
        '在夜间评估方面，现有分析基于POI营业时间数据，尚无针对夜间安全感的主观调查数据支撑。'
    )
    add_para(doc, safe_text(p51), indent=True)

    add_h2(doc, '5.2  治理建议')
    p52 = (
        '基于研究结论，提出从三个维度升级15分钟生活圈评估体系的治理建议。'
        '在空间维度上，建议将可达性幻觉指数与路网比率纳入评价指标体系，'
        '增设"实地步行可达性验证"作为新建居住区规划审批的前置条件。'
        '在质量维度上，建议将时间贫困指数纳入城中村更新单元规划的审批考量，'
        '增设"步行网络通透性评估"专章，'
        '防止出现"更新但断路"的悖论。'
        '在时间维度上，建议将夜间POI可用率纳入公共服务配套标准，'
        '明确纳入24小时便利店服务覆盖率和夜间步行照明连续性等韧性指标。'
    )
    add_para(doc, safe_text(p52), indent=True)

    add_h2(doc, '5.3  未来研究方向')
    p53 = (
        '未来研究可从以下方向深化：扩展街景样本至全域覆盖；'
        '纳入差异化步行速度修正系数，特别是针对老年人（0.6至0.8米/秒）、'
        '儿童及残障人士的专项修正；'
        '引入GPS轨迹数据或步行日记调查，校验模型预测路径与实际路径选择的一致性；'
        '叠加收入水平、住房权属、照护负担等社会经济属性，构建更完整的社会公平分析框架；'
        '将研究框架推广至广州天河区、北京朝阳区、上海浦东新区等高密度城区，检验发现的普适性；'
        '接入实时交通数据，构建动态更新的15分钟生活圈评估系统。'
    )
    add_para(doc, safe_text(p53), indent=True)

    new_page(doc)

    # ══════════════════════════════════════════════════════════════════
    # CONCLUSION
    # ══════════════════════════════════════════════════════════════════
    add_h1(doc, '结论')
    conclusion = (
        '我们基于深圳市南山区402个社区、69424条POI记录、61402条路网边数据和1166条建筑记录的实证数据，'
        '首次揭示城中村因密集步行巷道而具有优于高端社区的实际可达性这一反直觉发现，'
        '系统解构了高密度城市15分钟生活圈中"可达性幻觉"的要素、成因与时空表现。'
        '核心发现可归纳为三点：'
        '其一，可达性幻觉是高密度异质城市的内生特征而非数据误差，'
        '其根源在于宏观规划指标与微观建成环境品质的脱节；'
        '其二，幻觉具有强烈的社会空间选择性，与弱势群体聚居区及街道物理障碍高度耦合，'
        '构成了一种隐蔽的环境不公；'
        '其三，可达性具有显著的昼夜动态性，夜间环境恶化导致弱势群体的有效生活半径大幅收缩，'
        '形成时空双重剥夺。'
        '这些发现对当前15分钟城市实践提出了深刻警示：'
        '单纯追求设施覆盖率与距离阈值的达标，不仅无法实现真正的宜居目标，'
        '反而可能掩盖并固化既有的空间不平等。'
    )
    add_para(doc, safe_text(conclusion))

    new_page(doc)

    # ══════════════════════════════════════════════════════════════════
    # APPENDIX
    # ══════════════════════════════════════════════════════════════════
    add_h1(doc, '附录A 研究图件汇编')

    add_h2(doc, '附录A.1  Conference论文配套图件')

    # fig1 ~ fig10 conference
    conf_figs = [
        ('fig1_framework.png',      '图A1-1  研究框架与技术路线图（Conference版）'),
        ('fig2_euclidean_vs_network.png', '图A1-2  欧氏距离 vs 路网可达性对比'),
        ('fig3_study_area.png',     '图A1-3  研究区域：深圳市南山区全域'),
        ('fig4_illusion_scatter.png', '图A1-4  可达性幻觉四象限散点分布图'),
        ('fig5_type_analysis.png',  '图A1-5  四区建成环境类型障碍评分分析'),
        ('fig6_deprived_communities.png', '图A1-6  弱势群体剥夺热点空间耦合分析'),
        ('fig7_ai_distribution.png', '图A1-7  综合可达性指数（AI*）空间分布图'),
        ('fig8_day_night.png',      '图A1-8  昼夜时空可达性对比分析（Conference版）'),
        ('fig9_supply_demand.png',  '图A1-9  供需匹配三维框架分析'),
        ('fig10_streetview_methodology.png', '图A1-10  街景障碍物识别与分类方法'),
    ]
    for fname, cap in conf_figs:
        p = P('15分钟城市时间贫困研究','conference_paper','figures',fname)
        if os.path.exists(p):
            add_fig(doc, p, width=Cm(15), caption_text=cap)
            doc.add_paragraph()

    new_page(doc)
    add_h2(doc, '附录A.2  Paper论文核心图件')

    # fig1 ~ fig11 paper figures
    paper_figs = [
        ('fig1_framework.png',            '图A2-1  研究框架与技术路线图（Paper版）'),
        ('fig2_convergence.png',          '图A2-2  元启发式算法收敛性验证'),
        ('fig3_morphology.png',            '图A2-3  建筑形态类型分布'),
        ('fig4_aii_quadrant.png',          '图A2-4  可达性幻觉指数四象限分析'),
        ('fig5_radar.png',                '图A2-5  步行环境四维评分雷达图'),
        ('fig6_time_poverty.png',          '图A2-6  时间贫困指数空间聚类图'),
        ('fig7_building_data.png',        '图A2-7  建筑轮廓与形态类型数据'),
        ('fig8_network_analysis_sci.png',   '图A2-8  步行网络综合可达性幻觉分析（SCI版）'),
        ('fig9_saii_walkability_sci.png', '图A2-9  综合幻觉指数（SAII）与步行可达性多维关联'),
        ('fig10_poi_distribution_sci.png','图A2-10  10类生活服务设施空间分布（背景为SAII热力图）'),
        ('fig11_four_index_comparison_sci.png', '图A2-11  四指标并列空间分布详图'),
    ]
    for fname, cap in paper_figs:
        p = P('projects','15min-urban-accessibility','paper','figures',fname)
        if os.path.exists(p):
            add_fig(doc, p, width=Cm(15), caption_text=cap)
            doc.add_paragraph()

    new_page(doc)
    add_h2(doc, '附录A.3  Paper论文扩展分析图件')

    ext_figs = [
        ('fig_network_analysis_nanshan.png',    '图A3-1  南山区步行网络分析详图'),
        ('fig_poi_distribution_detail.png',   '图A3-2  POI分类详细分布图'),
        ('fig_saii_walkability_relationship.png','图A3-3  SAII与步行可达性关联分析'),
        ('fig_four_index_spatial_comparison.png','图A3-4  多维可达性指标空间分布对比'),
        ('fig_day_night_comparison_sci.png',     '图A3-5  昼夜时空可达性对比分析（6面板大图）'),
        ('fig_sv_obstacle_mosaic.png',          '图A3-6  街景障碍物标注样本（12个代表性场景）'),
        ('fig_sv_before_after_mosaic.png',        '图A3-7  街景标注原始图与障碍物语义分割对照示例'),
        ('dl_integration_results.png',            '图A3-8  深度学习步行环境评估集成结果'),
        ('section13_results.png',                '图A3-9  跨区障碍物检测与分类统计结果'),
    ]
    for fname, cap in ext_figs:
        p = P('projects','15min-urban-accessibility','paper','figures',fname)
        if not os.path.exists(p):
            p = P('projects','15min-urban-accessibility','paper',fname)
        if os.path.exists(p):
            add_fig(doc, p, width=Cm(15), caption_text=cap)
            doc.add_paragraph()

    new_page(doc)
    add_h2(doc, '附录A.4  街道断面分析详图（144条样本）')

    profile_files = sorted([f for f in os.listdir(SP) if f.endswith('.png')])
    print(f'Found {len(profile_files)} street profile images')

    # Insert in rows of 2
    for i in range(0, len(profile_files), 2):
        row_imgs = []
        for j in range(2):
            if i + j < len(profile_files):
                row_imgs.append(os.path.join(SP, profile_files[i + j]))
        if len(row_imgs) == 2:
            add_fig_grid(doc, row_imgs, widths=[Cm(7.5), Cm(7.5)],
                         caption_text=f'图A4-{i//2+1}  街道断面分析（{profile_files[i]} & {profile_files[i+1]}）')
        elif len(row_imgs) == 1:
            add_fig(doc, row_imgs[0], width=Cm(10),
                     caption_text=f'图A4-{i//2+1}  街道断面分析（{profile_files[i]}）')
        doc.add_paragraph()

    new_page(doc)
    add_h2(doc, '附录A.5  街景标注图像样本（60张精选标注图）')

    ac_files = sorted([f for f in os.listdir(AC) if f.endswith('.jpg')])
    print(f'Found {len(ac_files)} annotated_cn images')

    add_para(doc, safe_text(
        f'本附录收录精选{len(ac_files)}张街景标注图像，每张图像标注了拍摄朝向（E/W/N/S）、'
        '坐标位置及障碍物类型。'
    ))
    doc.add_paragraph()

    n_display = len(ac_files)
    n_cols = 4
    n_rows = (n_display + n_cols - 1) // n_cols
    for row in range(n_rows):
        row_paths = []
        for col in range(n_cols):
            idx = row * n_cols + col
            if idx < n_display:
                row_paths.append(os.path.join(AC, ac_files[idx]))
        if row_paths:
            widths = [Cm(3.7)] * len(row_paths)
            cap_text = f'图A5-{row+1}  街景标注样本（第{row*n_cols+1}至{row*n_cols+len(row_paths)}张，共{len(ac_files)}张）'
            add_fig_grid(doc, row_paths, widths=widths, caption_text=cap_text, n_cols=n_cols)
            doc.add_paragraph()

    new_page(doc)

    # ── Save ──────────────────────────────────────────────────────
    out_path = P('报告_final.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    print(f'Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}')
    return out_path

if __name__ == '__main__':
    build_report()
