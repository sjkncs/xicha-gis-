from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "报告_final_哈工大模板.docx"

FONT_SONG = "宋体"
FONT_LATIN = "Times New Roman"
BODY_SIZE = 12.0


CONCLUSION_PARAGRAPHS = [
    "本研究围绕高密度城市15分钟生活圈中“地图达标但体验失效”的关键问题展开。我们团队没有把生活圈评价停留在设施数量、服务半径或静态路网时间上，而是把居民在真实街道中遭遇的绕行、阻隔、夜间服务衰减、无障碍设施缺口和弱势群体通行成本纳入同一分析框架。基于深圳市南山区402个社区、69424条生活服务POI、61402条步行网络边、1166个建筑与街景评估点位，并结合南山、宝安、福田、龙华四区294张街景样本，本文从定量指标和视觉证据两个层面说明：15分钟生活圈的核心矛盾并不只是“服务有没有”，更在于“服务是否能够被不同群体以合理时间和安全成本真实使用”。因此，结论部分重点概括研究的主要结果、创新性贡献与后续展望，而不再重复摘要式的内容介绍。",
    "（1）提出了“可达性幻觉”这一面向高密度城市生活圈评价的新见解。我们团队将名义可达性与实际步行可达性之间的系统性偏差界定为可达性幻觉，并进一步指出，这种幻觉不是个别道路计算误差，也不是居民主观感受的偶然偏差，而是由欧氏距离假设、简化路网模型、封闭街区边界、交通基础设施切割、街道微观障碍和夜间服务衰减共同造成的结构性问题。通过这一概念，研究把传统“15分钟内是否有设施”的判断推进为“居民是否能够真实、连续、安全地到达设施”的判断，为生活圈评价提供了更贴近城市实际运行状态的理论视角。",
    "（2）建立了多源数据融合的可达性幻觉测度体系。我们团队将社区人口、POI服务供给、步行路网、建筑形态、街景图像和跨区样本整合到统一框架中，构建了SAII、TPI、AI*、SCR等互补指标，使服务供给、路径绕行、时间贫困和街景通行阻抗能够在社区尺度上被共同解释。与单纯依赖缓冲区或最短路径的方法相比，该体系的创新点在于把街景视觉识别结果转化为可进入路网阻抗和空间诊断的变量，使车辆占道、路障、围挡、杆柱、人行空间压缩、夜间可见性不足等过去难以进入模型的微观因素具备了量化表达能力。",
    "（3）揭示了南山区可达性幻觉的主要定量结果。研究结果表明，南山区生活服务供给密度较高，但真实步行体验并未同步达到理想状态。全样本平均路网比率约为1.42，意味着直线或名义路径所暗示的15分钟可达，在实际路网中往往会被拉长到约21分钟以上；85.1%的社区呈现正向可达性幻觉，说明“名义达标、现实折损”并非少数社区的局部问题，而是高密度建成环境中广泛存在的评价偏差。更重要的是，高幻觉社区并不总是位于设施稀缺地区，反而大量出现在服务密集但边界封闭、过街困难、微观障碍密集或夜间服务下降的空间单元中，这一结果从定量上否定了“设施越多即可达性越好”的简单判断。",
    "（4）阐明了可达性幻觉与时空公平之间的内在联系。我们团队发现，可达性幻觉具有明显的社会空间选择性：城中村边界、老旧住区、大型封闭社区周边、快速路和铁路沿线以及弱势群体居住或活动密集区域，更容易同时承受路径绕行、街道障碍和夜间安全感下降。也就是说，生活圈失效并不是平均地落在所有居民身上，而是会通过步行速度差异、照护负担、身体能力差异和夜间出行需求被进一步放大。由此，本文将15分钟生活圈研究从空间覆盖问题推进到时空正义问题，说明生活圈政策若忽视老人、儿童、残障人士、推婴儿车者和夜间劳动者的真实通行条件，就可能在统计上达标，却在实践中继续制造隐性的可达性不平等。",
    "（5）形成了面向治理应用的方法创新和技术路线。我们团队将网络分析、街景深度学习、供需匹配和跨区对比结合起来，不只是为了生成图件或指标，而是为了把“哪里不可达、为什么不可达、应当如何修复”连接成可操作的诊断链条。研究显示，南山区需要重点修复封闭边界、过街断点和微观障碍，宝安区需要处理快速路和产业大街区造成的尺度阻隔，福田区的连续人行网络提供了相对正向的对照经验，龙华区则更需要同步提升服务密度与夜间可用性。这种分区化解释使可达性幻觉指数不只是描述性结果，也可以转化为城市体检、生活圈补短板、无障碍设施建设和街景智能巡检的优先级依据。",
    "（6）指出了后续研究的拓展方向。受街景覆盖、群体行为数据和微观标注体系限制，本文仍有进一步深化空间。未来我们团队计划从三个方面继续推进：一是扩展街景采样与无障碍障碍物标注，重点补充台阶、坡道、盲道、减速带、路面破损和临时围挡等通用模型难以稳定识别的类别；二是引入GPS轨迹、步行日记、问卷访谈和主观安全感评价，校验模型预测路径与居民真实路径选择之间的差异；三是将南山区框架推广到广州、北京、上海等更多高密度城区，比较不同城市形态、道路制度和更新模式下可达性幻觉的生成机制。总体而言，本文的贡献在于提出了可达性幻觉的概念，建立了多源融合的测度方法，揭示了其定量分布和公平后果，并为后续动态化、精细化、面向治理的15分钟生活圈评价系统奠定了基础。",
]


def text(paragraph) -> str:
    return paragraph.text.strip()


def set_run_font(run) -> None:
    run.font.name = FONT_LATIN
    run.font.size = Pt(BODY_SIZE)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_SONG)
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:cs"), FONT_LATIN)


def apply_body_format(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        set_run_font(run)


def replace_paragraph_text(paragraph, value: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(value)
    set_run_font(run)
    apply_body_format(paragraph)


def remove_paragraph(paragraph) -> None:
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def main() -> None:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = ROOT / f"报告_final_哈工大模板.before_conclusion_expression_{stamp}.docx"
    shutil.copy2(DOCX, backup)

    doc = Document(DOCX)
    start = next(i for i, p in enumerate(doc.paragraphs) if text(p) == "结  论")
    end = next(i for i, p in enumerate(doc.paragraphs[start + 1 :], start + 1) if text(p) == "参考文献")

    body_paragraphs = doc.paragraphs[start + 1 : end]
    for idx, value in enumerate(CONCLUSION_PARAGRAPHS):
        if idx < len(body_paragraphs):
            replace_paragraph_text(body_paragraphs[idx], value)
        else:
            inserted = doc.paragraphs[end].insert_paragraph_before(value, style="Normal")
            apply_body_format(inserted)

    # Remove leftover non-empty old conclusion paragraphs and the accidental blank Heading 1 before references.
    doc = Document(DOCX) if False else doc
    start = next(i for i, p in enumerate(doc.paragraphs) if text(p) == "结  论")
    end = next(i for i, p in enumerate(doc.paragraphs[start + 1 :], start + 1) if text(p) == "参考文献")
    middle = doc.paragraphs[start + 1 : end]
    for paragraph in middle[len(CONCLUSION_PARAGRAPHS) :]:
        remove_paragraph(paragraph)

    doc.save(DOCX)

    doc = Document(DOCX)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    start = next(i for i, p in enumerate(doc.paragraphs) if text(p) == "结  论")
    end = next(i for i, p in enumerate(doc.paragraphs[start + 1 :], start + 1) if text(p) == "参考文献")
    conclusion_count = sum(1 for p in doc.paragraphs[start + 1 : end] if text(p))
    print("BACKUP", backup)
    print("OUTPUT", DOCX)
    print(
        "VALIDATION",
        {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "inline_shapes": len(doc.inline_shapes),
            "conclusion_paragraphs": conclusion_count,
            "has_team_subject": "我们团队" in all_text,
            "has_numbered_conclusion": all(f"（{i}）" in all_text for i in range(1, 7)),
            "bad_question_marks": "????????" in all_text,
            "replacement_char": "�" in all_text,
            "a5_drawings": sum(t._tbl.xml.count("<w:drawing") for t in doc.tables if len(t.rows) == 15 and len(t.columns) == 4),
        },
    )


if __name__ == "__main__":
    main()
