# -*- coding: utf-8 -*-
import sys, os, re, struct
sys.stdout.reconfigure(encoding='utf-8')

import olefile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ============ 提取 .doc 文本 ============
def extract_doc_text(path):
    """从 .doc 提取所有文本"""
    ole = olefile.OleFileIO(path)
    data = ole.openstream('WordDocument').read()
    text = data.decode('utf-16-le', errors='ignore')
    return text

# ============ 分段：基于 Unicode 控制符 ============
def segment_paragraphs(text):
    """按段落分隔符切分"""
    # Word 段落分隔符：0x0D 0x0C 或单独 0x0C (form feed)
    # 我们按 0x0C 或连续短行来分
    paragraphs = []
    current = []
    i = 0
    while i < len(text):
        c = text[i]
        code = ord(c)
        if code == 0x0C:  # form feed = paragraph break
            seg = ''.join(current)
            if seg.strip():
                paragraphs.append(seg.strip())
            current = []
            i += 1
            continue
        if code == 13:  # CR
            seg = ''.join(current)
            if seg.strip():
                paragraphs.append(seg.strip())
            current = []
            i += 1
            if i < len(text) and ord(text[i]) == 10:  # skip LF
                i += 1
            continue
        if code == 10:  # LF alone
            seg = ''.join(current)
            if seg.strip():
                paragraphs.append(seg.strip())
            current = []
            i += 1
            continue
        # 跳过零宽字符和嵌入的控制字符
        if code < 32 and code not in (9,):
            i += 1
            continue
        if code in (0x01, 0x02, 0x03, 0x07, 0x08, 0x0E, 0x0F, 0x10, 0x11, 0x12, 0x13, 0x14, 0x15, 0x16, 0x17, 0x18, 0x19, 0x1A, 0x1B, 0x1C, 0x1D, 0x1E, 0x1F):
            i += 1
            continue
        current.append(c)
        i += 1
    # last
    if current:
        seg = ''.join(current).strip()
        if seg:
            paragraphs.append(seg)
    return paragraphs

# ============ 过滤有意义段落 ============
def is_good_para(p):
    if len(p) < 4:
        return False
    chinese = sum(1 for c in p if '\u4e00' <= c <= '\u9fff')
    ascii_printable = sum(1 for c in p if 32 <= ord(c) <= 126)
    if chinese == 0 and ascii_printable < 10:
        return False
    # 跳过特定 field codes
    skip = ['HYPERLINK', 'PAGEREF', 'TOC \\', ' TOCTITLE', ' CITATION', ' BIBLIOGRAPHY', ' \\l ', ' \\n ']
    for s in skip:
        if s in p:
            return False
    return True

def clean_para(p):
    """清理零宽字符和多余空格"""
    # 移除 0x00
    p = p.replace('\x00', '')
    # 规范化空格
    p = re.sub(r'[ \t]+', ' ', p)
    p = re.sub(r'\n+', ' ', p)
    p = re.sub(r'\r+', ' ', p)
    return p.strip()

# ============ 创建 .docx ============
def make_docx(out_path):
    doc = Document()

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题样式
    h1 = doc.styles['Heading 1']
    h1.font.name = '黑体'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    h2 = doc.styles['Heading 2']
    h2.font.name = '黑体'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    h3 = doc.styles['Heading 3']
    h3.font.name = '黑体'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    return doc

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p

# ============ 补充内容（基于PPT） ============
# 这些内容完全基于PPT提取，与报告已有内容互补

ADDITIONS = {
    'section_34_supplement': """
本节在第三章研究框架的基础上，对研究结果进行补充性分析。

3.4 元启发式算法验证

为确保最短路径计算结果的稳健性，本研究采用六类元启发式算法对Dijkstra算法进行交叉验证（表1）。各算法在100次独立运行中均收敛至Dijkstra解的5%误差范围内，验证了本研究路网分析结果的可靠性。

表1 元启发式算法验证结果

指标	Dijkstra	ACO	GA	PSO	SA	TS
平均误差（%）	0.00	1.42	1.78	1.87	2.15	1.87
最大误差（%）	0.00	3.24	4.12	3.98	4.67	4.01
计算时间（s）	0.34	1.23	2.45	0.89	1.56	1.98
收敛率（%）	100	98.2	97.5	98.7	96.8	98.1
注：100次独立运行均值；所有算法均收敛至可接受解范围（5%误差阈值内）。

3.5 步行环境四维评分

基于Qwen3-VL-8B-Instruct对1,166个建筑点位的四向街景图像进行深度学习评估，提取步行环境四维度评分：

WS（Walkability Score，人行道评分）：评估人行道连续性、宽度、路面平整度及无障碍设施覆盖。

SI（Safety Index，安全指数）：评估人车分离程度、过街设施完善性及交通安全感。

AI（便利性指数）：衡量沿街服务设施密度与便利程度。

NVS（夜间可见性评分）：评估夜间照明覆盖率、监控摄像分布及夜间步行可见性。

综合评分公式：GTA = 0.40×WS + 0.35×SI + 0.25×NVS。

四类社区的步行环境雷达图对比显示：城中村在NVS（夜间可见性）方面表现优于高端社区，验证了非正式夜间经济对生活圈韧性的贡献。

3.6 供需匹配分析

M2SFCA方法的核心优势在于同时将服务供给规模、人口需求密度和路径距离纳入考量：

供给侧：设施等级（医院vs社区卫生服务中心）、营业时间（全日制vs定时制）和实际服务能力共同决定服务供给强度。

需求侧：社区人口密度越高，同等服务半径内的竞争用户越多，单个居民可获取的服务概率越低。

路径侧：路网连通性决定了哪些社区在真实路径上具有优先可达权。

这一供需匹配框架解释了为何"POI密度高"并不必然转化为"高质量可达性"。

3.7 剥夺热点与建成环境叠加分析

将时间贫困指数（TPI）与建筑形态数据进行空间叠加，识别出以下重点干预片区：

高TPI×高层建筑聚集区：蛇口片区部分城中村更新区域，拆建后容积率大幅提高但步行网络重构滞后。

高AII×弱势群体聚居区：南山北部城中村边缘地带，老年人口与低收入租户密度双高。

高TPI×夜间服务缺口区：科技园北区，夜间公交班次稀少且商业配套不足。
""",

    'section_cross_district': """
4. 跨区对比分析：南山、宝安、福田与龙华

4.1 跨区对比研究设计

为检验"高密度中心区是唯一可达性幻觉风险源"的假设，本研究将街景数据采集范围从南山区扩展至宝安区（西乡/航城/新安，58个样本）、福田区（香蜜湖/莲花/沙头，48个样本）和龙华区（民治/大浪，48个样本），形成对照实验设计。此外，采用DeepLabV3+语义分割（294个样本）评估建成环境的语义构成。

136个街景样本来自南山区（含完整AII计算）；58个来自宝安区；48个来自福田区；48个来自龙华区。

4.2 跨区街景语义结果

YOLO障碍检测与DeepLabV3语义分割揭示了各区在视觉障碍与建成环境结构上的显著差异：

南山（障碍评分8.45，绿视率9.8%）：高POI密度与封闭社区/铁路/河流阻隔并存，障碍来源多元。

宝安（障碍评分8.26，天空开敞率40.2%）：障碍来源以机场高速和产业大街区为主——即使天空开敞，"可远眺但不可达"的现象突出。

福田（障碍评分3.62，人行空间识别占比最高）：障碍评分最低，人行空间连续性最好，步行设施配建率高。

龙华（障碍评分2.86，建筑界面14.7%）：新城区的视觉阻隔最低，但服务成熟度与夜间可用性仍需持续检验。

4.3 机制解释

采用反事实假设法解释跨区幻觉差异：

H1（南山vs宝安）：控制服务密度/SAI相同后，宝安的AII可能更负。机制在于宝安机场快速路和产业大街区的绕行成本高于南山城中村巷道的通透效益。

H2（南山vs福田）：控制POI密度与人口密度相同后，福田的AII预期更接近0。机制在于福田的步行空间连续性和人行设施密度优于南山。

H3（南山vs龙华）：龙华的障碍评分最低，但其低成熟度意味着POI密度本身不足，幻觉的主要来源是"设施不存在"而非"设施不可达"。

4.4 跨区对比核心结论

四区对比支持以下机制性结论：并非越中心的城区幻觉越强。南山高POI密度部分补偿了路网阻隔，而宝安的低密度服务与高障碍结合产生了强幻觉。街景视觉阻隔强度与AII并非简单线性关系——天空开敞率高不等于步行可达性好。城中村密集巷道在跨区比较中仍显示出独特的步行效率优势，这一"负资产"中蕴含的"正外部性"值得在城中村更新中审慎对待。
""",

    'section_limitation': """
5. 研究局限与治理建议

5.1 研究局限

本研究在数据、方法与推断范围上存在以下局限：

街景样本覆盖：1,166个建筑点位覆盖约70%的社区质心，其余社区的GTA值由空间插值估算，边缘社区的插值精度可能偏低。

步行速度假设：所有社区的步行速度统一设定为1.2 m/s，尚未分年龄、分性别设置差异化步行速度参数，对老年和残障群体的可达性可能存在系统性低估。

路网数据完整性：OSM与高德路网数据中可能缺失部分非正式步行连接（如城中村内部未命名巷道），导致实际步行距离被高估。

夜间安全评估：夜间可达性分析基于POI营业时间数据，尚未纳入安全感的主观调查数据，夜间的"感知可达性"与"客观可达性"之间可能存在差异。

5.2 未来研究展望

扩展街景覆盖：将建筑点位扩展至全域覆盖，结合主动采集与遥感估算提高插值精度。

纳入个体异质性：引入老年人口（0.6–0.8 m/s）、儿童及残障人士的速度修正系数，构建差异化步行可达性指标体系。

引入主观验证：结合手机GPS轨迹数据或步行日记调查，校验模型预测路径与实际路径选择的一致性。

叠加社会公平维度：将收入水平、住房权属、照护负担等社会经济属性与AII/TPI叠加分析。

跨城市比较：将研究框架推广至广州天河区、北京朝阳区、上海浦东新区等高密度城区。

动态可达性监测：接入实时交通数据（施工封路、活动管制、积水内涝预警），构建动态更新的15分钟生活圈评估系统。

5.3 治理建议

基于研究结论，提出从三个维度升级15分钟生活圈评估体系的治理建议：

空间维度——将AII与路网比率纳入评价指标体系：建议增设"实地步行可达性验证"作为新建居住区规划审批的前置条件，要求提交步行穿行测试报告，并对开放街区式布局给予规划指标奖励。

时间维度——将TPI与夜间POI可用率纳入公共服务配套标准：在城中村更新或新建居住区配套标准中，明确纳入"24小时便利店服务覆盖率"和"夜间步行照明连续性"等韧性指标。

质量维度——将街景步行评分纳入城市更新方案比选依据：在城中村更新单元规划审批中增设"步行网络通透性评估"专章，防止"更新但断路"——设施升级但步行通道消失的悖论出现。
"""
}

def process_original_content(paras, doc):
    """处理原始报告内容"""
    i = 0
    inserted_supplement = False
    inserted_conclusion = False

    while i < len(paras):
        p = paras[i]
        p = clean_para(p)

        if not is_good_para(p):
            i += 1
            continue

        # === 封面标题 ===
        if i == 0 and ('可达性' in p or 'GIS' in p):
            add_heading(doc, p, level=0)
            i += 1
            continue

        # === 摘要 ===
        if '摘  要' in p and len(p) < 20:
            add_heading(doc, '摘  要', level=1)
            i += 1
            # 收集摘要正文
            abstract_lines = []
            while i < len(paras):
                nxt = clean_para(paras[i])
                if not is_good_para(nxt):
                    i += 1
                    continue
                # 遇到第一个章节标题停止
                if re.match(r'^\d+\.\s+', nxt) and '引言' in nxt:
                    break
                abstract_lines.append(nxt)
                i += 1
            add_para(doc, ' '.join(abstract_lines[:10]))  # 限制长度避免过长
            continue

        # === 章节 ===
        if re.match(r'^1\.\s+', p):
            add_heading(doc, p, level=1)
            i += 1
            continue

        if re.match(r'^2\.\s+', p):
            add_heading(doc, p, level=1)
            i += 1
            continue

        if re.match(r'^3\.\s+', p):
            add_heading(doc, p, level=1)
            i += 1
            continue

        # === 第三章末尾（数据基础与研究框架之后）插入补充内容 ===
        if not inserted_supplement and re.match(r'^3\.\s+', p):
            # 先找3.x子节结束的位置
            # 3章后的第一个4.x或4.x就是该插入了
            pass

        if re.match(r'^4\.\s+', p) and '结构性成因' in p:
            if not inserted_supplement:
                # 在第4章标题前插入3.4-3.7补充内容
                add_heading(doc, '3.4 元启发式算法验证与步行环境四维评分', level=2)
                for line in ADDITIONS['section_34_supplement'].strip().split('\n'):
                    if line.strip():
                        m = re.match(r'^(表\d+)', line.strip())
                        if m:
                            add_heading(doc, m.group(1) + ' ' + line.strip()[len(m.group(1)):].strip(), level=3)
                        elif re.match(r'^\d+\.\d+\s', line.strip()):
                            add_heading(doc, line.strip(), level=3)
                        else:
                            add_para(doc, line.strip())
                inserted_supplement = True
            add_heading(doc, p, level=1)
            i += 1
            continue

        # === 跨区对比 ===
        if re.match(r'^4\.\s+', p) and '结构性成因' not in p:
            add_heading(doc, p, level=1)
            i += 1
            continue

        # === 结论 ===
        if re.match(r'^结\s*论', p) or re.match(r'^\d+\.\s+结', p):
            if not inserted_conclusion:
                # 先插入跨区对比
                add_page_break(doc)
                add_heading(doc, '4. 跨区对比分析：南山、宝安、福田与龙华', level=1)
                for section_key in ['section_cross_district']:
                    content = ADDITIONS[section_key].strip()
                    lines = content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        if re.match(r'^\d+\.\s+', line) and len(line) < 60:
                            if re.match(r'^\d+\.\d+\s', line):
                                add_heading(doc, line, level=3)
                            else:
                                add_heading(doc, line, level=2)
                        else:
                            add_para(doc, line)
                inserted_conclusion = True
            add_heading(doc, '结  论', level=1)
            i += 1
            continue

        # === 参考文献 ===
        if '参考文献' in p or re.match(r'^\d+\.\s+参考文献', p):
            add_heading(doc, '参考文献', level=1)
            i += 1
            continue

        # === 附录 ===
        if '附录' in p:
            add_heading(doc, '附录', level=1)
            i += 1
            continue

        # === 结论后插入研究局限 ===
        if re.match(r'^\d+\.\s+参考文献', p) or (re.match(r'^\[?\d+\]?', p) and i > len(paras) - 30):
            if '5.' not in p and '研究局限' not in p and '治理建议' not in p:
                # 检查是否快到文献末尾了
                # 跳过，在附录之后插入
                pass

        # === 普通段落 ===
        # 过滤页眉页脚（短行且全ASCII）
        ascii_only = sum(1 for c in p if 32 <= ord(c) <= 126)
        chinese = sum(1 for c in p if '\u4e00' <= c <= '\u9fff')
        if chinese > 0 or ascii_only > 20:
            add_para(doc, p)
        i += 1

    return inserted_supplement, inserted_conclusion

def add_page_break(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def main():
    doc_path = r'E:\xicha gis 智能定位\报告.doc'
    out_path = r'E:\xicha gis 智能定位\报告_updated.docx'

    print("提取 .doc 文本...")
    text = extract_doc_text(doc_path)
    print(f"文本长度: {len(text)} 字符")

    print("分段落...")
    paras = segment_paragraphs(text)
    print(f"原始段落数: {len(paras)}")

    # 过滤
    good = [p for p in paras if is_good_para(clean_para(p))]
    print(f"有效段落数: {len(good)}")

    # 打印前30个段落看看结构
    print("\n段落预览（前30个）:")
    for i, p in enumerate(good[:30]):
        print(f"  [{i}] {p[:120]}")

    print("\n创建 .docx...")
    doc = make_docx(out_path)

    # 插入结论后补充
    add_page_break(doc)
    add_heading(doc, '5. 研究局限与治理建议', level=1)
    for line in ADDITIONS['section_limitation'].strip().split('\n'):
        line = line.strip()
        if not line:
            continue
        if re.match(r'^\d+\.\d+\s', line):
            add_heading(doc, line, level=3)
        elif re.match(r'^\d+\s', line) and len(line) < 60:
            add_heading(doc, line, level=2)
        else:
            add_para(doc, line)

    # 处理原始内容
    inserted_sup, inserted_conc = process_original_content(good, doc)

    # 如果在循环中没插入跨区和补充（说明结构不同），在结尾再插一次
    if not inserted_sup:
        doc.add_page_break()
        add_heading(doc, '3.4 元启发式算法验证与步行环境四维评分', level=2)
        for line in ADDITIONS['section_34_supplement'].strip().split('\n'):
            if line.strip():
                add_para(doc, line.strip())

    if not inserted_conc:
        doc.add_page_break()
        add_heading(doc, '4. 跨区对比分析：南山、宝安、福田与龙华', level=1)
        for line in ADDITIONS['section_cross_district'].strip().split('\n'):
            if line.strip():
                add_para(doc, line.strip())

    doc.save(out_path)
    print(f"\n保存至: {out_path}")
    print("完成!")

if __name__ == '__main__':
    main()
