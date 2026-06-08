# -*- coding: utf-8 -*-
"""Embed figures at the end of the document - simplest reliable approach"""
import sys, os
sys.path.insert(0, r'E:\xicha gis 智能定位')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r'E:\xicha gis 智能定位'

# Order of figures as they appear in the report
# caption -> relative path
FIGURES = [
    ('图1：综合可达性分布图', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig3_spatial.png'),
    ('图2：街景障碍物分布图', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig7_saii_analysis.png'),
    ('图3：三类幻觉维度示意图', 'projects\\15min-urban-accessibility\\paper\\figures\\fig1_framework.png'),
    ('图5：研究框架技术路线图', 'projects\\15min-urban-accessibility\\paper\\figures\\fig1_framework.png'),
    ('图6：时间贫困指数空间聚类图', 'projects\\15min-urban-accessibility\\paper\\figures\\fig6_time_poverty.png'),
    ('图7：综合可达性与幻觉指数双变量地图', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig3_spatial.png'),
    ('图8：四象限散点图', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig4_illusion_scatter.png'),
    ('图9：弱势群体剥夺热点图', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig6_deprived_communities.png'),
    ('图10：步行环境四维评分雷达图', 'projects\\15min-urban-accessibility\\paper\\figures\\fig5_radar.png'),
    ('图11：供需匹配三维框架图', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig9_supply_demand.png'),
    ('图12：时间贫困与幻觉指数相关性图', 'projects\\15min-urban-accessibility\\paper\\figures\\fig6_time_poverty.png'),
    ('图13：昼夜达标率对比图', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig8_day_night.png'),
    ('图14：四区障碍评分对比箱线图', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig7_saii_analysis.png'),
    ('图15：四区对比机制路径图', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig5_type_analysis.png'),
    ('图16：治理建议实施路径图', 'projects\\15min-urban-accessibility\\paper\\figures\\fig1_framework.png'),
]

def style_para(para, font_name='宋体', size=12, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False):
    para.alignment = align
    if para.runs:
        run = para.runs[0]
    else:
        run = para.add_run()
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def embed_figures(in_path, out_path):
    doc = Document(in_path)

    # Page break before figures section
    doc.add_page_break()

    # Section title
    p_title = doc.add_paragraph()
    style_para(p_title, '黑体', 16, WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    p_title.add_run('图件清单')

    doc.add_paragraph()

    total = 0
    missing = 0

    for caption, rel_path in FIGURES:
        img_path = os.path.join(BASE, rel_path)
        if not os.path.exists(img_path):
            print(f"  [MISSING] {rel_path}")
            missing += 1
            continue

        total += 1

        # Check if run.add_picture exists and works
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            run.add_picture(img_path, width=Inches(5.5))
        except AttributeError:
            # fallback: add_picture on paragraph
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(img_path, width=Inches(5.5))

        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f'图 {caption}')
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(10)
        r.font.italic = True

        # Small spacing
        doc.add_paragraph()

        fname = os.path.basename(img_path)
        print(f"  [{total}] {caption}")
        print(f"       <- {rel_path}")

    # Summary
    doc.add_paragraph()
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_note = p_note.add_run(f'（共嵌入 {total} 张图件，{missing} 张未找到对应文件）')
    r_note.font.name = '宋体'
    r_note._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r_note.font.size = Pt(10)

    print(f"\n=== Summary: {total} embedded, {missing} missing ===")
    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == '__main__':
    inp = os.path.join(BASE, '报告_final.docx')
    out = os.path.join(BASE, '报告_final_含图.docx')
    embed_figures(inp, out)
