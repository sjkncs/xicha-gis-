# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'E:\xicha gis 智能定位\报告_updated.docx')

# 删除段落65的重复标题
removed = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt == '4. 可达性幻觉的空间图谱与结构性成因':
        p._element.getparent().remove(p._element)
        removed += 1
        print(f"已删除重复标题: {txt}")

# 同时处理可能的空段落
empty_count = 0
paras_to_del = []
for p in doc.paragraphs:
    if not p.text.strip():
        empty_count += 1

print(f"空段落数: {empty_count}")

print(f"共删除 {removed} 个重复标题段落")

doc.save(r'E:\xicha gis 智能定位\报告_updated.docx')
print("已保存")

# 最终结构
print("\n=== 最终完整结构 ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    style = p.style.name
    prefix = '[H]' if 'Heading' in style else '[P]'
    if txt:
        print(f'{prefix} {i:3d} | {txt[:120]}')
