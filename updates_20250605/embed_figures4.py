# -*- coding: utf-8 -*-
"""Embed figures into the final report - simple python-docx approach"""
import sys, os
sys.path.insert(0, r'E:\xicha gis 智能定位')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r'E:\xicha gis 智能定位'

# Caption -> image path (only include files that exist)
FIGURES = [
    ('（图1：综合可达性分布图）', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig3_spatial.png'),
    ('（图2：街景障碍物分布图）', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig7_saii_analysis.png'),
    ('（图3：三类幻觉维度示意图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig1_framework.png'),
    ('（图5：研究框架技术路线图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig1_framework.png'),
    ('（图6：时间贫困指数空间聚类图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig6_time_poverty.png'),
    ('（图7：综合可达性与幻觉指数双变量地图）', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig3_spatial.png'),
    ('（图8：四象限散点图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig4_illusion_scatter.png'),
    ('（图9：弱势群体剥夺热点图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig6_deprived_communities.png'),
    ('（图10：步行环境四维评分雷达图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig5_radar.png'),
    ('（图11：供需匹配三维框架图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig9_supply_demand.png'),
    ('（图12：时间贫困与幻觉指数相关性图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig6_time_poverty.png'),
    ('（图13：昼夜达标率对比图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig8_day_night.png'),
    ('（图14：四区障碍评分对比箱线图）', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig7_saii_analysis.png'),
    ('（图15：四区对比机制路径图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig5_type_analysis.png'),
    ('（图16：治理建议实施路径图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig1_framework.png'),
]

def set_font(run, name='宋体', size=10, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.italic = italic

def add_figure(doc, img_relpath, caption, width=Inches(5.5)):
    img_abspath = os.path.join(BASE, img_relpath)
    if not os.path.exists(img_abspath):
        print(f"  [SKIP] Not found: {img_relpath}")
        return False

    # Image paragraph
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(img_abspath, width=width)

    # Caption paragraph
    cap_num = caption.replace('（', '').replace('）', '')
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run(f'图 {cap_num}')
    set_font(r, italic=True)

    print(f"  [OK] {caption} <- {img_relpath.split(chr(92))[-1]}")
    return True

def embed_figures(input_path, output_path):
    doc = Document(input_path)
    total = 0
    skipped = 0

    # Find paragraphs containing figure captions
    for i, para in enumerate(doc.paragraphs):
        for caption, img_rel in FIGURES:
            if caption in para.text:
                if add_figure(doc, img_rel, caption):
                    total += 1
                else:
                    skipped += 1

    print(f"\nFigures embedded: {total}, skipped (not found): {skipped}")
    doc.save(output_path)
    print(f"Saved to: {output_path}")

if __name__ == '__main__':
    inp = os.path.join(BASE, '报告_final.docx')
    out = os.path.join(BASE, '报告_final_含图.docx')
    embed_figures(inp, out)
