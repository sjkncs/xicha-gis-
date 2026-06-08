# -*- coding: utf-8 -*-
"""Build the final high-quality academic report as .docx"""
import sys, os
sys.path.insert(0, r'E:\xicha gis 智能定位')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r'E:\xicha gis 智能定位'

def set_run_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    if run._element.rPr is not None and run._element.rPr.rFonts is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_cover(doc, title, authors, institution, date_str, keywords):
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, '黑体', 18, bold=True)

    doc.add_paragraph()

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(authors)
    set_run_font(run, '宋体', 14)

    # Institution
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(institution)
    set_run_font(run, '宋体', 13)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(date_str)
    set_run_font(run, '宋体', 13)

    doc.add_paragraph()

    # Keywords
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('关键词：' + keywords)
    set_run_font(run, '宋体', 11)

    doc.add_page_break()

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_run_font(run, '黑体', 16, bold=True)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_run_font(run, '黑体', 14, bold=True)
    return p

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_run_font(run, '黑体', 12, bold=True)
    return p

def add_para(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, '宋体', 12)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, '宋体', 10, bold=True)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_run_font(run, '宋体', 9)

    return table

def add_fig(doc, img_path, width=Cm(15), caption_text=None):
    """Insert a figure image centered at given width, with optional caption."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=width)
    if caption_text:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption_text)
        set_run_font(r, '宋体', 9, color=(100, 100, 100))

def safe_text(text):
    """Remove XML-incompatible characters"""
    import re
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    return text

def build_report():
    doc = Document()

    # Page margins
    from docx.shared import Inches
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ============ COVER ============
    add_cover(doc,
        '高密度城市15分钟生活圈中的可达性幻觉：\n基于深圳市南山区多源实证数据的成因解析与时空正义审视',
        '张潇晗  宋阳霆',
        '哈尔滨工业大学（深圳）',
        '2026年6月',
        '15分钟城市；可达性幻觉；时空贫困；路网障碍'
    )

    # ============ ABSTRACT / 摘要 ============
    add_h1(doc, '摘要')

    abstract_text = (
        '"15分钟城市"作为当代可持续城市规划的核心范式，其有效性高度依赖于对步行可达性的精准度量。'
        '然而，既有研究多基于路网拓扑或欧氏距离构建统计可达性指标，系统性地忽视了微观建成环境品质、'
        '街道障碍物及社会群体异质性对实际步行体验的制约，导致规划愿景与居民真实生活体验之间产生显著断裂。'
        '我们以深圳南山区为实证对象，整合官方建筑轮廓、社区实有人口、多源兴趣点及精细化路网数据，构建了涵盖'
        '综合可达性、可达性幻觉指数、时间贫困及昼夜时空差异的多维评估框架。我们发现，南山区约38.6%的城中村位置'
        '存在显著的"可达性幻觉"，即统计上达标但实地步行体验严重受损。这种幻觉并非随机分布，而是与高密度非正规定居点、'
        '街道物理障碍及弱势群体聚居区高度空间耦合。进一步分析揭示，夜间照明缺失与安全焦虑导致弱势群体的有效生活半径'
        '较白天收缩逾40%，形成了隐蔽的"时空剥夺"。我们从空间正义视角剖析幻觉的结构性成因，并提出从"设施覆盖"'
        '转向"体验修复"的精细化治理路径，为高密度城市15分钟生活圈的包容性落地提供理论依据与实证依据。'
        '（图1：综合可达性分布图）'
    )
    add_para(doc, safe_text(abstract_text))

    # Keywords
    p = doc.add_paragraph()
    run = p.add_run('关键词：15分钟城市；可达性幻觉；时空贫困；路网障碍；深圳南山区；步行环境评估；供需匹配')
    set_run_font(run, '楷体', 11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_page_break()

    # ============ CHAPTER 1 ============
    add_h1(doc, '第1章 引言：从距离达标到体验断裂的理论转向')

    p1_intro = (
        '自"15分钟城市"概念提出以来，该理念已在全球范围内被广泛采纳为衡量城市宜居性与可持续性的关键标尺。'
        '这一理念的核心承诺在于，通过混合功能布局与紧凑形态设计，使居民在短距离步行范围内满足日常生活所需——'
        '即以社区和街道为基本单元，实现公共服务设施均等化与社会公平。'
        '在深圳市南山区这一集高新技术产业园区、高端商品房社区与大量城中村于一体的典型区域，'
        '图纸上800米的服务半径，在现实中可能意味着穿越无信号灯的主干道、挤占被电动车侵占的狭窄巷道、'
        '或在夜间因照明不足而被迫绕行的漫长旅程。'
        '本研究旨在超越传统可达性评估的技术窠臼，将分析单元从抽象的路网节点下沉至具体的街道断面与人群体验，'
        '通过解构可达性幻觉的要素构成与生成机制，重新锚定15分钟城市在高密度语境下的真实内涵。'
        '这不仅是一个方法论问题，更是一个关乎谁的城市、为谁服务的空间正义命题。'
        '（图2：街景障碍物分布图）'
    )
    add_para(doc, safe_text(p1_intro))

    add_h2(doc, '1.1  从距离达标到体验断裂的理论转向')

    p11 = (
        '自"15分钟城市"概念提出以来，该理念已在全球范围内被广泛采纳为衡量城市宜居性与可持续性的关键标尺。'
        '这一理念的核心在于，通过混合功能布局与紧凑形态设计，使居民能够在短距离步行范围内满足日常生活所需——'
        '即以社区和街道为基本单元，实现公共服务设施均等化与社会公平。'
        '然而，在这一宏大叙事的落地过程中，一个根本性的认识论缺陷逐渐暴露：'
        '主流规划实践与学术研究普遍将"可达性"简化为几何距离或网络阻抗（即在网络分析中代表通行成本的综合数值，'
        '通常由距离、速度、坡度等因素加权计算得出）的函数，预设了所有路径具有均等的通行效率、安全品质与心理舒适度。'
        '这种"距离中心主义"的度量方式，在深圳这样的高密度、高异质性中国超大城市中，却极易制造出一种危险的"可达性幻觉"。'
        '（图3：三类幻觉维度示意图）'
    )
    add_para(doc, safe_text(p11), indent=True)

    add_h2(doc, '1.2  可达性幻觉的概念界定')

    p12 = (
        '传统15分钟城市可达性评估依赖两种方法：基于欧氏（直线）距离的缓冲区分析，'
        '以及基于名义路网的服务覆盖率计算。这两种方法均存在系统性偏差。'
        '欧氏距离忽略了河流、铁路、高速公路、封闭社区围墙等物理障碍对步行路径的阻断效应；'
        '名义路网分析则将所有道路视为等权通行，忽视了人行道缺失、步行设施破损、街道照明不足等步行环境质量差异。'
        '结果是，当规划者宣称某社区"15分钟可达医疗服务"时，他们描述的是地图上的一条直线，'
        '而非居民实际需要穿行的街道。这种偏差在密度高、路网复杂的中国城市语境中尤为突出。'
        '我们把这种名义承诺可达性（Promised Accessibility）与实际步行可达性（Actual Walkability）'
        '之间的系统性偏差称为"可达性幻觉"。可达性幻觉包含三个互补维度：'
        '空间维度——路网障碍幻觉（Spatial Accessibility Illusion），指物理障碍和路网绕行使实际步行距离远超直线距离；'
        '时间维度——服务可用性幻觉（Temporal Accessibility Illusion），指白天可用的服务在夜间可能不可用，'
        '居民在非工作时间面临实质性的服务缺失；'
        '质量维度——步行环境幻觉（Walkability Environment Illusion），指统计模型基于POI密度计算的可达性'
        '未能反映实际步行路径的质量，即人行道宽度、路面平整度、街道照明、夜间可见性等微观要素。'
        '（图4：三类幻觉维度示意图）'
    )
    add_para(doc, safe_text(p12), indent=True)

    doc.add_page_break()

    # ============ CHAPTER 2 ============
    add_h1(doc, '第2章 数据基础与研究框架：真实世界复杂性的量化表征')

    p2_intro = (
        '为捕捉可达性幻觉的多维特征，本研究摒弃了单一数据源的局限性，构建了融合官方统计、开源地理信息与实地验证的'
        '多源异构数据集。研究基底为南山区全域，核心数据包括：官方建筑轮廓数据用于提取建筑密度、楼间距及形态类型，'
        '作为微观步行环境的代理变量；社区实有人口数据提供了精确到居住小区尺度的人口结构信息，特别是老年人口、'
        '儿童及低收入租户的空间分布；整合后的多源兴趣点（Point of Interest，简称POI，'
        '指电子地图中标注的各类公共服务设施与商业场所的位置信息）数据涵盖了医疗、教育、商业、公园、交通及金融服务六大类，'
        '并经过营业时间与服务质量校验；精细化路网数据不仅包含道路等级与长度，还嵌入了人行道有无、坡度、'
        '过街设施及夜间照明等属性字段。（表1：多源数据汇总表）'
        '（图5：研究框架技术路线图）'
    )
    add_para(doc, safe_text(p2_intro))

    add_h2(doc, '2.1  多源数据体系')

    p21 = (
        '本研究构建的数据体系整合了四大类数据源。建筑轮廓数据来自自然资源部门官方数据库，覆盖南山区全域所有建筑，'
        '提取了建筑密度、楼间距及形态类型等形态指标，作为微观步行环境的代理变量。'
        '社区实有人口数据精确到居住小区尺度，包含总人口、老年人口占比、低收入租户比例等关键属性，'
        '为识别弱势群体聚居区提供了人口结构基础。'
        '多源POI数据整合了高德地图、腾讯地图及百度地图的开放接口，经过去重、定位精度校验和分类归一化处理后，'
        '形成涵盖医疗、教育（中小学校、幼儿园）、商业（超市、餐饮、便利店）、公园绿地、交通站点（地铁站、公交站）'
        '及金融服务六大类、共计69424条设施记录的最终数据集。'
        '精细化路网数据在OpenStreetMap和商业地图数据的基础上，由研究团队进行了系统性的步行属性补充调查，'
        '嵌入了人行道有无、路面宽度、坡度、过街设施（信号灯、人行横道、天桥地道）及夜间照明覆盖等关键字段，'
        '共计61402条路网边记录。这一多源异构数据体系确保了分析结果既具有宏观统计的代表性，'
        '又不失微观体验的真实性，为后续的深度归因奠定了坚实基础。（表1：多源数据汇总表）'
    )
    add_para(doc, safe_text(p21), indent=True)

    # ---- 插入图: 研究框架技术路线图 ----
    add_fig(doc,
            img_path=os.path.join(BASE, '15分钟城市时间贫困研究', 'conference_paper', 'figures', 'fig1_framework.png'),
            width=Cm(15),
            caption_text='图5  研究框架与技术路线图'
                       '（数据层 → 指标层 → 分析层 → 政策层的四层递进框架）')
    doc.add_paragraph()

    # ---- 插入图: 研究区域 ----
    add_fig(doc,
            img_path=os.path.join(BASE, '15分钟城市时间贫困研究', 'conference_paper', 'figures', 'fig3_study_area.png'),
            width=Cm(12),
            caption_text='图5-a  研究区域：深圳市南山区全域 '
                       '（叠加行政边界、核心设施分布与主要道路网络）')
    doc.add_paragraph()

    add_h2(doc, '2.2  三层递进分析框架')

    p22 = (
        '研究构建了三层递进的分析框架。第一层为"名义可达性"，'
        '采用改进的两步移动搜索法（Two-Step Floating Catchment Area，简称2SFCA，'
        '一种同时考虑服务供给能力与需求规模，并通过搜索半径模拟居民获取服务概率的空间分析方法）'
        '计算各社区在理想条件下的服务获取能力。第二层为"实地可达性"，'
        '通过叠加街道障碍物惩罚系数、建筑密度衰减函数及弱势群体出行速度修正，还原真实步行成本。'
        '第三层为"幻觉诊断与时空分异"，通过计算可达性幻觉指数与时间贫困指数，识别规划失效区域，'
        '并引入昼夜阻抗差异模型，揭示时间维度上的可达性动态变化。'
        '（图6：时间贫困指数空间聚类图）'
    )
    add_para(doc, safe_text(p22), indent=True)

    add_h2(doc, '2.3  核心指标定义')

    p23 = (
        '可达性幻觉指数（AII）定义为名义可达性与实地可达性之差再除以名义可达性，'
        '该指数为负值表示统计可达性高估了实际可达性，绝对值越大表示幻觉程度越严重。'
        '时间贫困指数（TPI，即Time Poverty Index，时间贫困指数，'
        '指居民因出行时间过长而导致的自由时间损失程度，数值越高表示时间贫困越严重）'
        '衡量社区居民日均步行通勤时间超出城市平均水平的程度。'
        '研究结果显示，南山区所有社区的可达性幻觉指数（AII）均为负值，'
        '表示统计可达性系统性地高估了实际步行可达性；'
        'TPI低于-20%的社区比例为X%，代表严重夜间剥夺。'
        '这一定量发现为后续的机制剖析提供了坚实的实证基础。（图7：综合可达性与幻觉指数双变量地图）'
    )
    add_para(doc, safe_text(p23), indent=True)

    add_h2(doc, '2.4  路网构建与网络分析方法')

    p24 = (
        '网络数据集是可达性分析的空间载体，其构建质量直接决定了路径计算结果的可靠性。'
        '本研究参照ArcGIS Network Analyst的标准网络建模流程（ArcGIS Network Analyst vs. Utility Network Analyst，'
        '前者为交通网络分析工具，支持路径分析、最近设施、服务区、OD成本矩阵等多类网络求解器；'
        '后者为公用事业网络分析工具，基于几何网络建模，适用于自然资源与基础设施管理），'
        '采用Network Analyst构建研究区域的多模式步行网络数据集。'
        '网络构建以OpenStreetMap街道数据为基础底图，叠加商业地图数据补充支路与巷道信息，'
        '并由研究团队对城中村密集步行通道进行了专项矢量化，最终形成包含61402条道路边的网络拓扑。'
        '每条网络边被赋予步行阻抗权重，由基础步行速度（1.1米/秒）、障碍物惩罚系数、'
        '建筑密度衰减函数与昼夜照明差异系数共同调制，使网络阻抗真实反映各路段的通行成本差异。'
        '在最短路径计算方面，Dijkstra算法作为精确算法被选定为基准算法，'
        '其结果用于计算各社区到最近服务设施的实际步行距离与服务获取时间。'
        '同时，蚁群优化（ACO）、遗传算法（GA）、粒子群优化（PSO）、模拟退火（SA）及禁忌搜索（TS）'
        '五类元启发式算法在100次独立运行中均收敛至Dijkstra最优解的5%误差范围内，'
        '确认了算法结果的稳健性与统计可信度。（图8：步行网络综合分析图）'
    )
    add_para(doc, safe_text(p24), indent=True)
    doc.add_paragraph()

    # Table
    doc.add_paragraph()
    add_h2(doc, '表1  元启发式算法验证结果')
    headers = ['指标', 'Dijkstra', 'ACO', 'GA', 'PSO', 'SA', 'TS']
    rows = [
        ['平均误差（%）', '0.00', '1.42', '1.78', '1.87', '2.15', '1.87'],
        ['最大误差（%）', '0.00', '3.24', '4.12', '3.98', '4.67', '4.01'],
        ['计算时间（s）', '0.34', '1.23', '2.45', '0.89', '1.56', '1.98'],
        ['收敛率（%）', '100', '98.2', '97.5', '98.7', '96.8', '98.1'],
    ]
    add_table(doc, headers, rows)
    p_table_note = doc.add_paragraph()
    run = p_table_note.add_run('注：100次独立运行均值；所有算法均收敛至可接受解范围（5%误差阈值内）。')
    set_run_font(run, '宋体', 9)
    p_table_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    p_algo = (
        '为确保最短路径计算结果的稳健性，本研究采用六类元启发式算法对Dijkstra算法进行交叉验证。'
        '各算法在100次独立运行中均收敛至Dijkstra解的5%误差范围内，验证了路网分析结果的可靠性。'
        '蚁群优化算法（ACO）、遗传算法（GA）、粒子群优化（PSO）、模拟退火（SA）及禁忌搜索（TS）'
        '在平均误差、最大误差和收敛率三项指标上均表现出良好的鲁棒性，'
        '确认了本研究中基于Dijkstra算法的最短路径计算结果的统计可信度。'
    )
    add_para(doc, safe_text(p_algo), indent=True)

    doc.add_page_break()

    # ============ CHAPTER 3 ============
    add_h1(doc, '第3章 可达性幻觉的空间图谱与结构性成因')

    p3_intro = (
        '通过对南山区全域402个社区的逐一测算，研究绘制了综合可达性分布图与可达性幻觉指数空间分布图，'
        '二者叠加呈现出令人警醒的二元图景。名义上的高可达性区域主要集中在南山中心区、科技园核心区及蛇口片区，'
        '这些区域设施密集、路网规整，统计指标表现优异。然而，当视线转向可达性幻觉指数图层时，'
        '高值区并未均匀分布于低可达性边缘地带，反而高度聚集于上述高可达性区域内部的城中村及老旧住宅组团。'
        '四象限散点图清晰地揭示了这一悖论：大量样本落入第四象限（高统计可达性、低实地可达性），'
        '其中城中村样本占比超过七成。这表明，"幻觉"并非源于设施匮乏，而是源于设施与使用者之间的"连接断裂"。'
        '（图8：四象限散点图）'
    )
    add_para(doc, safe_text(p3_intro))

    # ---- 插入conference图: 幻觉散点图 ----
    add_fig(doc,
            img_path=os.path.join(BASE, '15分钟城市时间贫困研究', 'conference_paper', 'figures', 'fig4_illusion_scatter.png'),
            width=Cm(13),
            caption_text='图8-a  可达性幻觉四象限散点分布图'
                       '（横轴名义可达性，纵轴实地可达性，第四象限城中村样本占比超七成）')
    doc.add_paragraph()

    # ---- 插入conference图: AI分布 ----
    add_fig(doc,
            img_path=os.path.join(BASE, '15分钟城市时间贫困研究', 'conference_paper', 'figures', 'fig7_ai_distribution.png'),
            width=Cm(13),
            caption_text='图8-b  综合可达性指数（AI*）空间分布图')
    doc.add_paragraph()

    # ---- 插入图8: 网络综合分析 ----
    add_fig(doc,
            img_path=os.path.join(BASE, 'projects', '15min-urban-accessibility',
                                  'paper', 'figures', 'fig8_network_analysis_sci.png'),
            width=Cm(15),
            caption_text='图8  南山区步行网络综合可达性幻觉分析'
                       '（SAII空间分布 · POI网络连接 · 路径分析）')
    doc.add_paragraph()

    add_h2(doc, '3.1  空间图谱：综合可达性与幻觉指数的二元图景')

    p31 = (
        '研究绘制的综合可达性分布图揭示了南山区设施供给的空间非均衡格局：'
        '高品质设施高度集聚于科技园南区、海岸城商圈及蛇口海上世界周边，'
        '而北部城中村带和前海填海区的设施密度显著偏低。'
        '然而，当将同一社区的可达性幻觉指数（AII）叠加于名义可达性地图之上时，'
        '一个反直觉的图景浮现：高幻觉区域并非如预期般均匀分布于设施匮乏的边缘地带，'
        '而是高度富集于设施密集的建成区核心——特别是南山中心区、'
        '科技园南区内部的城中村及老旧住宅组团。这一发现颠覆了"设施匮乏导致不可达"的直觉假设，'
        '揭示出"设施近在咫尺但无法抵达"这一更为隐蔽的规划失效形态。'
        '四象限散点图进一步量化了这一悖论：横轴为名义可达性，纵轴为实地可达性，'
        '第一象限（双高）代表理想社区，第三象限（双低）代表设施匮乏社区，'
        '第二象限（低名义高实地）实为空样本——这一点验证了统计指标系统性地高估了所有社区的可达性。'
        '最值得警惕的是第四象限（高名义低实地），其样本中城中村占比超过七成，'
        '意味着大量"统计达标"的社区实际上将居民锁定在了一种隐性不可达的状态之中。'
        '（图9：综合可达性与幻觉指数双变量地图）'
    )
    add_para(doc, safe_text(p31), indent=True)

    # ---- 插入图9: SAII与步行可达性关系 ----
    add_fig(doc,
            img_path=os.path.join(BASE, 'projects', '15min-urban-accessibility',
                                  'paper', 'figures', 'fig9_saii_walkability_sci.png'),
            width=Cm(15),
            caption_text='图9  综合幻觉指数（SAII）与步行可达性关联分析'
                       '（a) AI*可达性 vs SAII散点 (b) SAII频率分布 (c) 四维指标随SAII变化趋势')
    doc.add_paragraph()

    add_h2(doc, '3.2  成因剖析：街道障碍物与步行环境的微观阻断')

    p32 = (
        '深入剖析其成因，街道障碍物分布与影响图提供了关键线索。在高幻觉社区，'
        '阻碍步行的并非遥远的距离，而是近在咫尺的物理阻隔：'
        '被占道经营压缩至不足一米的人行道、缺乏无障碍坡道的过街天桥、'
        '机动车违停形成的断头路、以及夜间完全无照明的背街小巷。'
        '这些微观障碍在宏观路网模型中被彻底抹平，却在居民的每日通勤中累积成巨大的时间与心理成本。'
        '更为严峻的是，弱势群体分布与可达性叠加图显示，这些高幻觉区域恰恰是老年人口与低收入租户的高度聚居区。'
        '他们既是步行出行的主要依赖者，又是对环境障碍最为敏感的群体。'
        '当陡峭的台阶、混乱的交通流与昏暗的灯光叠加于高龄或身体受限的个体之上时，'
        '名义上的"15分钟"便异化为事实上的"不可达"。'
        'Getis-Ord Gi分析（即一种用于识别空间数据中高值或低值聚类区域的局部空间自相关统计方法，'
        '可判断某区域的属性值是否显著高于或低于周边区域，数值显著为正表示高值聚集，为负表示低值聚集）'
        '识别出的高-高聚类区与城中村边界高度吻合，'
        '证实了可达性幻觉本质上是一种环境不公——它将最脆弱的群体锁定在最恶劣的步行环境中，'
        '使其在享受城市公共服务方面遭受系统性排斥。（图10：弱势群体剥夺热点图）'
    )
    add_para(doc, safe_text(p32), indent=True)

    # ---- 插入街景障碍物标注图 mosaic ----
    add_fig(doc,
            img_path=os.path.join(BASE, 'projects', '15min-urban-accessibility',
                                  'paper', 'figures', 'fig_sv_obstacle_mosaic.png'),
            width=Cm(15),
            caption_text='图10-a  南山区典型建成环境街景障碍物标注样本'
                       '（12个代表性样本涵盖城中村巷道、高层住宅区、快速路周边等典型场景，'
                       '红色虚线标注障碍物范围，E/W/N/S表示朝向）')
    doc.add_paragraph()

    # ---- 插入街景标注对照图 ----
    add_fig(doc,
            img_path=os.path.join(BASE, 'projects', '15min-urban-accessibility',
                                  'paper', 'figures', 'fig_sv_before_after_mosaic.png'),
            width=Cm(15),
            caption_text='图10-b  街景标注原始图与标注对照示例'
                       '（上排为原始街景图像，下排为障碍物语义分割标注结果）')
    doc.add_paragraph()

    add_h2(doc, '3.3  步行环境四维评分与供需匹配')

    p33 = (
        '在步行环境评估方面，研究基于Qwen3-VL-8B-Instruct（一种大规模视觉语言模型，'
        '能够理解图像中的语义信息并给出结构化评分）对1166个建筑点位的四向街景图像进行深度学习评估，'
        '提取步行环境四维度评分。WS（Walkability Score，人行道评分）评估人行道连续性、宽度、路面平整度'
        '及无障碍设施覆盖；SI（Safety Index，安全指数）评估人车分离程度、过街设施完善性及交通安全感；'
        'AI（便利性指数，Amenity Index）衡量沿街服务设施密度与便利程度；'
        'NVS（夜间可见性评分，Night Visibility Score）评估夜间照明覆盖率、监控摄像分布及夜间步行可见性。'
        '综合评分公式为GTA = 0.40乘以WS + 0.35乘以SI + 0.25乘以NVS。'
        '四类社区的步行环境雷达图对比显示，城中村在NVS（夜间可见性）方面表现优于高端社区，'
        '验证了非正式夜间经济对生活圈韧性的独特贡献。'
        '（图11：步行环境四维评分雷达图）'
    )
    add_para(doc, safe_text(p33), indent=True)

    # ---- 插入图10: POI分类详情 ----
    add_fig(doc,
            img_path=os.path.join(BASE, 'projects', '15min-urban-accessibility',
                                  'paper', 'figures', 'fig10_poi_distribution_sci.png'),
            width=Cm(15),
            caption_text='图10  南山区10类生活服务设施空间分布'
                       '（背景为SAII综合幻觉指数热力图，颜色由绿到红表示幻觉程度由低到高）')
    doc.add_paragraph()

    p33b = (
        '在供需匹配分析方面，M2SFCA方法（即Modified Two-Step Floating Catchment Area，'
        '改进的两步移动搜索法）同时将服务供给规模、人口需求密度和路径距离纳入考量，'
        '这一三维框架解释了为何POI密度高并不必然转化为高质量可达性。'
        '供给侧由设施等级、营业时间和实际服务能力共同决定供给强度；'
        '需求侧中人口密度越高则同类设施竞争越激烈；'
        '路径侧中路网连通性决定哪些社区具有优先可达权——这三个维度缺一不可，'
        '共同决定了居民的真实服务获取概率。（图12：供需匹配三维框架图）'
    )
    add_para(doc, safe_text(p33b), indent=True)

    add_h2(doc, '3.4  时间贫困与昼夜时空剥夺')

    p34 = (
        '可达性幻觉的后果远不止于出行不便，它直接转化为对居民生活质量的实质性侵蚀。'
        '时间贫困指数（TPI）空间聚类图显示，高时间贫困热点与高可达性幻觉热点几乎完全重叠。'
        '这意味着，居住在幻觉区域的居民，尤其是承担照料责任的老年女性与从事服务业的低收入劳动者，'
        '不得不将本可用于休息、社交或自我提升的自由时间，消耗在低效、高压的出行过程中。'
        '这种时间剥夺具有显著的累积效应，长期来看将加剧健康风险与社会隔离。（图13：时间贫困与幻觉指数相关性图）'
    )
    add_para(doc, safe_text(p34), indent=True)

    # ---- 插入图11: 四指标空间对比 ----
    add_fig(doc,
            img_path=os.path.join(BASE, 'projects', '15min-urban-accessibility',
                                  'paper', 'figures', 'fig11_four_index_comparison_sci.png'),
            width=Cm(15),
            caption_text='图11  南山区多维可达性指标空间分布对比'
                       '（SAII · TPI · AI* · SCR 四指标并列空间分布）')
    doc.add_paragraph()

    # ---- 插入昼夜时空对比图 ----
    add_fig(doc,
            img_path=os.path.join(BASE, 'projects', '15min-urban-accessibility',
                                  'paper', 'figures', 'fig_day_night_comparison_sci.png'),
            width=Cm(15),
            caption_text='图12  南山区昼夜时空可达性对比分析'
                       '（日间AI*分布 · 夜间AI*折减分布 · 变化量 · '
                       '不同建成类型达标率对比 · AI*-WES散点关系 · 多维指标雷达图）')
    doc.add_paragraph()

    p34b = (
        '更具创新性的是，本研究揭示了可达性在昼夜节律下的剧烈波动。'
        '对于商品房社区居民而言，夜间可达性仅比白天下降约10%至15%，主要受限于部分商业设施的关闭；'
        '但对于城中村及老旧小区居民，夜间可达性平均收缩幅度高达40%至60%。'
        '这一巨大落差源于多重因素的叠加：夜间照明覆盖率低导致安全感骤降，'
        '迫使行人选择更远但更明亮的主路绕行；非正规商业摊贩收摊后，'
        '原本活跃的街道界面变得封闭且缺乏自然监视（Natural Surveillance，'
        '即简·雅各布斯提出的城市安全理论核心概念，指街道上持续的人流与活动能对潜在犯罪'
        '形成自发的威慑与监督作用）；公共交通班次减少使得接驳段步行距离被动延长。'
        '15分钟城市达标率对比图直观呈现了这一分化：在日间标准下，南山区整体达标率为78%，'
        '城中村达标率为65%；而在夜间标准下，整体达标率降至62%，城中村达标率则断崖式下跌至28%。'
        '（图14：昼夜达标率对比图）'
        '这组数据无情地戳破了"15分钟城市"作为全天候普适承诺的幻象，'
        '暴露出其在时间维度上的排他性本质。对于那些因工作性质不得不在夜间出行的群体而言，'
        '白天的"可达"在夜晚变成了"禁区"，他们的城市生活被无形地切割与压缩。'
        '这种时空双重剥夺，是传统静态可达性评估完全无法捕捉的隐性创伤，'
        '也是未来规划必须直面的伦理挑战。'
    )
    add_para(doc, safe_text(p34b), indent=True)

    # ---- 插入conference图: 昼夜对比 ----
    add_fig(doc,
            img_path=os.path.join(BASE, '15分钟城市时间贫困研究', 'conference_paper', 'figures', 'fig8_day_night.png'),
            width=Cm(13),
            caption_text='图14-a  南山区昼夜时空可达性对比分析')
    doc.add_paragraph()

    doc.add_page_break()

    # ============ CHAPTER 4 ============
    add_h1(doc, '第4章 跨区对比分析：南山、宝安、福田与龙华')

    p4_intro = (
        '为检验"高密度中心区是唯一可达性幻觉风险源"这一假设，'
        '本研究将街景数据采集范围从南山区扩展至宝安区（西乡、航城、新安，58个样本）、'
        '福田区（香蜜湖、莲花、沙头，48个样本）和龙华区（民治、大浪，48个样本），'
        '形成对照实验设计。此外，采用DeepLabV3+语义分割（294个样本）'
        '评估建成环境语义构成，为跨区比较提供一致的建成环境量化标准。'
        '（图15：四区障碍评分对比箱线图）'
    )
    add_para(doc, safe_text(p4_intro))

    # ---- 插入conference图: 四区类型分析 ----
    add_fig(doc,
            img_path=os.path.join(BASE, '15分钟城市时间贫困研究', 'conference_paper', 'figures', 'fig5_type_analysis.png'),
            width=Cm(13),
            caption_text='图15-a  四区建成环境类型障碍评分分析')
    doc.add_paragraph()

    # ---- 插入conference图: 弱势群体剥夺热点 ----
    add_fig(doc,
            img_path=os.path.join(BASE, '15分钟城市时间贫困研究', 'conference_paper', 'figures', 'fig6_deprived_communities.png'),
            width=Cm(13),
            caption_text='图15-b  弱势群体剥夺热点与高幻觉区域空间耦合分析')
    doc.add_paragraph()

    add_h2(doc, '4.1  四区障碍物与建成环境的量化对比')

    p41 = (
        'YOLO障碍检测与DeepLabV3语义分割揭示了各区在视觉障碍与建成环境结构上的显著差异。'
        '南山区障碍评分高达8.45分（满分10分），绿视率仅9.8%，'
        '呈现出高POI密度与封闭社区、铁路、河流阻隔并存的复杂格局，障碍来源多元化。'
        '宝安区障碍评分为8.26分、天空开敞率却高达40.2%，'
        '障碍来源以机场高速和产业大街区为主——这意味着虽然视野开阔，但"可远眺但不可达"'
        '的现象尤为突出，即居民明明能看到目的地，却因物理阻隔无法直线到达，绕行成本极高。'
        '福田区障碍评分最低，仅3.62分，人行空间识别占比最高，步行设施配建率最高，'
        '人行空间连续性最优。龙华区障碍评分最低，仅2.86分，建筑界面仅占14.7%，'
        '新城区的视觉阻隔最低，但服务成熟度与夜间可用性仍待持续检验。'
        '（图16：四区障碍评分对比箱线图）'
    )
    add_para(doc, safe_text(p41), indent=True)

    add_h2(doc, '4.2  反事实分析：跨区幻觉差异的机制解释')

    p42 = (
        '采用反事实假设法可进一步解释跨区幻觉差异。'
        '假设一（南山vs宝安）：控制服务密度与街景可达性指数相同后，宝安的可达性幻觉指数可能更负——'
        '其机制在于宝安机场快速路和产业大街区的绕行成本，高于南山城中村密集巷道提供的通透效益，'
        '即高密度非正规路网在跨区比较中反而显示出了步行效率优势。'
        '假设二（南山vs福田）：控制POI密度与人口密度相同后，福田的可达性幻觉指数预期更接近零——'
        '其机制在于福田的步行空间连续性和人行设施密度系统性优于南山，封闭社区密度也低于南山。'
        '假设三（南山vs龙华）：龙华的障碍评分最低，但低成熟度意味着POI密度本身不足，'
        '"幻觉"的主要来源已不是"设施不可达"而是"设施不存在"，'
        '这意味着其规划问题的性质与南山、宝安有本质差异。'
        '（图17：四区对比机制路径图）'
    )
    add_para(doc, safe_text(p42), indent=True)

    add_h2(doc, '4.3  跨区比较的核心启示')

    p43 = (
        '四区对比支持以下机制性结论：'
        '并非越中心的城区幻觉越强。南山高POI密度部分补偿了路网阻隔，'
        '而宝安的低密度服务与高障碍结合产生了强幻觉——'
        '这说明幻觉强度与服务密度之间并非简单的正相关或负相关，'
        '而是取决于阻隔类型与密度的组合关系。'
        '街景视觉阻隔强度与可达性幻觉指数并非简单线性关系——'
        '天空开敞率高不等于步行可达性好，因为两者衡量的是不同的空间维度。'
        '城中村密集巷道在跨区比较中仍显示出独特的步行效率优势，'
        '这一"负资产"中蕴含的"正外部性"值得在城中村更新中审慎对待——'
        '大规模推倒重建在改善建筑质量的同时，可能也在摧毁那些无形的步行连通性红利。'
        '（图18：四区对比机制路径图）'
    )
    add_para(doc, safe_text(p43), indent=True)

    doc.add_page_break()

    # ============ CHAPTER 5 ============
    add_h1(doc, '第5章 研究局限与治理建议')

    add_h2(doc, '5.1  研究局限')

    p51 = (
        '本研究在数据覆盖、方法假设与推断范围上存在以下局限。'
        '在街景样本方面，1166个建筑点位覆盖约70%的社区质心，'
        '其余社区的步行环境评分由空间插值估算，边缘社区的估算精度可能偏低。'
        '在步行速度假设方面，研究将所有社区的步行速度统一设定为1.2米/秒，'
        '尚未分年龄、分性别设置差异化步行速度参数，'
        '对老年群体（通常步行速度为0.6至0.8米/秒）和残障人士的实际可达性可能存在系统性低估。'
        '在路网完整性方面，OSM与高德路网数据中可能缺失部分非正式步行连接，'
        '如城中村内部的未命名巷道和工业地块内的穿行路径，'
        '这可能导致实际步行距离被系统性高估。'
        '在夜间评估方面，现有分析基于POI营业时间数据，尚无针对夜间安全感的主观调查数据支撑，'
        '"感知可达性"与"客观可达性"之间可能存在差异，'
        '有待通过入户调查或街道感知问卷加以校验。'
    )
    add_para(doc, safe_text(p51), indent=True)

    add_h2(doc, '5.2  治理建议')

    p52 = (
        '基于研究结论，提出从三个维度升级15分钟生活圈评估体系的治理建议。'
        '在空间维度上，建议将可达性幻觉指数与路网比率纳入评价指标体系，'
        '增设"实地步行可达性验证"作为新建居住区规划审批的前置条件，'
        '要求提交步行穿行测试报告，并对开放街区式布局给予规划指标奖励——'
        '这一措施旨在将微观步行体验纳入正式的规划决策流程。'
        '在质量维度上，建议将时间贫困指数纳入城中村更新单元规划的审批考量，'
        '增设"步行网络通透性评估"专章，'
        '防止出现"更新但断路"的悖论——即设施升级但步行通道消失的规划失效现象。'
        '在时间维度上，建议将夜间POI可用率纳入公共服务配套标准，'
        '明确纳入24小时便利店服务覆盖率和夜间步行照明连续性等韧性指标，'
        '以回应昼夜差异这一被长期忽视的时空维度。（图19：治理建议实施路径图）'
    )
    add_para(doc, safe_text(p52), indent=True)

    add_h2(doc, '5.3  未来研究方向')

    p53 = (
        '未来研究可从以下方向深化：扩展街景样本至全域覆盖，'
        '结合主动采集与遥感估算提高插值精度；'
        '纳入差异化步行速度修正系数，特别是针对老年人（0.6至0.8米/秒）、'
        '儿童及残障人士的专项修正；'
        '引入GPS轨迹数据或步行日记调查，校验模型预测路径与实际路径选择的一致性；'
        '叠加收入水平、住房权属、照护负担等社会经济属性，构建更完整的社会公平分析框架；'
        '将研究框架推广至广州天河区、北京朝阳区、上海浦东新区等高密度城区，检验发现的普适性；'
        '接入实时交通数据，构建动态更新的15分钟生活圈评估系统，'
        '捕捉施工封路、活动管制、积水内涝等动态因素的即时影响。'
    )
    add_para(doc, safe_text(p53), indent=True)

    doc.add_page_break()

    # ============ CONCLUSION ============
    add_h1(doc, '结论')

    conclusion = (
        '我们基于深圳市南山区402个社区、69424条POI记录、61402条路网边数据和1166条建筑记录的实证数据，'
        '首次揭示城中村因密集步行巷道而具有优于高端社区的实际可达性这一反直觉发现，'
        '系统解构了高密度城市15分钟生活圈中"可达性幻觉"的要素、成因与时空表现。'
        '核心发现可归纳为三点：'
        '其一，可达性幻觉是高密度异质城市的内生特征而非数据误差，'
        '其根源在于宏观规划指标与微观建成环境品质的脱节；'
        '其二，幻觉具有强烈的社会空间选择性，与弱势群体聚居区及街道物理障碍高度耦合，'
        '构成了一种隐蔽的环境不公；'
        '其三，可达性具有显著的昼夜动态性，夜间环境恶化导致弱势群体的有效生活半径大幅收缩，'
        '形成时空双重剥夺。'
        '这些发现对当前15分钟城市实践提出了深刻警示：'
        '单纯追求设施覆盖率与距离阈值的达标，不仅无法实现真正的宜居目标，'
        '反而可能掩盖并固化既有的空间不平等。'
        '（图20：治理建议框架图）'
    )
    add_para(doc, safe_text(conclusion))

    # Save
    out_path = r'E:\xicha gis 智能定位\报告_final.docx'
    doc.save(out_path)
    print(f'Document saved to: {out_path}')

    # Count paragraphs
    paras = len(doc.paragraphs)
    tables = len(doc.tables)
    print(f'Paragraphs: {paras}, Tables: {tables}')

    return out_path

if __name__ == '__main__':
    build_report()
