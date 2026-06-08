# -*- coding: utf-8 -*-
"""Embed figures using python-docx native API only"""
import sys, os
sys.path.insert(0, r'E:\xicha gis 智能定位')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

BASE = r'E:\xicha gis 智能定位'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

FIGURES = [
    ('（图1：综合可达性分布图）', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig3_spatial.png'),
    ('（图2：街景障碍物分布图）', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig7_saii_analysis.png'),
    ('（图3：三类幻觉维度示意图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig1_framework.png'),
    ('（图5：研究框架技术路线图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig1_framework.png'),
    ('（图6：时间贫困指数空间聚类图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig6_time_poverty.png'),
    ('（图7：综合可达性与幻觉指数双变量地图）', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig3_spatial.png'),
    ('（图8：四象限散点图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig4_illusion_scatter.png'),
    ('（图9：弱势群体剥夺热点图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig6_deprived_communities.png'),
    ('（图10：步行环境四维评分雷达图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig5_radar.png'),
    ('（图11：供需匹配三维框架图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig9_supply_demand.png'),
    ('（图12：时间贫困与幻觉指数相关性图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig6_time_poverty.png'),
    ('（图13：昼夜达标率对比图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig8_day_night.png'),
    ('（图14：四区障碍评分对比箱线图）', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig7_saii_analysis.png'),
    ('（图15：四区对比机制路径图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig5_type_analysis.png'),
    ('（图16：治理建议实施路径图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig1_framework.png'),
]

def style_caption_para(para, text):
    """Style a caption paragraph using python-docx API"""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)
    run.font.italic = True

def embed_figures(input_path, output_path):
    doc = Document(input_path)
    body = doc.element.body

    # First pass: collect all caption matches
    matches = []  # (para_idx, caption, img_rel)
    for i, para in enumerate(doc.paragraphs):
        for caption, img_rel in FIGURES:
            if caption in para.text:
                img_path = os.path.join(BASE, img_rel)
                matches.append((i, caption, img_path))
                print(f"Found [{i}]: {caption}")
                break

    print(f"\nMatched {len(matches)} captions")

    # Second pass: add images/captions at the end, collecting their XML elements
    fig_elems = []  # (after_idx, [elem1, elem2])

    for after_idx, caption, img_path in matches:
        if not os.path.exists(img_path):
            print(f"  SKIP: not found {img_path}")
            continue

        cap_num = caption.replace('（', '').replace('）', '')
        cap_text = f'图 {cap_num}'

        # Add image paragraph
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_img = p_img.add_run()
        r_img.add_picture(img_path, width=Inches(5.5))
        img_elem = p_img._element

        # Add caption paragraph
        p_cap = doc.add_paragraph()
        style_caption_para(p_cap, cap_text)
        cap_elem = p_cap._element

        # Collect elements for insertion
        fig_elems.append((after_idx, [img_elem, cap_elem]))
        print(f"  Added: {caption} <- {os.path.basename(img_path)}")

    # Third pass: move elements to correct positions (in reverse order)
    # The image/caption paragraphs are at the end of body
    # We need to move each to just after its caption paragraph
    body_children = list(body)

    for after_idx, (img_elem, cap_elem) in reversed(fig_elems):
        children = list(body)
        # Find the caption paragraph index in current body
        try:
            target_idx = list(body).index(children[after_idx] if after_idx < len(children) else body[-1])
        except:
            target_idx = after_idx

        # Remove img and cap from end
        body.remove(img_elem)
        body.remove(cap_elem)

        # Insert after target
        # Find target paragraph (the one containing the caption text)
        # Try to find by matching text content
        target_p = None
        for j, child in enumerate(list(body)):
            p_texts = []
            for t in child.iter(f'{{{W_NS}}}t'):
                if t.text:
                    p_texts.append(t.text)
            if any('图' in ''.join(p_texts) and '15分钟' in ''.join(p_texts) or
                    '幻觉' in ''.join(p_texts) or '可达性' in ''.join(p_texts) for _ in [1]):
                # Check if it's the right paragraph by looking for specific text
                full_text = ''.join(p_texts)
                for caption, _ in FIGURES:
                    if caption in full_text:
                        target_p = child
                        target_idx = j
                        break
            if target_p:
                break

        body.insert(target_idx + 1, img_elem)
        body.insert(target_idx + 2, cap_elem)

    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    print(f"Total figures: {len(fig_elems)}")

if __name__ == '__main__':
    inp = os.path.join(BASE, '报告_final.docx')
    out = os.path.join(BASE, '报告_final_含图.docx')
    embed_figures(inp, out)
