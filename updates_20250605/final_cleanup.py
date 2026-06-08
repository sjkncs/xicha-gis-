# -*- coding: utf-8 -*-
import sys, re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'E:\xicha gis 智能定位\报告_updated.docx')
paras_to_remove = []

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if not txt:
        continue
    chinese = sum(1 for c in txt if '\u4e00' <= c <= '\u9fff')
    readable = sum(1 for c in txt if (32 <= ord(c) <= 126) or ('\u4e00' <= c <= '\u9fff'))

    # 乱码段落：可读字符比例很低
    if readable < len(txt) * 0.4 and len(txt) > 10:
        paras_to_remove.append(p)
        print(f"删除乱码段落 {i}: {txt[:60]}")
        continue

    # 哈尔滨工业大学工学博士学位论文
    if '哈尔滨工业大学工学博士学位论文' in txt:
        paras_to_remove.append(p)
        print(f"删除页眉段落 {i}: {txt[:60]}")
        continue

    # [3][4][5] 参考文献混入结论
    if re.match(r'\[\d+\]\s', txt) and i > 40:
        paras_to_remove.append(p)
        print(f"删除文献段落 {i}: {txt[:60]}")
        continue

# 去重
seen = set()
for p in paras_to_remove:
    pid = id(p._element)
    if pid not in seen:
        seen.add(pid)
        try:
            p._element.getparent().remove(p._element)
        except:
            pass

doc.save(r'E:\xicha gis 智能定位\报告_updated.docx')
print("\n已保存最终文档")

# 最终结构
print("\n=== 最终文档结构 ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt:
        style = p.style.name
        prefix = '[H]' if 'Heading' in style else '[P]'
        print(f'{prefix} {i:3d} | {txt[:120]}')
