from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
TEMPLATE_DOC = ROOT / "1.2哈尔滨工业大学（深圳）全日制硕士研究生学位论文书写范例(1).doc"
SOURCE_REPORT = ROOT / "报告_final.docx"
OUTPUT_DOCX = ROOT / "报告_final_哈工大模板.docx"
TEMP_DIR = ROOT / "_hit_template_work"
TEMP_TEMPLATE_DOCX = TEMP_DIR / "hit_template_converted.docx"
TEMP_BODY_DOCX = TEMP_DIR / "report_body_only.docx"


WD_FORMAT_DOCX = 16
WD_DO_NOT_SAVE = 0
WD_ALIGN_LEFT = 0
WD_ALIGN_CENTER = 1
WD_ALIGN_JUSTIFY = 3
WD_WITHIN_TABLE = 12
WD_COLOR_BLACK = 0


def normalize_text(text: str) -> str:
    return text.replace("\r", "").replace("\x07", "").strip()


def collect_source_text() -> tuple[list[str], str]:
    doc = Document(SOURCE_REPORT)
    abstract_started = False
    abstract_parts: list[str] = []
    keywords = ""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text == "摘要":
            abstract_started = True
            continue
        if abstract_started and text.startswith("关键词"):
            keywords = text
            break
        if abstract_started:
            abstract_parts.append(text)
    if not abstract_parts:
        raise RuntimeError("未能从报告_final.docx提取摘要正文。")
    if not keywords:
        keywords = "关键词：15分钟城市；可达性幻觉；时空贫困；路网障碍；深圳南山区；步行环境评估"
    return abstract_parts, keywords


def find_paragraph_start(doc, pattern: str, after_start: int = 0, require_heading: bool = False) -> int:
    regex = re.compile(pattern)
    for paragraph in doc.Paragraphs:
        if paragraph.Range.Start < after_start:
            continue
        text = normalize_text(paragraph.Range.Text)
        if not regex.match(text):
            continue
        if require_heading:
            style_name = str(paragraph.Style.NameLocal)
            if "标题" not in style_name and "Heading" not in style_name:
                continue
        return paragraph.Range.Start
    raise RuntimeError(f"未找到匹配段落：{pattern}")


def find_paragraph(doc, pattern: str, after_start: int = 0):
    regex = re.compile(pattern)
    for paragraph in doc.Paragraphs:
        if paragraph.Range.Start < after_start:
            continue
        text = normalize_text(paragraph.Range.Text)
        if regex.match(text):
            return paragraph
    return None


def safe_style(doc, name: str):
    try:
        return doc.Styles(name)
    except Exception:
        return None


def set_range_font(range_obj, size: float | None = None, bold: bool | None = None) -> None:
    range_obj.Font.NameFarEast = "宋体"
    range_obj.Font.NameAscii = "Times New Roman"
    range_obj.Font.NameOther = "Times New Roman"
    range_obj.Font.Color = WD_COLOR_BLACK
    if size is not None:
        range_obj.Font.Size = size
    if bold is not None:
        range_obj.Font.Bold = -1 if bold else 0


def replace_first_matching_paragraph(doc, pattern: str, replacement_lines: list[str], style_name: str | None = None) -> None:
    paragraph = find_paragraph(doc, pattern)
    if paragraph is None:
        raise RuntimeError(f"未找到模板占位段落：{pattern}")
    rng = paragraph.Range
    rng.Text = "\r".join(replacement_lines) + "\r"
    inserted = doc.Range(rng.Start, rng.Start + len("\r".join(replacement_lines)) + len(replacement_lines))
    if style_name:
        style = safe_style(doc, style_name)
        if style is not None:
            inserted.Style = style
    set_range_font(inserted, size=12)


def create_body_doc(word) -> None:
    src = word.Documents.Open(str(SOURCE_REPORT), False, True, False)
    try:
        body_start = find_paragraph_start(src, r"^第1章", require_heading=True)
        body_rng = src.Range(body_start, src.Content.End - 1)
        body_doc = word.Documents.Add()
        try:
            body_doc.Range().FormattedText = body_rng.FormattedText
            body_doc.SaveAs2(str(TEMP_BODY_DOCX), FileFormat=WD_FORMAT_DOCX)
        finally:
            body_doc.Close(False)
    finally:
        src.Close(False)


def convert_template(word) -> None:
    doc = word.Documents.Open(str(TEMPLATE_DOC), False, True, False)
    try:
        doc.SaveAs2(str(TEMP_TEMPLATE_DOCX), FileFormat=WD_FORMAT_DOCX)
    finally:
        doc.Close(False)


def replace_body_with_report(doc) -> int:
    body_start = find_paragraph_start(doc, r"^第1章", require_heading=True)
    body_rng = doc.Range(body_start, doc.Content.End - 1)
    body_rng.Delete()
    insert_rng = doc.Range(body_start, body_start)
    insert_rng.InsertFile(str(TEMP_BODY_DOCX))
    return body_start


def apply_template_styles(doc, body_start: int) -> None:
    heading1 = safe_style(doc, "Heading 1")
    heading2 = safe_style(doc, "Heading 2")
    heading3 = safe_style(doc, "Heading 3")
    body_style = safe_style(doc, "Body Text First Indent") or safe_style(doc, "Normal")
    normal = safe_style(doc, "Normal")
    plain = safe_style(doc, "Plain Text") or normal

    for paragraph in doc.Paragraphs:
        text = normalize_text(paragraph.Range.Text)
        if not text:
            continue
        try:
            if paragraph.Range.Information(WD_WITHIN_TABLE):
                continue
        except Exception:
            pass

        has_image = paragraph.Range.InlineShapes.Count > 0
        if has_image:
            paragraph.Alignment = WD_ALIGN_CENTER
            continue

        if text in {"摘  要", "摘要"} and heading1 is not None:
            paragraph.Style = heading1
            set_range_font(paragraph.Range, bold=True)
            continue

        if paragraph.Range.Start >= body_start:
            if re.match(r"^第\d+章", text) or text in {"结  论", "结论", "参考文献"}:
                if heading1 is not None:
                    paragraph.Style = heading1
                set_range_font(paragraph.Range, bold=True)
            elif re.match(r"^\d+\.\d+\.\d+", text):
                if heading3 is not None:
                    paragraph.Style = heading3
                set_range_font(paragraph.Range, bold=True)
            elif re.match(r"^\d+\.\d+", text) or re.match(r"^附录[A-Z]\.\d+", text):
                if heading2 is not None:
                    paragraph.Style = heading2
                set_range_font(paragraph.Range, bold=True)
            elif re.match(r"^(图|表)[A-Za-z0-9一二三四五六七八九十A-Z\-]+", text):
                if normal is not None:
                    paragraph.Style = normal
                paragraph.Alignment = WD_ALIGN_CENTER
                set_range_font(paragraph.Range, size=10.5)
            elif text.startswith(("图表说明", "表格说明")):
                if normal is not None:
                    paragraph.Style = normal
                paragraph.Alignment = WD_ALIGN_LEFT
                set_range_font(paragraph.Range, size=10.5)
            else:
                if body_style is not None:
                    paragraph.Style = body_style
                paragraph.Alignment = WD_ALIGN_JUSTIFY
                # 只清除文字颜色和中英文字体，字号由模板正文样式控制。
                paragraph.Range.Font.NameFarEast = "宋体"
                paragraph.Range.Font.NameAscii = "Times New Roman"
                paragraph.Range.Font.NameOther = "Times New Roman"
                paragraph.Range.Font.Color = WD_COLOR_BLACK
        else:
            # 摘要、关键词与目录前文字保持模板区域样式，只清除红色。
            if plain is not None and (text.startswith("关键词") or len(text) > 40):
                paragraph.Style = plain
            paragraph.Range.Font.Color = WD_COLOR_BLACK


def apply_table_format(doc) -> None:
    table_grid = safe_style(doc, "Table Grid")
    for table in doc.Tables:
        try:
            if table_grid is not None:
                table.Style = table_grid
        except Exception:
            pass
        set_range_font(table.Range, size=10.5)
        table.Range.Font.Bold = 0


def fit_images(doc) -> None:
    max_width_pt = 455  # about 16.05 cm
    for shape in doc.InlineShapes:
        try:
            if shape.Width > max_width_pt:
                ratio = max_width_pt / shape.Width
                shape.Width = max_width_pt
                shape.Height = shape.Height * ratio
        except Exception:
            continue


def update_toc_and_fields(doc) -> None:
    for toc in doc.TablesOfContents:
        try:
            toc.Update()
            toc.Range.Font.Color = WD_COLOR_BLACK
            toc.Range.Font.NameFarEast = "宋体"
            toc.Range.Font.NameAscii = "Times New Roman"
        except Exception:
            pass
    try:
        doc.Fields.Update()
    except Exception:
        pass


def run_word_job(job) -> None:
    import win32com.client  # type: ignore

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        job(word)
    finally:
        try:
            word.Quit()
        except Exception:
            pass


def validate_output() -> dict[str, int | bool]:
    doc = Document(OUTPUT_DOCX)
    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    red_runs = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            color = run.font.color.rgb
            if color is not None and str(color).upper().startswith("FF"):
                red_runs += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        color = run.font.color.rgb
                        if color is not None and str(color).upper().startswith("FF"):
                            red_runs += 1
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "red_runs": red_runs,
        "has_old_template_topic": "气体润滑轴承" in all_text or "FLUENT软件" in all_text or "多孔质" in all_text,
        "has_report_title": "高密度城市15分钟生活圈中的可达性幻觉" in all_text,
        "has_bad_question_marks": "????????" in all_text,
        "has_replacement_char": "�" in all_text,
    }


def main() -> None:
    TEMP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if OUTPUT_DOCX.exists():
        shutil.copy2(OUTPUT_DOCX, ROOT / f"报告_final_哈工大模板.before_{stamp}.docx")

    abstract_parts, keywords = collect_source_text()

    run_word_job(convert_template)
    run_word_job(create_body_doc)
    shutil.copy2(TEMP_TEMPLATE_DOCX, OUTPUT_DOCX)

    def fill_doc(word) -> None:
        doc = None
        try:
            doc = word.Documents.Open(str(OUTPUT_DOCX), False, False, False)
            replace_first_matching_paragraph(doc, r"^摘要是论文内容", abstract_parts, "Plain Text")
            replace_first_matching_paragraph(doc, r"^关键词：关键词1", [keywords], "Plain Text")
            body_start = replace_body_with_report(doc)
            apply_template_styles(doc, body_start)
            apply_table_format(doc)
            fit_images(doc)
            update_toc_and_fields(doc)
            doc.SaveAs2(str(OUTPUT_DOCX), FileFormat=WD_FORMAT_DOCX)
        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    pass

    run_word_job(fill_doc)

    print("OUTPUT", OUTPUT_DOCX)
    print("VALIDATION", validate_output())


if __name__ == "__main__":
    main()
