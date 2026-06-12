from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "报告_final_哈工大模板.docx"


FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_LATIN = "Times New Roman"

SIZE_H1_CN = 18.0  # 小二
SIZE_H1_LATIN = 12.0  # 小四
SIZE_H2 = 15.0  # 小三
SIZE_H3 = 14.0  # 四号
SIZE_BODY = 12.0  # 小四
SIZE_CAPTION = 10.5  # 五号


CHAPTER_INTROS = {
    2: "本章作为全文实证分析的基础，重点说明研究区域、数据来源、指标体系和技术路线。通过把社区、人口、POI、路网、街景和建筑形态等数据统一到社区尺度，本章为后续可达性幻觉识别提供可计算、可解释的证据基础。",
    3: "本章进入核心实证分析，围绕南山区内部的可达性幻觉空间分布、街道障碍物机制、步行环境评分和时间贫困效应展开。章节逻辑由宏观空间图谱推进到微观街道阻断，再回到社区尺度的综合诊断。",
    4: "本章通过南山、宝安、福田和龙华四区对照，检验可达性幻觉机制是否只属于单一区域个案，还是能够在不同建成环境和服务供给格局中重复出现。跨区比较有助于区分地方特殊性与可推广规律。",
    5: "本章在前文实证结果基础上讨论研究局限、治理建议和未来研究方向。其目的不是简单罗列政策口号，而是把可达性幻觉识别结果转化为步行网络修复、无障碍设施补短板和生活圈评价方法改进的操作路径。",
}

CHAPTER_SUMMARIES = {
    2: "本章完成了研究区域、数据体系、技术路线和路网建模方法的说明。通过把名义可达性、实际步行阻抗和街景环境变量放在同一分析框架中，研究得以从“设施是否存在”的覆盖率问题推进到“居民是否能够真实到达”的机制问题。",
    3: "本章表明，可达性幻觉在南山区并非边缘现象，而是由服务集聚、路网绕行、街道障碍和夜间可用性共同塑造的空间过程。宏观指标与微观街景证据相互印证，说明高设施密度并不必然等于高真实可达性。",
    4: "本章通过跨区比较说明，可达性幻觉既具有南山区高密度建成环境的地方特征，也具有可迁移的机制基础。不同城区的障碍物类型和街道形态存在差异，但名义可达性与实际通行体验之间的偏差具有共同解释框架。",
    5: "本章总结了数据、模型和外推层面的局限，并提出针对步行网络、无障碍设施、街景监测和生活圈评价体系的治理建议。后续研究需要继续强化微观障碍物识别、弱势群体出行行为和动态服务可用性的联合建模。",
}


def style_name(paragraph) -> str:
    return paragraph.style.name if paragraph.style is not None else ""


def text(paragraph) -> str:
    return paragraph.text.strip()


def is_heading(paragraph, level: int | None = None) -> bool:
    name = style_name(paragraph)
    if level is None:
        return name.startswith("Heading")
    return name == f"Heading {level}"


def body_start_index(doc: Document) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if is_heading(paragraph, 1) and text(paragraph).startswith("第1章"):
            return idx
    raise RuntimeError("未找到正文第1章标题。")


def next_body_h1_index(doc: Document, start_idx: int) -> int:
    for idx in range(start_idx + 1, len(doc.paragraphs)):
        paragraph = doc.paragraphs[idx]
        if is_heading(paragraph, 1) and (
            text(paragraph).startswith("第")
            or text(paragraph).startswith("结")
            or text(paragraph).startswith("参考")
            or text(paragraph).startswith("附录")
        ):
            return idx
    return len(doc.paragraphs)


def chapter_bounds(doc: Document, chapter_no: int) -> tuple[int, int]:
    start = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if is_heading(paragraph, 1) and text(paragraph).startswith(f"第{chapter_no}章"):
            start = idx
            break
    if start is None:
        raise RuntimeError(f"未找到第{chapter_no}章。")
    return start, next_body_h1_index(doc, start)


def set_rfonts(run, east_asia: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = FONT_LATIN
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:cs"), FONT_LATIN)


def is_ascii_segment(value: str) -> bool:
    return all(ord(ch) < 128 for ch in value)


def segments_for_h1(value: str) -> list[tuple[str, bool]]:
    if not value:
        return []
    pieces: list[tuple[str, bool]] = []
    current = value[0]
    current_ascii = is_ascii_segment(value[0])
    for ch in value[1:]:
        ascii_ch = is_ascii_segment(ch)
        if ascii_ch == current_ascii:
            current += ch
        else:
            pieces.append((current, current_ascii))
            current = ch
            current_ascii = ascii_ch
    pieces.append((current, current_ascii))
    return pieces


def set_heading_text(paragraph, value: str, level: int) -> None:
    paragraph.clear()
    paragraph.style = f"Heading {level}"
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for segment, ascii_only in segments_for_h1(value):
            run = paragraph.add_run(segment)
            set_rfonts(run, FONT_HEI, SIZE_H1_LATIN if ascii_only else SIZE_H1_CN, bold=True)
    elif level == 2:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(value)
        set_rfonts(run, FONT_HEI, SIZE_H2, bold=True)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(value)
        set_rfonts(run, FONT_HEI, SIZE_H3, bold=True)


def set_body_text(paragraph, east_asia: str = FONT_SONG, size_pt: float = SIZE_BODY, bold: bool = False) -> None:
    if not paragraph.runs and paragraph.text:
        paragraph.add_run(paragraph.text)
    for run in paragraph.runs:
        set_rfonts(run, east_asia, size_pt, bold=bold)


def insert_before(paragraph, value: str, style: str = "Normal"):
    new_paragraph = paragraph.insert_paragraph_before(value, style=style)
    return new_paragraph


def title_without_number(value: str) -> str:
    return re.sub(r"^\d+\.\d+(?:\.\d+)?\s*", "", value).strip()


def ensure_first_chapter(doc: Document) -> None:
    start, end = chapter_bounds(doc, 1)
    set_heading_text(doc.paragraphs[start], "第1章 绪论", 1)

    # Insert 1.1 before the first content paragraph when the chapter currently starts directly with body text.
    first_after_h1 = doc.paragraphs[start + 1]
    if not (is_heading(first_after_h1, 2) and text(first_after_h1).startswith("1.1") and "课题背景" in text(first_after_h1)):
        inserted = insert_before(first_after_h1, "1.1  课题背景及研究的目的和意义", "Heading 2")
        set_heading_text(inserted, "1.1  课题背景及研究的目的和意义", 2)

    start, end = chapter_bounds(doc, 1)
    h2s = [p for p in doc.paragraphs[start + 1 : end] if is_heading(p, 2)]
    desired = [
        "1.1  课题背景及研究的目的和意义",
        "1.2  15分钟生活圈与可达性幻觉研究现状",
        "1.3  可达性幻觉的概念界定与研究缺口",
    ]
    for paragraph, value in zip(h2s[:3], desired):
        set_heading_text(paragraph, value, 2)

    start, end = chapter_bounds(doc, 1)
    has_main_content = any(is_heading(p, 2) and text(p).startswith("1.4") and "主要研究内容" in text(p) for p in doc.paragraphs[start:end])
    if not has_main_content:
        next_h1 = doc.paragraphs[end]
        additions = [
            ("1.4  本文的主要研究内容", "Heading 2"),
            ("围绕“名义可达性是否等于真实可达性”这一核心问题，本文按照由理论界定到数据建模、由空间诊断到跨区验证、由机制解释到治理建议的顺序展开。全文主要研究内容包括以下五个方面：", "Normal"),
            ("（1）构建可达性幻觉的理论框架。本文将传统15分钟生活圈评价中被忽略的路网绕行、微观障碍、夜间服务可用性和弱势群体通行成本纳入统一概念框架，明确可达性幻觉的内涵、表现形式和测度对象。", "Normal"),
            ("（2）建立多源数据融合与指标计算体系。研究整合社区、人口、POI、步行路网、建筑形态、街景识别和跨区样本数据，形成SAII、TPI、AI*、SCR等指标，用于刻画服务供给、真实路径和街道环境之间的差异。", "Normal"),
            ("（3）识别南山区可达性幻觉的空间格局。通过社区尺度空间分析、四象限诊断和街景障碍物识别，揭示服务密集区内部仍可能出现“设施近但不可达”的结构性失效。", "Normal"),
            ("（4）开展跨区对比与机制验证。以南山、宝安、福田和龙华为对照样本，比较不同建成环境中的障碍物类型、步行环境质量和时间贫困差异，检验研究框架的可迁移性。", "Normal"),
            ("（5）提出面向治理的评价改进方向。研究最终服务于生活圈评价方法、步行网络修复、无障碍设施补短板和街景智能监测，为高密度城市公共服务公平治理提供可操作依据。", "Normal"),
        ]
        for value, style in additions:
            paragraph = insert_before(next_h1, value, style)
            if style == "Heading 2":
                set_heading_text(paragraph, value, 2)


def ensure_chapter_intro_summary(doc: Document, chapter_no: int) -> None:
    start, end = chapter_bounds(doc, chapter_no)
    first_after_h1 = doc.paragraphs[start + 1]
    original_first_was_h2 = is_heading(first_after_h1, 2)

    if not (is_heading(first_after_h1, 2) and text(first_after_h1).startswith(f"{chapter_no}.1") and "引言" in text(first_after_h1)):
        intro_h2 = insert_before(first_after_h1, f"{chapter_no}.1  引言", "Heading 2")
        set_heading_text(intro_h2, f"{chapter_no}.1  引言", 2)
        if original_first_was_h2:
            intro_body = insert_before(first_after_h1, CHAPTER_INTROS[chapter_no], "Normal")
            set_body_text(intro_body, FONT_SONG, SIZE_BODY)

    start, end = chapter_bounds(doc, chapter_no)
    h2s = [p for p in doc.paragraphs[start + 1 : end] if is_heading(p, 2)]
    content_h2s = [p for p in h2s if "引言" not in text(p) and "本章小结" not in text(p)]
    counter = 2
    for paragraph in content_h2s:
        title = title_without_number(text(paragraph))
        if title:
            set_heading_text(paragraph, f"{chapter_no}.{counter}  {title}", 2)
            counter += 1

    start, end = chapter_bounds(doc, chapter_no)
    has_summary = any(is_heading(p, 2) and "本章小结" in text(p) for p in doc.paragraphs[start:end])
    if not has_summary:
        next_h1 = doc.paragraphs[end]
        heading_text = f"{chapter_no}.{counter}  本章小结"
        summary_h2 = insert_before(next_h1, heading_text, "Heading 2")
        set_heading_text(summary_h2, heading_text, 2)
        summary_p = insert_before(next_h1, CHAPTER_SUMMARIES[chapter_no], "Normal")
        set_body_text(summary_p, FONT_SONG, SIZE_BODY)


def normalize_all_headings(doc: Document) -> None:
    start = body_start_index(doc)
    paragraphs = doc.paragraphs
    for idx, paragraph in enumerate(paragraphs):
        value = text(paragraph)
        if paragraph.paragraph_format is not None:
            paragraph.paragraph_format.widow_control = True
        if not value:
            if paragraph._p.xml.find("w:drawing") >= 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if style_name(paragraph).startswith("toc"):
            continue
        if is_heading(paragraph, 1):
            set_heading_text(paragraph, value, 1)
        elif is_heading(paragraph, 2):
            set_heading_text(paragraph, value, 2)
        elif is_heading(paragraph, 3):
            set_heading_text(paragraph, value, 3)
        elif paragraph._p.xml.find("w:drawing") >= 0 and not value:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif re.match(r"^[（(]\d+[）)]", value):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_body_text(paragraph, FONT_HEI, SIZE_BODY, bold=True)
        elif re.match(r"^(图|表)[A-Za-z0-9一二三四五六七八九十A-Z\-]*", value) or value.startswith(("图表说明", "表格说明")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if value.startswith(("图", "表")) and not value.startswith(("图表说明", "表格说明")) else WD_ALIGN_PARAGRAPH.LEFT
            set_body_text(paragraph, FONT_SONG, SIZE_CAPTION)
        elif paragraph._p.xml.find("w:drawing") >= 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_body_text(paragraph, FONT_SONG, SIZE_CAPTION)
        elif idx >= start or value.startswith(("关键词", "“15分钟城市”", "本研究以", "本文的贡献")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_body_text(paragraph, FONT_SONG, SIZE_BODY)
        else:
            # Cover and directory front matter keep template layout, but remove red/color drift and set stable CJK/Latin fonts.
            set_body_text(paragraph, FONT_SONG, SIZE_BODY)


def normalize_tables(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_body_text(paragraph, FONT_SONG, SIZE_CAPTION)


def update_toc_with_word(path: Path) -> None:
    import win32com.client  # type: ignore

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    doc = None
    try:
        doc = word.Documents.Open(str(path), False, False, False)
        try:
            for toc in doc.TablesOfContents:
                toc.Update()
                toc.Range.Font.NameFarEast = FONT_SONG
                toc.Range.Font.NameAscii = FONT_LATIN
                toc.Range.Font.NameOther = FONT_LATIN
                toc.Range.Font.Color = 0
        except Exception:
            pass
        try:
            doc.Fields.Update()
        except Exception:
            pass
        doc.SaveAs2(str(path), FileFormat=16)
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        try:
            word.Quit()
        except Exception:
            pass


def validate(path: Path) -> dict[str, int | bool | list[str]]:
    doc = Document(path)
    headings = [text(p) for p in doc.paragraphs if is_heading(p, 1) or is_heading(p, 2)]
    all_text = "\n".join(p.text for p in doc.paragraphs)
    checks = {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "a4_captions": sum(1 for p in doc.paragraphs if text(p).startswith("图A4-") and "街道断面分析" in text(p)),
        "a5_drawings": sum(t._tbl.xml.count("<w:drawing") for t in doc.tables if len(t.rows) == 15 and len(t.columns) == 4),
        "bad_question_marks": "????????" in all_text,
        "replacement_char": "�" in all_text,
        "old_template_terms": any(term in all_text for term in ["气体润滑轴承", "FLUENT软件", "多孔质", "试件渗透率"]),
        "key_headings": [h for h in headings if h.startswith(("第1章", "1.1", "1.4", "2.1", "2.6", "3.1", "3.7", "4.1", "4.4", "5.1", "5.5"))][:20],
    }
    return checks


def main() -> None:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = ROOT / f"报告_final_哈工大模板.before_structure_format_{stamp}.docx"
    shutil.copy2(DOCX, backup)

    doc = Document(DOCX)
    ensure_first_chapter(doc)
    for chapter_no in [2, 3, 4, 5]:
        ensure_chapter_intro_summary(doc, chapter_no)
    normalize_all_headings(doc)
    normalize_tables(doc)
    doc.save(DOCX)
    update_toc_with_word(DOCX)

    print("BACKUP", backup)
    print("OUTPUT", DOCX)
    print("VALIDATION", validate(DOCX))


if __name__ == "__main__":
    main()
