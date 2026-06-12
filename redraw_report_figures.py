from __future__ import annotations

import math
import re
import shutil
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from matplotlib import font_manager
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Rectangle
from PIL import Image


ROOT = Path(__file__).resolve().parent
REPORT = ROOT / "报告_final.docx"
PAPER_FIG_DIR = ROOT / "projects" / "15min-urban-accessibility" / "paper" / "figures"
CONFERENCE_FIG_DIR = ROOT / "projects" / "15min-urban-accessibility" / "conference_paper" / "figures"
LEGACY_CONFERENCE_FIG_DIR = ROOT / "15分钟城市时间贫困研究" / "conference_paper" / "figures"
OVERLEAF_FIG_DIR = (
    ROOT
    / "papers"
    / "conference-slides"
    / "会议论文"
    / "15min可达性幻觉"
    / "overleaf_paper"
    / "figures"
)
ACCESSIBILITY_CSV = ROOT / "projects" / "15min-urban-accessibility" / "accessibility_results.csv"
PROFILE_DIR = ROOT / "projects" / "15min-urban-accessibility" / "v2_real_data" / "street_profiles"
PROFILE_SUMMARY_CSV = PROFILE_DIR / "profile_summary.csv"


FONT_CANDIDATES = [
    Path("C:/Windows/Fonts/msyh.ttc"),
    Path("C:/Windows/Fonts/simhei.ttf"),
    Path("C:/Windows/Fonts/simsun.ttc"),
    Path("C:/Windows/Fonts/Deng.ttf"),
]


USAGE_COLORS = {
    "Residential": "#4C78A8",
    "Commercial": "#E45756",
    "Industrial": "#F58518",
    "Infrastructure": "#54A24B",
    "Mixed Residential": "#B279A2",
    "Mixed_Area": "#B279A2",
    "Other": "#8D6E63",
    "Unknown": "#9E9E9E",
}

USAGE_CN = {
    "Residential": "住宅建筑",
    "Commercial": "商业/办公",
    "Industrial": "工业/产业",
    "Infrastructure": "基础设施",
    "Mixed Residential": "混合居住",
    "Mixed_Area": "混合区域",
    "Other": "其他建筑",
    "Unknown": "未知用途",
}


def configure_chinese_font() -> tuple[str, font_manager.FontProperties]:
    for font_path in FONT_CANDIDATES:
        if font_path.exists():
            font_manager.fontManager.addfont(str(font_path))
            prop = font_manager.FontProperties(fname=str(font_path))
            family = prop.get_name()
            plt.rcParams["font.family"] = family
            plt.rcParams["font.sans-serif"] = [family, "Microsoft YaHei", "SimHei", "SimSun"]
            plt.rcParams["axes.unicode_minus"] = False
            return family, prop
    prop = font_manager.FontProperties()
    plt.rcParams["axes.unicode_minus"] = False
    return "default", prop


FONT_FAMILY, FONT_PROP = configure_chinese_font()


def add_text(ax, x, y, text, **kwargs):
    kwargs.setdefault("fontproperties", FONT_PROP)
    return ax.text(x, y, text, **kwargs)


def set_title(ax, title, **kwargs):
    kwargs.setdefault("fontproperties", FONT_PROP)
    return ax.set_title(title, **kwargs)


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def copy_with_relative_path(src: Path, backup_root: Path) -> None:
    if not src.exists():
        return
    rel = src.relative_to(ROOT)
    dst = backup_root / rel
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)


def backup_inputs(target_paths: list[Path], profile_paths: list[Path]) -> Path:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_root = ROOT / "figure_redraw_backups" / stamp
    backup_root.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REPORT, ROOT / f"报告_final.before_figure_redraw_{stamp}.docx")
    for path in target_paths:
        copy_with_relative_path(path, backup_root)
    for path in profile_paths:
        copy_with_relative_path(path, backup_root)
    copy_with_relative_path(PROFILE_SUMMARY_CSV, backup_root)
    return backup_root


def rounded_box(ax, x, y, w, h, label, fc, ec="#283445", lw=1.5, fs=17):
    box = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.018,rounding_size=0.035",
        facecolor=fc,
        edgecolor=ec,
        linewidth=lw,
    )
    ax.add_patch(box)
    add_text(
        ax,
        x + w / 2,
        y + h / 2,
        label,
        ha="center",
        va="center",
        fontsize=fs,
        color="#13202F",
        linespacing=1.35,
        weight="bold",
    )
    return box


def arrow(ax, x1, y1, x2, y2, color="#2F3E46", text=None):
    arr = FancyArrowPatch(
        (x1, y1),
        (x2, y2),
        arrowstyle="-|>",
        mutation_scale=18,
        linewidth=1.8,
        color=color,
        connectionstyle="arc3,rad=0.0",
    )
    ax.add_patch(arr)
    if text:
        add_text(
            ax,
            (x1 + x2) / 2,
            (y1 + y2) / 2 + 0.035,
            text,
            ha="center",
            va="center",
            fontsize=11,
            color=color,
            bbox=dict(boxstyle="round,pad=0.18", facecolor="#FFFFFF", edgecolor="none", alpha=0.9),
        )


def draw_framework(path: Path, variant_label: str = "国家自然科学基金项目研究路线") -> None:
    ensure_parent(path)
    fig, ax = plt.subplots(figsize=(15.8, 9.2), dpi=220)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("#F7F8FA")
    ax.set_facecolor("#F7F8FA")

    add_text(
        ax,
        0.5,
        0.955,
        "15分钟城市可达性幻觉识别：研究框架与技术路线图",
        ha="center",
        va="center",
        fontsize=26,
        weight="bold",
        color="#122033",
    )
    add_text(
        ax,
        0.5,
        0.912,
        variant_label,
        ha="center",
        va="center",
        fontsize=14,
        color="#52616B",
    )

    stages = [
        (
            "01 数据基础",
            ["社区与人口", "建筑轮廓与形态", "POI生活服务", "步行路网", "街景与障碍样本"],
            "#D9EAF7",
        ),
        (
            "02 名义可达性",
            ["15分钟服务承诺", "M2SFCA供需匹配", "欧氏距离与覆盖率", "服务容量归一化"],
            "#E6F4EA",
        ),
        (
            "03 实际步行可达性",
            ["路网最短路径", "绕行比 Rnet", "夜间服务衰减", "台阶/坡道/路障阻抗"],
            "#FFF3D9",
        ),
        (
            "04 幻觉诊断与治理",
            ["SAII 综合幻觉指数", "TPI 时间贫困指数", "AI* 实际可达性", "SCR 街景通行阻抗", "治理分区与优先级"],
            "#FCE3E1",
        ),
    ]
    xs = [0.055, 0.295, 0.535, 0.775]
    w = 0.18
    y_top = 0.785
    header_h = 0.07
    item_h = 0.052
    gap = 0.015

    for idx, (title, items, color) in enumerate(stages):
        x = xs[idx]
        rounded_box(ax, x, y_top, w, header_h, title, color, fs=16)
        y = y_top - gap - item_h
        for item in items:
            rounded_box(ax, x + 0.013, y, w - 0.026, item_h, item, "#FFFFFF", ec="#CBD5DE", lw=1.0, fs=12)
            y -= item_h + gap * 0.72
        if idx < len(stages) - 1:
            arrow(ax, x + w + 0.012, y_top + header_h / 2, xs[idx + 1] - 0.012, y_top + header_h / 2, text=["清洗融合", "路径约束", "指标合成"][idx])

    y_mid = 0.28
    rounded_box(
        ax,
        0.06,
        y_mid,
        0.22,
        0.11,
        "理论问题\n地图承诺是否等于真实可达？",
        "#E8EEF6",
        ec="#AAB7C4",
        fs=13,
    )
    rounded_box(
        ax,
        0.365,
        y_mid,
        0.25,
        0.11,
        "机制识别\n空间绕行 + 时间贫困 + 环境阻抗",
        "#EAF3EA",
        ec="#A9BEA9",
        fs=13,
    )
    rounded_box(
        ax,
        0.70,
        y_mid,
        0.23,
        0.11,
        "治理输出\n风险热区、断点修复、无障碍优先清单",
        "#F7E7E3",
        ec="#CFA99F",
        fs=13,
    )
    arrow(ax, 0.285, y_mid + 0.055, 0.358, y_mid + 0.055, color="#546A7B")
    arrow(ax, 0.625, y_mid + 0.055, 0.692, y_mid + 0.055, color="#546A7B")

    note = (
        "读图说明：流程按“需求与供给数据 -> 名义可达性 -> 真实路径与微观阻抗 -> 幻觉诊断与治理分区”逐级推进；"
        "箭头表示因果递进关系，指标用于把规划承诺与居民实际步行体验之间的偏差量化。"
    )
    add_text(
        ax,
        0.5,
        0.095,
        note,
        ha="center",
        va="center",
        fontsize=12.5,
        color="#2E4057",
        wrap=True,
        bbox=dict(boxstyle="round,pad=0.55", facecolor="#FFFFFF", edgecolor="#D8DEE6", linewidth=1.0),
    )

    plt.savefig(path, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


def cleaned_profile_points(profile_df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for _, row in profile_df.iterrows():
        lat = float(row.get("avg_lat", np.nan))
        lng = float(row.get("avg_lng", np.nan))
        if abs(lat) > 90 and 20 <= lng <= 25:
            lat, lng = lng, lat
        if 20 <= lat <= 25 and 112 <= lng <= 115:
            clean = row.copy()
            clean["clean_lat"] = lat
            clean["clean_lng"] = lng
            rows.append(clean)
    if not rows:
        return pd.DataFrame(columns=list(profile_df.columns) + ["clean_lat", "clean_lng"])
    return pd.DataFrame(rows)


def nearest_scr_for_communities(communities: pd.DataFrame, profiles: pd.DataFrame) -> np.ndarray:
    if profiles.empty:
        return np.full(len(communities), 0.5)
    c_lat = communities["lat"].to_numpy(float)
    c_lng = communities["lng"].to_numpy(float)
    p_lat = profiles["clean_lat"].to_numpy(float)
    p_lng = profiles["clean_lng"].to_numpy(float)
    scr = profiles["scr_estimate"].to_numpy(float)
    values = []
    for lat, lng in zip(c_lat, c_lng):
        scale = max(math.cos(math.radians(lat)), 0.2)
        dist2 = ((p_lng - lng) * 111.0 * scale) ** 2 + ((p_lat - lat) * 111.0) ** 2
        values.append(float(scr[int(np.argmin(dist2))]))
    return np.asarray(values)


def draw_spatial_comparison(path: Path, sci_style: bool = False) -> None:
    ensure_parent(path)
    df = pd.read_csv(ACCESSIBILITY_CSV)
    profiles = cleaned_profile_points(pd.read_csv(PROFILE_SUMMARY_CSV))
    df = df.copy()
    df["AI_star_proxy"] = (df["A_day_norm"].astype(float) + df["A_night_norm"].astype(float)) / 2.0
    df["TPI_intensity"] = np.where(df["TPI"].astype(float) < 0, -df["TPI"].astype(float), df["TPI"].astype(float))
    df["SCR_proxy"] = nearest_scr_for_communities(df, profiles)
    pop = df["population"].fillna(df["population"].median()).astype(float)
    pop_size = 28 + 95 * (pop - pop.min()) / max(pop.max() - pop.min(), 1.0)

    panels = [
        ("SAII 综合幻觉指数", "SAII", "高值表示名义承诺与真实可达性偏差更强", "YlOrRd"),
        ("TPI 时间贫困强度", "TPI_intensity", "以 -TPI/绝对偏差表达，数值越大表示时间损失越重", "Oranges"),
        ("AI* 实际综合可达性", "AI_star_proxy", "昼夜可达性归一化均值，数值越高表示服务更容易到达", "Greens"),
        ("SCR 街景通行阻抗", "SCR_proxy", "邻近街道断面估计，数值越高表示人行空间受压越明显", "PuRd"),
    ]

    fig, axes = plt.subplots(2, 2, figsize=(15.2, 11.2), dpi=220)
    fig.patch.set_facecolor("#FFFFFF")
    axes = axes.ravel()
    for ax, (title, col, desc, cmap) in zip(axes, panels):
        values = df[col].astype(float).to_numpy()
        lo = np.nanpercentile(values, 2)
        hi = np.nanpercentile(values, 98)
        if not np.isfinite(lo) or not np.isfinite(hi) or lo == hi:
            lo, hi = np.nanmin(values), np.nanmax(values) + 1e-6
        sc = ax.scatter(
            df["lng"],
            df["lat"],
            c=values,
            cmap=cmap,
            s=pop_size,
            vmin=lo,
            vmax=hi,
            edgecolors="#FFFFFF",
            linewidths=0.55,
            alpha=0.88,
        )
        ax.grid(True, color="#E2E7EE", linewidth=0.8)
        ax.set_facecolor("#FAFBFC")
        set_title(ax, title, fontsize=16, weight="bold", pad=10)
        ax.set_xlabel("经度", fontsize=11, fontproperties=FONT_PROP)
        ax.set_ylabel("纬度", fontsize=11, fontproperties=FONT_PROP)
        ax.tick_params(labelsize=9)
        for label in ax.get_xticklabels() + ax.get_yticklabels():
            label.set_fontproperties(FONT_PROP)
        cb = fig.colorbar(sc, ax=ax, shrink=0.82, pad=0.012)
        cb.ax.tick_params(labelsize=9)
        for label in cb.ax.get_yticklabels():
            label.set_fontproperties(FONT_PROP)
        cb.set_label("指标值", fontsize=10, fontproperties=FONT_PROP)
        add_text(
            ax,
            0.02,
            0.02,
            desc,
            transform=ax.transAxes,
            ha="left",
            va="bottom",
            fontsize=10,
            color="#334155",
            bbox=dict(boxstyle="round,pad=0.25", facecolor="#FFFFFF", edgecolor="#D7DEE8", alpha=0.92),
        )

    heading = "多维可达性指标空间分布对比（SAII · TPI · AI* · SCR）"
    if sci_style:
        heading = "四指标并列空间分布详图（SCI版）"
    fig.suptitle(heading, fontsize=22, fontproperties=FONT_PROP, weight="bold", y=0.988)
    fig.text(
        0.5,
        0.028,
        "说明：每个点代表一个社区，点大小近似表示人口规模；颜色越深表示该子图对应指标强度越高。SCR为邻近街道断面估计值，用于补足社区表中缺失的街景通行阻抗字段。",
        ha="center",
        va="center",
        fontsize=11.5,
        color="#334155",
        fontproperties=FONT_PROP,
    )
    fig.tight_layout(rect=(0.02, 0.055, 0.98, 0.955))
    plt.savefig(path, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


def usage_key(raw: str) -> str:
    if pd.isna(raw):
        return "Unknown"
    raw = str(raw).strip()
    return raw if raw in USAGE_COLORS else "Other"


def risk_level(scr: float, eww: float) -> tuple[str, str]:
    if scr >= 0.62 or eww <= 2.5:
        return "较高通行风险", "#C0392B"
    if scr >= 0.45 or eww <= 3.5:
        return "中等通行风险", "#D9822B"
    return "较低通行风险", "#2E7D32"


def distribute_floors(row: pd.Series) -> np.ndarray:
    n = int(max(row.get("n_buildings", 1), 1))
    min_f = int(max(row.get("min_floors", 1), 1))
    max_f = int(max(row.get("max_floors", min_f), min_f))
    total = int(max(row.get("total_floors", n * min_f), n * min_f))
    seed = int(re.sub(r"\D", "", str(row.get("profile_id", "0"))) or "0")
    rng = np.random.default_rng(seed)
    if n == 1:
        return np.array([max_f])
    values = rng.triangular(min_f, max(row.get("avg_floors", min_f), min_f), max_f + 0.1, n)
    floors = np.clip(np.rint(values), min_f, max_f).astype(int)
    floors[0] = max_f
    floors[-1] = min_f
    for _ in range(2000):
        diff = total - int(floors.sum())
        if diff == 0:
            break
        if diff > 0:
            idxs = np.where(floors < max_f)[0]
            if len(idxs) == 0:
                break
            floors[int(rng.choice(idxs))] += 1
        else:
            idxs = np.where(floors > min_f)[0]
            if len(idxs) == 0:
                break
            floors[int(rng.choice(idxs))] -= 1
    return floors


def profile_image_path(row: pd.Series, existing: dict[str, Path]) -> Path:
    profile_id = str(row["profile_id"])
    if profile_id in existing:
        return existing[profile_id]
    safe_road = str(row.get("road_name", "未命名道路")).replace("/", "_").replace("\\", "_")
    return PROFILE_DIR / f"{profile_id}_{safe_road}.png"


def draw_profile(row: pd.Series, path: Path) -> None:
    ensure_parent(path)
    usage = usage_key(row.get("dominant_usage", "Unknown"))
    usage_color = USAGE_COLORS[usage]
    usage_label = USAGE_CN.get(usage, "其他建筑")
    floors = distribute_floors(row)
    n = len(floors)
    scr = float(row.get("scr_estimate", 0.0))
    eww = float(row.get("eww_estimate", 0.0))
    level, level_color = risk_level(scr, eww)
    road_name = str(row.get("road_name", "未命名道路"))
    profile_id = str(row.get("profile_id", "PROFILE"))

    fig = plt.figure(figsize=(14.8, 8.4), dpi=210, facecolor="#FFFFFF")
    gs = fig.add_gridspec(
        3,
        4,
        height_ratios=[0.22, 1.0, 0.55],
        width_ratios=[1.1, 1.1, 1.1, 0.92],
        hspace=0.18,
        wspace=0.20,
    )
    ax_title = fig.add_subplot(gs[0, :])
    ax_section = fig.add_subplot(gs[1, :3])
    ax_stats = fig.add_subplot(gs[1, 3])
    ax_bar = fig.add_subplot(gs[2, :2])
    ax_legend = fig.add_subplot(gs[2, 2:])

    for ax in [ax_title, ax_stats, ax_legend]:
        ax.axis("off")

    add_text(
        ax_title,
        0.01,
        0.60,
        f"街道断面分析：{profile_id}  {road_name}",
        transform=ax_title.transAxes,
        ha="left",
        va="center",
        fontsize=22,
        weight="bold",
        color="#102A43",
    )
    add_text(
        ax_title,
        0.99,
        0.60,
        level,
        transform=ax_title.transAxes,
        ha="right",
        va="center",
        fontsize=15,
        weight="bold",
        color="#FFFFFF",
        bbox=dict(boxstyle="round,pad=0.35", facecolor=level_color, edgecolor=level_color),
    )

    max_floor = max(int(np.max(floors)), 1)
    ax_section.set_xlim(0, 100)
    ax_section.set_ylim(0, max_floor + 8)
    ax_section.set_facecolor("#F8FAFC")
    ax_section.grid(axis="y", color="#E2E8F0", linewidth=0.7)
    ax_section.set_ylabel("建筑楼层数", fontsize=12, fontproperties=FONT_PROP)
    ax_section.set_xticks([])
    ax_section.tick_params(labelsize=9)
    for label in ax_section.get_yticklabels():
        label.set_fontproperties(FONT_PROP)
    for spine in ax_section.spines.values():
        spine.set_visible(False)

    ax_section.add_patch(Rectangle((0, 0), 100, 1.2, facecolor="#DDE6ED", edgecolor="none", zorder=0))
    ax_section.add_patch(Rectangle((32, 0), 36, 2.0, facecolor="#686F78", edgecolor="none", zorder=1))
    ax_section.add_patch(Rectangle((18, 0), 14, 1.65, facecolor="#B9D7EA", edgecolor="none", zorder=2))
    ax_section.add_patch(Rectangle((68, 0), 14, 1.65, facecolor="#B9D7EA", edgecolor="none", zorder=2))
    ax_section.add_patch(Rectangle((12, 0), 6, 1.35, facecolor="#B8E0B8", edgecolor="none", zorder=2))
    ax_section.add_patch(Rectangle((82, 0), 6, 1.35, facecolor="#B8E0B8", edgecolor="none", zorder=2))
    add_text(ax_section, 50, 1.05, "机动车道", ha="center", va="center", fontsize=12, color="#FFFFFF", weight="bold")
    add_text(ax_section, 25, 1.92, "人行空间", ha="center", va="bottom", fontsize=11, color="#1E3A5F", weight="bold")
    add_text(ax_section, 75, 1.92, "人行空间", ha="center", va="bottom", fontsize=11, color="#1E3A5F", weight="bold")

    left_count = n // 2
    right_count = n - left_count
    left_xs = np.linspace(2, 16, max(left_count, 1))
    right_xs = np.linspace(84, 98, max(right_count, 1))
    bar_width = min(2.1, 14 / max(max(left_count, right_count), 1))
    building_positions = list(left_xs[:left_count]) + list(right_xs[:right_count])
    for x, f in zip(building_positions, floors):
        ax_section.add_patch(
            Rectangle(
                (x - bar_width / 2, 1.2),
                bar_width,
                max(float(f), 0.8),
                facecolor=usage_color,
                edgecolor="#FFFFFF",
                linewidth=0.8,
                alpha=0.92,
            )
        )
    add_text(ax_section, 8.5, max_floor + 4.1, "左侧建筑界面", ha="center", va="center", fontsize=11, color="#334155")
    add_text(ax_section, 91.5, max_floor + 4.1, "右侧建筑界面", ha="center", va="center", fontsize=11, color="#334155")
    add_text(
        ax_section,
        50,
        max_floor + 4.5,
        f"建筑颜色表示主导用途：{usage_label}；柱高表示楼层数",
        ha="center",
        va="center",
        fontsize=12.5,
        color="#102A43",
        bbox=dict(boxstyle="round,pad=0.35", facecolor="#FFFFFF", edgecolor="#CBD5E1"),
    )

    metrics = [
        ("建筑数", f"{int(row.get('n_buildings', 0))} 栋"),
        ("总楼层", f"{int(row.get('total_floors', 0))} 层"),
        ("平均楼层", f"{float(row.get('avg_floors', 0.0)):.1f} 层"),
        ("最高楼层", f"{int(row.get('max_floors', 0))} 层"),
        ("主导用途", usage_label),
        ("SCR估计", f"{scr:.2f}"),
        ("EWW估计", f"{eww:.1f} m"),
    ]
    y = 0.94
    for label, value in metrics:
        ax_stats.add_patch(
            FancyBboxPatch(
                (0.04, y - 0.095),
                0.92,
                0.075,
                boxstyle="round,pad=0.018,rounding_size=0.03",
                facecolor="#F8FAFC",
                edgecolor="#D9E2EC",
            )
        )
        add_text(ax_stats, 0.10, y - 0.058, label, transform=ax_stats.transAxes, ha="left", va="center", fontsize=10.5, color="#52616B")
        add_text(ax_stats, 0.91, y - 0.058, value, transform=ax_stats.transAxes, ha="right", va="center", fontsize=11.5, color="#102A43", weight="bold")
        y -= 0.116

    bar_labels = ["最低楼层", "平均楼层", "最高楼层"]
    bar_values = [float(row.get("min_floors", 0)), float(row.get("avg_floors", 0)), float(row.get("max_floors", 0))]
    bars = ax_bar.bar(bar_labels, bar_values, color=["#9FB3C8", "#4C78A8", "#1F4E79"], width=0.55)
    ax_bar.set_title("楼层强度对比", fontsize=14, fontproperties=FONT_PROP, weight="bold")
    ax_bar.set_ylabel("楼层数", fontsize=11, fontproperties=FONT_PROP)
    ax_bar.grid(axis="y", color="#E2E8F0", linewidth=0.8)
    ax_bar.spines[["top", "right"]].set_visible(False)
    for tick in ax_bar.get_xticklabels() + ax_bar.get_yticklabels():
        tick.set_fontproperties(FONT_PROP)
        tick.set_fontsize(10)
    for bar, value in zip(bars, bar_values):
        ax_bar.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + max(bar_values) * 0.03 + 0.2,
            f"{value:.1f}",
            ha="center",
            va="bottom",
            fontsize=10.5,
            fontproperties=FONT_PROP,
            color="#102A43",
        )

    ax_legend.set_xlim(0, 1)
    ax_legend.set_ylim(0, 1)
    add_text(ax_legend, 0.02, 0.88, "建筑用途颜色图例", ha="left", va="center", fontsize=14, weight="bold", color="#102A43")
    legend_items = [
        ("Residential", "住宅建筑"),
        ("Commercial", "商业/办公"),
        ("Industrial", "工业/产业"),
        ("Infrastructure", "基础设施"),
        ("Mixed Residential", "混合居住"),
        ("Other", "其他/未知"),
    ]
    for i, (key, label) in enumerate(legend_items):
        x = 0.03 + (i % 3) * 0.31
        y = 0.62 - (i // 3) * 0.28
        ax_legend.add_patch(Rectangle((x, y), 0.045, 0.11, facecolor=USAGE_COLORS[key], edgecolor="#FFFFFF"))
        add_text(ax_legend, x + 0.058, y + 0.055, label, ha="left", va="center", fontsize=11.5, color="#334155")
    add_text(
        ax_legend,
        0.02,
        0.08,
        "SCR表示人行空间受机动车道或障碍物挤压后的通行阻抗估计；EWW表示有效步行宽度估计。",
        ha="left",
        va="bottom",
        fontsize=10.5,
        color="#52616B",
        wrap=True,
    )

    plt.savefig(path, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)


def draw_all_profiles() -> list[Path]:
    summary = pd.read_csv(PROFILE_SUMMARY_CSV)
    existing = {}
    for path in PROFILE_DIR.glob("PROFILE_*.png"):
        match = re.match(r"(PROFILE_\d+)", path.stem)
        if match:
            existing[match.group(1)] = path
    output_paths = []
    for _, row in summary.iterrows():
        path = profile_image_path(row, existing)
        draw_profile(row, path)
        output_paths.append(path)
    return output_paths


def set_word_run_font(run, size_pt: float | None = None, bold: bool | None = None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def set_paragraph_text(paragraph, text: str, size_pt: float = 10.5) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_word_run_font(run, size_pt=size_pt)


def find_caption(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    return None


def next_paragraph(doc: Document, paragraph):
    paragraphs = doc.paragraphs
    for idx, item in enumerate(paragraphs):
        if item._p is paragraph._p and idx + 1 < len(paragraphs):
            return paragraphs[idx + 1]
    return None


def remove_drawing_paragraphs_before(paragraph, max_remove: int = 3) -> int:
    removed = 0
    parent = paragraph._p.getparent()
    prev = paragraph._p.getprevious()
    while prev is not None and removed < max_remove:
        xml = prev.xml
        if prev.tag.endswith("}p") and "w:drawing" in xml:
            to_remove = prev
            prev = prev.getprevious()
            parent.remove(to_remove)
            removed += 1
            continue
        break
    return removed


def insert_images_before(paragraph, image_paths: list[Path], width_cm: float) -> None:
    for image_path in image_paths:
        new_p = paragraph.insert_paragraph_before()
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = new_p.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))


def replace_single_figure(doc: Document, prefix: str, image_path: Path, width_cm: float, note: str) -> bool:
    caption = find_caption(doc, prefix)
    if caption is None:
        return False
    remove_drawing_paragraphs_before(caption, max_remove=3)
    insert_images_before(caption, [image_path], width_cm=width_cm)
    note_p = next_paragraph(doc, caption)
    if note_p is not None and note_p.text.strip().startswith("图表说明"):
        set_paragraph_text(note_p, note)
    return True


def profile_paths_from_caption(caption_text: str) -> list[Path]:
    match = re.search(r"（(.+?)）", caption_text)
    if not match:
        return []
    names = [name.strip() for name in match.group(1).split("&")]
    paths = []
    for name in names:
        path = PROFILE_DIR / name
        if path.exists():
            paths.append(path)
        else:
            prefix = re.match(r"(PROFILE_\d+)", name)
            if prefix:
                found = list(PROFILE_DIR.glob(prefix.group(1) + "_*.png"))
                if found:
                    paths.append(found[0])
    return paths


def replace_profile_figures(doc: Document) -> int:
    captions = [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip().startswith("图A4-") and "街道断面分析" in p.text
    ]
    replaced = 0
    note = (
        "图表说明：该组街道断面图以大字号中文重绘，柱高表示建筑楼层数，建筑颜色表示主导用途：蓝色=住宅建筑，红色=商业/办公，"
        "橙色=工业/产业，绿色=基础设施，紫色=混合居住，棕/灰色=其他或未知。SCR表示人行空间通行阻抗估计，EWW表示有效步行宽度估计；"
        "两项指标共同用于判断人行通道是否连续、是否被建筑界面、机动车道或临时障碍压缩。"
    )
    for caption_text in captions:
        caption = None
        for p in doc.paragraphs:
            if p.text.strip() == caption_text:
                caption = p
                break
        if caption is None:
            continue
        image_paths = profile_paths_from_caption(caption_text)
        if not image_paths:
            continue
        remove_drawing_paragraphs_before(caption, max_remove=3)
        insert_images_before(caption, image_paths, width_cm=16.1)
        note_p = next_paragraph(doc, caption)
        if note_p is not None and note_p.text.strip().startswith("图表说明"):
            set_paragraph_text(note_p, note)
        replaced += 1
    return replaced


def replace_report_images() -> dict[str, int]:
    doc = Document(REPORT)
    framework = PAPER_FIG_DIR / "fig1_framework.png"
    conference_framework = CONFERENCE_FIG_DIR / "fig1_framework.png"
    spatial = PAPER_FIG_DIR / "fig_four_index_spatial_comparison.png"
    spatial_sci = PAPER_FIG_DIR / "fig11_four_index_comparison_sci.png"

    framework_note = (
        "图表说明：该图以“数据基础—名义可达性—实际步行可达性—幻觉诊断与治理”的层进结构概括研究路线。"
        "左侧数据层提供社区、人口、POI、路网和街景证据；中间两层分别计算规划承诺和真实路径阻抗；右侧输出SAII、TPI、AI*、SCR等指标及治理分区，箭头表示前后因果递进。"
    )
    spatial_note = (
        "图表说明：该图并列展示SAII、TPI、AI*和SCR四类指标的社区空间分布。每个点代表一个社区，点大小近似反映人口规模，颜色越深表示对应指标强度越高；"
        "SAII和TPI用于识别可达性幻觉与时间贫困热点，AI*表示昼夜综合可达性水平，SCR表示邻近街道断面推断的人行通行阻抗。"
    )

    results = {
        "framework": 0,
        "spatial": 0,
        "profiles": 0,
    }
    for prefix, image in [
        ("图5", framework),
        ("图A1-1", conference_framework),
        ("图A2-1", framework),
    ]:
        if replace_single_figure(doc, prefix, image, 16.1, framework_note):
            results["framework"] += 1
    for prefix, image in [
        ("图19", spatial),
        ("图20", spatial_sci),
        ("图A2-11", spatial_sci),
        ("图A3-4", spatial),
    ]:
        if replace_single_figure(doc, prefix, image, 16.1, spatial_note):
            results["spatial"] += 1
    results["profiles"] = replace_profile_figures(doc)
    doc.save(REPORT)
    return results


def validate_report() -> dict[str, int | bool]:
    doc = Document(REPORT)
    profile_caps = [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip().startswith("图A4-") and "街道断面分析" in p.text
    ]
    a5_tables = [table for table in doc.tables if len(table.rows) == 15 and len(table.columns) == 4]
    a5_drawings = sum(table._tbl.xml.count("<w:drawing") for table in a5_tables)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "profile_captions": len(profile_caps),
        "a5_grid_tables": len(a5_tables),
        "a5_grid_drawings": a5_drawings,
        "has_bad_question_marks": "????????" in all_text,
        "has_replacement_char": "�" in all_text,
    }


def image_health(paths: list[Path]) -> list[str]:
    rows = []
    for path in paths:
        with Image.open(path) as img:
            rows.append(f"{path.relative_to(ROOT)} | {img.size[0]}x{img.size[1]} | {path.stat().st_size} bytes")
    return rows


def main() -> None:
    framework_targets = [
        PAPER_FIG_DIR / "fig1_framework.png",
        CONFERENCE_FIG_DIR / "fig1_framework.png",
        LEGACY_CONFERENCE_FIG_DIR / "fig1_framework.png",
        OVERLEAF_FIG_DIR / "fig1_framework.png",
        OVERLEAF_FIG_DIR / "fig1_framework_sci.png",
    ]
    spatial_targets = [
        PAPER_FIG_DIR / "fig_four_index_spatial_comparison.png",
        PAPER_FIG_DIR / "fig11_four_index_comparison_sci.png",
    ]
    profile_paths = sorted(PROFILE_DIR.glob("PROFILE_*.png"))
    backup_root = backup_inputs(framework_targets + spatial_targets, profile_paths)

    print(f"Font: {FONT_FAMILY}")
    print(f"Backup: {backup_root}")

    draw_framework(PAPER_FIG_DIR / "fig1_framework.png", "Paper版 / 正文研究框架")
    draw_framework(CONFERENCE_FIG_DIR / "fig1_framework.png", "Conference版 / 技术路线图")
    for dst in [LEGACY_CONFERENCE_FIG_DIR / "fig1_framework.png", OVERLEAF_FIG_DIR / "fig1_framework.png"]:
        ensure_parent(dst)
        shutil.copy2(CONFERENCE_FIG_DIR / "fig1_framework.png", dst)
    draw_framework(OVERLEAF_FIG_DIR / "fig1_framework_sci.png", "SCI版 / 多源证据链与指标体系")

    draw_spatial_comparison(PAPER_FIG_DIR / "fig_four_index_spatial_comparison.png", sci_style=False)
    draw_spatial_comparison(PAPER_FIG_DIR / "fig11_four_index_comparison_sci.png", sci_style=True)

    generated_profiles = draw_all_profiles()
    replace_results = replace_report_images()
    validation = validate_report()

    sample_paths = [
        PAPER_FIG_DIR / "fig1_framework.png",
        PAPER_FIG_DIR / "fig_four_index_spatial_comparison.png",
        generated_profiles[0],
        generated_profiles[-1],
    ]
    print("Replaced:", replace_results)
    print("Validation:", validation)
    print("Generated profile images:", len(generated_profiles))
    print("Image health:")
    for row in image_health(sample_paths):
        print("  " + row)


if __name__ == "__main__":
    main()
