# -*- coding: utf-8 -*-
import sys, re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(r'E:\xicha gis 智能定位\报告_updated.docx')
print(f"清理前段落数: {len(doc.paragraphs)}")

# 找出所有要处理的段落
paragraphs = doc.paragraphs

i = 0
while i < len(paragraphs):
    p = paragraphs[i]
    txt = p.text.strip()
    chinese = sum(1 for c in txt if '\u4e00' <= c <= '\u9fff')
    # 清理段落[1] "标题" -> 实际标题
    if txt == '标题':
        for run in p.runs:
            if run.text.strip() == '标题':
                run.text = ''
        # 找下一个实际标题
        i += 1
        while i < len(paragraphs):
            next_p = paragraphs[i]
            next_txt = next_p.text.strip()
            if '可达性幻觉' in next_txt and chinese > 10:
                for run in next_p.runs:
                    pass
                break
            i += 1
        i += 1
        continue

    # 删除乱码段落（含大量不可见控制字符的）
    if chinese == 0 and len(txt) > 0:
        # 统计可读字符
        readable = sum(1 for c in txt if 32 <= ord(c) <= 126 or '\u4e00' <= c <= '\u9fff')
        if readable < len(txt) * 0.3 or (len(txt) > 20 and readable < 5):
            p._element.getparent().remove(p._element)
            i += 1
            continue

    # 删除哈尔滨工业大学工学博士学位论文乱码行
    if '哈尔滨工业大学工学博士学位论文' in txt:
        p._element.getparent().remove(p._element)
        i += 1
        continue

    # 删除"结论"段落中混入的参考文献条目（[3][4][5]）
    if re.match(r'\[\d+\]\s', txt) and i > 50:
        p._element.getparent().remove(p._element)
        i += 1
        continue

    i += 1

print(f"清理后段落数: {len(doc.paragraphs)}")

# 现在重新组织结构：在第4章"结构性成因"（重复标题）前插入跨区对比
# 先找第4章的位置
section4_idx = None
for i, p in enumerate(doc.paragraphs):
    if '结构性成因' in p.text and 'Heading' in p.style.name:
        section4_idx = i
        break

if section4_idx is not None:
    print(f"找到第4章在段落 {section4_idx}")
    print(f"  标题: {doc.paragraphs[section4_idx].text}")

# 重新读取清理后的段落
paras_data = []
for p in doc.paragraphs:
    paras_data.append({'text': p.text, 'style': p.style.name, 'element': p._element})

# 在第4章标题前插入跨区对比内容
# 找第4章段落（结构性成因后的那个重复标题）
target_idx = None
for i, pd in enumerate(paras_data):
    if '结构性成因' in pd['text'] and 'Heading' in pd['style']:
        target_idx = i
        break

if target_idx is not None:
    print(f"在段落 {target_idx} 前插入跨区对比")

    from docx.oxml import OxmlElement as OE
    from docx.oxml.ns import qn as Q

    def make_heading_para(doc, text, lvl=2):
        p = doc.add_paragraph()
        p.style = doc.styles[f'Heading {lvl}']
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.bold = True
        run._r.rPr.rFonts.set(Q('w:eastAsia'), '黑体')
        return p

    def make_para(doc, text):
        p = doc.add_paragraph(text)
        return p

    # 构建跨区对比内容
    cross_content = [
        ('h', '4. 跨区对比分析：南山、宝安、福田与龙华', 1),
        ('h', '4.1 跨区对比研究设计', 2),
        ('p', '为检验"高密度中心区是唯一可达性幻觉风险源"的假设，本研究将街景数据采集范围从南山区扩展至宝安区（西乡/航城/新安，58个样本）、福田区（香蜜湖/莲花/沙头，48个样本）和龙华区（民治/大浪，48个样本），形成对照实验设计。此外，采用DeepLabV3+语义分割（294个样本）评估建成环境语义构成。'),
        ('h', '4.2 跨区街景语义结果', 2),
        ('p', 'YOLO障碍检测与DeepLabV3语义分割揭示了各区在视觉障碍与建成环境结构上的显著差异：'),
        ('p', '南山（障碍评分8.45，绿视率9.8%）：高POI密度与封闭社区/铁路/河流阻隔并存，障碍来源多元。'),
        ('p', '宝安（障碍评分8.26，天空开敞率40.2%）：障碍来源以机场高速和产业大街区为主——即天空开敞，"可远眺但不可达"的现象突出。'),
        ('p', '福田（障碍评分3.62，人行空间识别占比最高）：障碍评分最低，人行空间连续性最好，步行设施配建率高。'),
        ('p', '龙华（障碍评分2.86，建筑界面14.7%）：新城区的视觉阻隔最低，但服务成熟度与夜间可用性仍需持续检验。'),
        ('h', '4.3 机制解释', 2),
        ('p', '采用反事实假设法解释跨区幻觉差异：'),
        ('p', 'H1（南山vs宝安）：控制服务密度/SAI相同后，宝安的AII可能更负。机制在于宝安机场快速路和产业大街区的绕行成本高于南山城中村巷道的通透效益。'),
        ('p', 'H2（南山vs福田）：控制POI密度与人口密度相同后，福田的AII预期更接近0。机制在于福田的步行空间连续性和人行设施密度优于南山。'),
        ('p', 'H3（南山vs龙华）：龙华的障碍评分最低，但其低成熟度意味着POI密度本身不足，幻觉的主要来源是"设施不存在"而非"设施不可达"。'),
        ('h', '4.4 跨区对比核心结论', 2),
        ('p', '四区对比支持以下机制性结论：并非越中心的城区幻觉越强。南山高POI密度部分补偿了路网阻隔，而宝安的低密度服务与高障碍结合产生了强幻觉。街景视觉阻隔强度与AII并非简单线性关系——天空开敞率高不等于步行可达性好。城中村密集巷道在跨区比较中仍显示出独特的步行效率优势，这一"负资产"中蕴含的"正外部性"值得在城中村更新中审慎对待。'),
    ]

    # 在第4章结构性成因前插入
    target_element = paras_data[target_idx]['element']
    parent = target_element.getparent()
    insert_pos = list(parent).index(target_element)

    for item in reversed(cross_content):
        if item[0] == 'h':
            _, text, lvl = item
            new_p = OE('w:p')
            new_pPr = OE('w:pPr')
            new_pStyle = OE('w:pStyle')
            new_pStyle.set(Q('w:val'), f'Heading{lvl}')
            new_pPr.append(new_pStyle)
            new_r = OE('w:r')
            new_rPr = OE('w:rPr')
            new_rFonts = OE('w:rFonts')
            new_rFonts.set(Q('w:eastAsia'), '黑体')
            new_rFonts.set(Q('w:ascii'), '黑体')
            new_rPr.append(new_rFonts)
            new_b = OE('w:b')
            new_rPr.append(new_b)
            new_sz = OE('w:sz')
            new_sz.set(Q('w:val'), str(28 if lvl == 1 else (24 if lvl == 2 else 22)))
            new_rPr.append(new_sz)
            new_t = OE('w:t')
            new_t.text = text
            new_r.append(new_rPr)
            new_r.append(new_t)
            new_p.append(new_pPr)
            new_p.append(new_r)
        else:
            _, text = item
            new_p = OE('w:p')
            new_r = OE('w:r')
            new_rPr = OE('w:rPr')
            new_rFonts = OE('w:rFonts')
            new_rFonts.set(Q('w:eastAsia'), '宋体')
            new_rFonts.set(Q('w:ascii'), '宋体')
            new_rPr.append(new_rFonts)
            new_sz = OE('w:sz')
            new_sz.set(Q('w:val'), '24')
            new_rPr.append(new_sz)
            new_t = OE('w:t')
            new_t.text = text
            new_r.append(new_rPr)
            new_r.append(new_t)
            new_p.append(new_r)
        parent.insert(insert_pos, new_p)

    print("跨区对比内容已插入")

# 重命名标题段落
for p in doc.paragraphs:
    if p.text.strip() == '标题':
        # 清空标题段落
        for run in p.runs:
            run.text = ''
        # 删除此空段落
        p._element.getparent().remove(p._element)
        break

doc.save(r'E:\xicha gis 智能定位\报告_updated.docx')
print("已保存")

# 打印最终结构
print("\n=== 最终文档结构 ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt:
        style = p.style.name
        prefix = '[H]' if 'Heading' in style else '[P]'
        print(f'{prefix} {i:3d} | {txt[:120]}')
