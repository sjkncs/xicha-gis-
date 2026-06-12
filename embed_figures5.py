# -*- coding: utf-8 -*-
"""Embed figures into the final report - document.add_picture + element move"""
import sys, os
sys.path.insert(0, r'E:\xicha gis 智能定位')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

BASE = r'E:\xicha gis 智能定位'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

FIGURES = [
    ('（图1：综合可达性分布图）', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig3_spatial.png'),
    ('（图2：街景障碍物分布图）', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig7_saii_analysis.png'),
    ('（图3：三类幻觉维度示意图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig1_framework.png'),
    ('（图5：研究框架技术路线图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig1_framework.png'),
    ('（图6：时间贫困指数空间聚类图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig6_time_poverty.png'),
    ('（图7：综合可达性与幻觉指数双变量地图）', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig3_spatial.png'),
    ('（图8：四象限散点图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig4_illusion_scatter.png'),
    ('（图9：弱势群体剥夺热点图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig6_deprived_communities.png'),
    ('（图10：步行环境四维评分雷达图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig5_radar.png'),
    ('（图11：供需匹配三维框架图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig9_supply_demand.png'),
    ('（图12：时间贫困与幻觉指数相关性图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig6_time_poverty.png'),
    ('（图13：昼夜达标率对比图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig8_day_night.png'),
    ('（图14：四区障碍评分对比箱线图）', 'projects\\15min-urban-accessibility\\v2_real_data\\p8_fig7_saii_analysis.png'),
    ('（图15：四区对比机制路径图）', 'projects\\15min-urban-accessibility\\conference_paper\\figures\\fig5_type_analysis.png'),
    ('（图16：治理建议实施路径图）', 'projects\\15min-urban-accessibility\\paper\\figures\\fig1_framework.png'),
]

def set_font(run, name='宋体', size=10, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.italic = italic

def add_figure_paragraphs(doc, img_relpath, caption, width=Inches(5.5)):
    """Add image + caption paragraphs to end of doc, returns the elements"""
    img_abspath = os.path.join(BASE, img_relpath)
    if not os.path.exists(img_abspath):
        return None

    elems = []
    cap_num = caption.replace('（', '').replace('）', '')

    # Add picture to document (appends to end)
    inline_shape = doc.add_picture(img_abspath, width=width)
    shape_elem = inline_shape._inline
    # Wrap shape in a paragraph element
    p_img_elem = etree.Element(f'{{{W_NS}}}p')
    r_img_elem = etree.SubElement(p_img_elem, f'{{{W_NS}}}r')
    r_img_elem.append(shape_elem)
    elems.append(p_img_elem)

    # Caption paragraph
    p_cap_elem = etree.Element(f'{{{W_NS}}}p')
    pPr = etree.SubElement(p_cap_elem, f'{{{W_NS}}}pPr')
    jc = etree.SubElement(pPr, f'{{{W_NS}}}jc', attrib={'val': 'center'})
    r_cap = etree.SubElement(p_cap_elem, f'{{{W_NS}}}r')
    rPr = etree.SubElement(r_cap, f'{{{W_NS}}}rPr')
    rFonts = etree.SubElement(rPr, f'{{{W_NS}}}rFonts', attrib={'w:eastAsia': '宋体', 'w:ascii': '宋体'})
    sz = etree.SubElement(rPr, f'{{{W_NS}}}sz', attrib={'val': '20'})
    i_elem = etree.SubElement(rPr, f'{{{W_NS}}}i')
    t = etree.SubElement(r_cap, f'{{{W_NS}}}t')
    t.text = f'图 {cap_num}'
    elems.append(p_cap_elem)

    return elems

def insert_elems_after_para(doc, para_idx, elems):
    """Insert XML elements after paragraph at para_idx"""
    body = doc.element.body
    children = list(body)
    if para_idx < len(children):
        ref_elem = children[para_idx]
        for i, elem in enumerate(elems):
            body.insert(list(body).index(ref_elem) + 1 + i, elem)

def embed_figures(input_path, output_path):
    doc = Document(input_path)
    body = doc.element.body
    total = 0
    skipped = 0

    # Collect (para_idx, elems) pairs
    insertions = []
    for i, para in enumerate(doc.paragraphs):
        for caption, img_rel in FIGURES:
            if caption in para.text:
                elems = add_figure_paragraphs(doc, img_rel, caption)
                if elems:
                    insertions.append((i, elems, caption))
                    print(f"  Found [{i}]: {caption}")
                else:
                    print(f"  SKIP [{i}]: {caption} (not found)")
                break

    # Insert in reverse order to maintain indices
    for i, elems, caption in reversed(insertions):
        insert_elems_after_para(doc, i, elems)
        total += 1

    print(f"\nFigures embedded: {total}")
    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == '__main__':
    inp = os.path.join(BASE, '报告_final.docx')
    out = os.path.join(BASE, '报告_final_含图.docx')
    embed_figures(inp, out)
