# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'E:\xicha gis 智能定位\报告_updated.docx')

# 精确删除含有乱码字符串的段落
target = '㻴㻶㻸㻺㻼㽀㾎㾐㾒㾔㾖㾶'
removed = 0
for p in doc.paragraphs:
    txt = p.text
    if target in txt:
        p._element.getparent().remove(p._element)
        removed += 1
        print(f"已删除: {txt[:60]}")

print(f"共删除 {removed} 个段落")

doc.save(r'E:\xicha gis 智能定位\报告_updated.docx')
print("已保存")

# 最终验证
print("\n=== 最终结构（全部段落）===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    style = p.style.name
    prefix = '[H]' if 'Heading' in style else '[P]'
    if txt:
        print(f'{prefix} {i:3d} | {txt[:120]}')
