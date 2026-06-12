#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
修复报告附录A.5中的街景标注图像
将 annotated_cn 目录下的60张图片以4列网格布局正确插入
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os
import glob

def add_chinese_font(run, font_name='微软雅黑'):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def find_appendix_section(doc, target_text='A.5'):
    """找到附录A.5的位置"""
    for i, para in enumerate(doc.paragraphs):
        if target_text in para.text and '街景标注图像样本' in para.text:
            return i
    return None

def remove_content_until_next_section(doc, start_idx):
    """删除从附录A.5标题到下一个附录标题之间的所有内容"""
    paragraphs_to_remove = []
    
    # 找到需要删除的段落范围
    for i in range(start_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        # 如果遇到下一个附录标题（如 A.6, B.1 等），停止
        if para.style.name.startswith('Heading') or \
           (para.text.strip().startswith('A.') and para.text.strip()[2].isdigit()) or \
           (para.text.strip().startswith('B.') and para.text.strip()[2].isdigit()):
            break
        paragraphs_to_remove.append(para)
    
    # 删除段落
    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)
    
    return len(paragraphs_to_remove)

def insert_images_in_grid(doc, insert_after_idx, image_paths, cols=4):
    """
    在指定位置后插入图片网格
    
    Args:
        doc: Document对象
        insert_after_idx: 插入位置的段落索引
        image_paths: 图片路径列表
        cols: 每行的列数
    """
    # 计算行数
    num_images = len(image_paths)
    rows = (num_images + cols - 1) // cols
    
    # 在指定段落后插入新段落作为表格的容器
    insert_position = doc.paragraphs[insert_after_idx]._element
    
    # 创建表格
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # 将表格移动到正确的位置
    insert_position.addnext(table._element)
    
    # 设置表格属性
    for row in table.rows:
        row.height = Inches(1.8)
    
    # 填充图片
    for idx, img_path in enumerate(image_paths):
        row_idx = idx // cols
        col_idx = idx % cols
        cell = table.rows[row_idx].cells[col_idx]
        
        # 清空单元格
        cell.text = ''
        
        # 添加图片
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        try:
            run = paragraph.add_run()
            run.add_picture(img_path, width=Inches(1.5))
            
            # 添加图片文件名作为说明（可选）
            # para_caption = cell.add_paragraph()
            # para_caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # filename = os.path.basename(img_path)
            # run_caption = para_caption.add_run(filename)
            # run_caption.font.size = Pt(7)
            # add_chinese_font(run_caption)
            
        except Exception as e:
            print(f"无法插入图片 {img_path}: {e}")
    
    print(f"成功插入 {num_images} 张图片，布局为 {rows} 行 × {cols} 列")

def main():
    # 路径配置
    doc_path = r'e:\xicha gis 智能定位\报告_final.docx'
    annotated_dir = r'e:\xicha gis 智能定位\自选年份\gpu_scripts\results\annotated_cn'
    output_path = r'e:\xicha gis 智能定位\报告_final_fixed.docx'
    
    print("正在加载文档...")
    doc = Document(doc_path)
    
    # 找到附录A.5的位置
    print("正在查找附录A.5...")
    appendix_idx = find_appendix_section(doc, 'A.5')
    
    if appendix_idx is None:
        print("错误：未找到附录A.5标题")
        return
    
    print(f"找到附录A.5，位于段落 {appendix_idx}")
    
    # 删除附录A.5下的旧内容
    print("正在删除旧内容...")
    removed_count = remove_content_until_next_section(doc, appendix_idx)
    print(f"已删除 {removed_count} 个段落")
    
    # 获取所有标注图片（按文件名排序）
    image_files = sorted(glob.glob(os.path.join(annotated_dir, '*_cn.jpg')))
    print(f"找到 {len(image_files)} 张标注图片")
    
    if len(image_files) == 0:
        print("错误：未找到标注图片")
        return
    
    # 插入图片网格
    print("正在插入图片网格...")
    insert_images_in_grid(doc, appendix_idx, image_files, cols=4)
    
    # 保存文档
    print(f"正在保存文档到: {output_path}")
    doc.save(output_path)
    print("完成！")
    print(f"\n请检查文件: {output_path}")

if __name__ == '__main__':
    main()
